#!/usr/bin/env python
"""
Usage: python monsters.py

# Sources:
# ezkajii's Monster Compendium.xls for most of the monster data
# highmage's dnd35.db for psionic powers
# dndtools' dnd.sqlite for most of the spell data

# CreatureInfo.csv from HeroForge on GitHub has many fewer monsters, does not have eg tirbana, probably not worth it
"""

from __future__ import print_function
import os,sys,os.path
import shutil
import itertools
import collections
import argparse
#os.chdir('xlrd-1.0.0')
sys.path.append(os.path.join(os.getcwd(), 'xlrd-1.0.0') )
#print(sys.path)
import re
import datetime
import time
import math
import sqlite3
import csv
import xlrd # https://github.com/python-excel/xlrd
# https://www.blog.pythonlibrary.org/2014/04/30/reading-excel-spreadsheets-with-python-and-xlrd/
#import ipdb
import cProfile,pstats
import time

import urllib.request, zipfile
import requests
import defusedxml.ElementTree
import defusedxml.sax
import ast
import logging
# These two lines enable debugging at httplib level (requests->urllib3->http.client)
# You will see the REQUEST, including HEADERS and DATA, and RESPONSE with HEADERS but without DATA.
# The only thing missing will be the response.body which is not logged.
try:
    import http.client as http_client
except ImportError:
    # Python 2
    import httplib as http_client
http_client.HTTPConnection.debuglevel = 1
import pycurl

import pandas

#import docx

rulebook_abbreviations = {'MM1':'Monster Manual v.3.5', 'MMI': 'Monster Manual v.3.5', 'MM':'Monster Manual v.3.5',
 'Planar':'Planar Handbook', 'PlH':'Planar Handbook',
 'ToHS':'Towers of High Sorcery',
 'CotSQ':'City of the Spider Queen', 'CSQ':'City of the Spider Queen',
 'WD':'City of Splendors: Waterdeep', 'CSW':'City of Splendors: Waterdeep',
 'City': 'CityScape', #'Ci': 'CityScape',
 'Sharn': 'Sharn - City of Towers', 'SCoT': 'Sharn City of Towers',
 'Aldriv':"Aldriv's Revenge",
 'Exp':'Expedition to the Demonweb Pits', 'EDP':'Expedition to the Demonweb Pits',
 'ExRav': 'Expedition to Castle Ravenloft', 'ExGre': 'Expedition to the Ruins of Greyhawk',
 'ExUn':'Expedition to Undermountain',
 'Dunge':'Dungeonscape', 'Du':'Dungeonscape',
 'PHB': "Player's Handbook v.3.5",
 'DMG': "Dungeon Master's Guide v.3.5", 'DMG2': "Dungeon Master's Guide II",
 'Psi':'Psionics Handbook (Web Enhancement)',
 'CoV':'Champions of Valor',
 'Loona':'Loona, Port of Intrigue',
 'ToH':'Tomb of Horror',
 'MMII': 'Monster Manual II', #'MM2':'Monster Manual 2',
 #'MM3':'Monster Manual 3',
 'MMIV': 'Monster Manual IV', #'MM4':'Monster Manual 4',
 'MMV': 'Monster Manual V', #'MM5':'Monster Manual 5',
 'FF':'Fiend Folio',
 'ElderE':'Elder Evils',
 'RoS':'Races of Stone', 'RotW':'Races of the Wild', 'RDr':'Races of the Dragon',
 'RoW': 'Races of the Wild',
 # Races of the Wild excerpt Chordevoc http://archive.wizards.com/default.asp?x=dnd/ex/20050204a&page=5
 'ELH':'Epic Level Handbook',
 'PoC':'Price of Courage',
 'MT':"Midnight's Terror", 'RTF':'Return to the Temple of the Frog',
 'ECS':'Eberron Campaign Setting', 'MoE': 'Magic of Eberron', 'EH': "Explorer's Handbook",
 #'SoS':'Spectre of Sorrows',
 'SpcSor': 'Spectre of Sorrows',
 'Sarlo': 'Secrets of Sarlona', 'SoS': 'Secrets of Sarlona',
 "Xen'd":"Secrets of Xen'drik", 'SX':"Secrets of Xen'drik",
 'FoW':'The Forge of War', 'FW':'The Forge of War',
 'FN': 'Five Nations', '5Nat': 'Five Nations',
 'SotAC':'Secrets of the Alubelok Coast', 'Sheep':"Sheep's Clothing",
 'A&EG':'Arms & Equipment Guide',
 'SvgSp':'Savage Species (Web Enhancement)', 'SS':'Savage Species',
 'LEoF':'Lost Empires of Faerun', 'LE':'Lost Empires of Faerun',
 'Kruk':'The Lost Tomb of Kruk-Ma-Kali', 'SD':'Stone Dead',
 'BoED':'Book of Exalted Deeds', 'FC1':'Fiendish Codex I', 'FC2': 'Fiendish Codex II',
 'HoA':'Fiendish Codex I',
 'SoCo':"Something's Cooking",
 'Draco':'Draconomicon',
 'MH':'Miniatures Handbook',
 'BoKr':'Bestiary of Krynn', 'DoK':'Dragons of Krynn',
 'MgoF':'Magic of Faerun', 'RoF':'Races of Faerun', 'MoF':'Monsters of Faerun',
 'PGTF': "Player's Guide to Faerun",
 # Magic of Faerun was updated in Player's Guide to Faerun: Crypt Spawn: Undead (augmented [previous type]); +8/+12 (for sample); LA +2. Do not recalculate attack bonus, saves, or skill points. Add darkvision 60 ft. to special qualities. Spectral Mage: Undead (augmented [previous type], incorporeal); +2/-; LA +6. Do not recalculate attack bonus, saves, or skill points; add darkvision 60 ft. to special qualities. Skills: A spectral mage gains a +8 racial bonus on Hide and Intimidate checks.
 # Monsters of Faerun was updated in Player's Guide to Faerun: Curst: Undead; +5/+8; 5 ft./5 ft.; LA +3. Revenant: Undead; +3/+5 (for sample); 5 ft./5 ft.; 5/magic; LA -. Add darkvision 60 ft. to special qualities. Replace regeneration with fast healing; add undead traits to special qualities. Fast Healing (Ex): A revenant regains lost hit points at the rate of 3 per round, except for damage dealt by fire, as long as it has at least 1 hit point. Fast healing does not restore hit points lost from starvation, thirst, or suffocation, and it does not allow the revenant to regrow or reattach lost body parts. Yuan-Ti, Tainted One: Monstrous humanoid; +3/+4 (for sample); LA +2. Use poison 1/day and polymorph 3/day; spell resistance changes to 12 + 1 per two levels; add darkvision 60 ft. to special qualities.
 'GotP':'Garden of the Plantmaster',
 'StSt':'The Standing Stone', 'StS':'The Standing Stone',
 'BoBS':'Bastion of Broken Souls', 'Forge': 'The Forge of Fury',
 'DrC':'Dragon Compendium', 'MoI':'Magic of Incarnum',
 'HoD':'Harvest of Darkness', 'EnvIm':'Environmental Impact',
 'DDen':"Dangerous Denizens - The Monsters of Tellene", 'DD':"Dangerous Denizens - The Monsters of Tellene",
 'KoD': 'Key of Destiny', 'BoVD': 'Book of Vile Darkness', #'ToB': 'Tome of Battle',
 'EPH': 'Expanded Psionics Handbook', # many available http://www.d20srd.org/indexes/psionicMonsters.htm but SRD does not appear in their lists
 #'Web': 'Web content',
 'TME': "The Mind's Eye",
 'ItDL': "Into the Dragon's Lair", 'FoN': 'Force of Nature', 'AoM': 'Age of Mortals',
 'ToM': 'Tome of Magic', 'TM': 'Tome of Magic',
 'ShSo': 'Shining South',
 'GW': 'Ghostwalk', 'Gh': 'Ghostwalk',
 'SaD': 'Stand and Deliver', 'HoB': 'Heroes of Battle',
 'C.Ps': 'Complete Psionic', 'OA': 'Oriental Adventures', #'LM': 'Libris Mortis',
 'Under': 'Underdark', 'Und': 'Underdark',
 'DotU':'Drow of the Underdark', 'DrU':'Drow of the Underdark',
 'C.Ar': 'Complete Arcane', 'CAr': 'Complete Arcane',
 'ToBV': 'The Treasure of the Black Veils', 'DS': 'Desert Sands', 'LDR': 'Lest Darkness Rise',
 'DCS': 'Dragonlance Campaign Setting',
 'HOS': 'Holy Order of the Stars',
 'MotP': 'Manual of the Planes', 'MP': 'Manual of the Planes', 'MoP': 'Manual of the Planes',
 r'A\d\d': 'Dragon Magazine Annual 00/01',
 r'\d\d\d': 'Dragon Magazine',
 'LoMys': 'Lands of Mystery',
 'DrM': 'Dragon Magic', 'DM': 'Dragon Magic',
 'CoR': 'Champions of Ruin', 'CR': 'Champions of Ruin',
 'LoM': 'Lords of Madness', 'WndW': 'The Secret of the Windswept Wall',
 'Sand': 'Sandstorm', 'Sa': 'Sandstorm',
 'Frost': 'Frostburn', 'Fro': 'Frostburn',
 'Storm':'Stormwrack', 'Sto':'Stormwrack',
 'C.War': 'Complete Warrior', 'C.Adv': 'Complete Adventurer', 'C.Sco': 'Complete Scoundrel', 'C.Mag': 'Complete Mage',
 'CW': 'Complete Warrior', 'CS': 'Complete Scoundrel',
 'S&S': 'Song & Silence',
 'SlvSk': 'The Silver Skeleton', 'RTEE': 'Return to the Temple of Elemental Evil',
 'HoH': 'Heroes of Horror', 'HH': 'Heroes of Horror',
 'TVoS': 'The Vessel of Stars',
 'Serp': 'Serpent Kingdoms', 'SK': 'Serpent Kingdoms',
 'SM': 'Silver Marches', 'BaS': 'Blood and Shadows - The Dark Elves of Tellene',
 'UE': 'Unapproachable East', 'Una': 'Unapproachable East',
 'BoKR': 'Bestiary of Krynn Revised',
 'Deities': 'Deities and Demigods', 'FRCS': 'Forgotten Realms Campaign Setting',
 'F&P': 'Faiths and Pantheons', #'FP': 'Faiths and Pantheons', # 3.0
 'MIC': 'Magic Item Compendium',
 }
# 108 total sources
 # if the abbreviation is just a number then it's a Dragon magazine number

# until get proper version numbers and dates for rulebooks, quick fix for monsters only just linear-order the relevant rulebooks
rulebook_priority = ('EPH', 'MM1', 'MM4', 'LoM', 'Sand', 'MH', 'MM2', 'C.Ps', 'OA', 'SaD', 'GotP', 'DrC', '318', '306')
# Stand and Deliver is a first level adventure set in the Kingdoms of Kalamar campaign setting. This adventure is designed for use with the revised (3.5) edition of the rules
# Garden of the Plantmaster: Campaign Resource and Adventure (Dungeons & Dragons: Kingdoms of Kalamar) Paperback - April, 2003
  # this pegs Dragon #309 as the start of 3.5ed http://www.enworld.org/forum/showthread.php?9651-DRAGON-Magazine-monster-index!
  # Dragon #309 (War, Incursion) from July 2003 was the first D&D 3.5 issue. https://rpg.stackexchange.com/questions/14892/in-what-issue-did-dragon-dungeon-magazine-transition-to-3-5e-rules
# Miniatures Handbook is 3.5 2003, Sandstorm is March 1, 2005

TERRAIN_NAMES = ('desert', 'forest', 'hill', 'mountain', 'plain', 'marsh', 'aquatic')

CLASS_ABBREVIATIONS = {'sor': 'Sorcerer', 'wiz': 'Wizard', 'clr': 'Cleric', 'drd': 'Druid', 'brd': 'Bard', 'bkgd': 'Blackguard', 'pal': 'Paladin', 'rng': 'Ranger'}



def fraction_to_negative(string):
  """All fractions of Hit Dice etc happen to be of the form 1/#,
  so they can be encoded as negative integers.
  """
  try:
    return int(string)
  except ValueError as E:
    if string[:2] == '1/':
      return -int(string[2:])
    else:
      raise

def integer_or_non(string):
  if string == '-': return None
  else: return int(string)

goodEvilToInt = {'G':1, 'E':-1, 'N':0, 'X':0}
#lawChaosToInt = {'L':1, 'C':-1, 'N':0, 'X':0}
# X really means any, like angels are any good, maybe really I should have an intersection table of monster_ids and alignments

spellNameErrors = {'resistance to energy':'Resist Energy', # Angel, Solar
  'protection from flame':'Protection from Energy (fire only)',
  'mass enlarge':'Mass Enlarge Person', # Greater Barghest
  # from 3.5 update booklet:
  'enlarge':'Enlarge Person',
  'reduce':'Reduce Person',
  'charm person or animal':'Charm Animal',
  'change self':'Disguise Self',
  'mass charm':'Charm Monster, Mass',
  'circle of doom':'Inflict Light Wounds, Mass',
  'mass haste':'Haste',
  'create/destroy water':'Create Water', # not really an error, but easier to handle blue dragon this way
  'summon djinni':'Summon Monster VII (djinni only)',
  'colossal vermin':'Giant Vermin (can increase from Large to Gargantuan and from Huge to Colossal)',
  'permanent creation (food for 2d12 people, 52 cu.ft. of soft goods, or 18 cu.ft. of wooden items; or 400bls of metal items at 10 times the duration of such items produced by major creation)':'Major Creation',
  'transmute rock to mud / mud to rock':'transmute rock to mud (or mud to rock)', # not really an error, but easier to handle copper dragon this way
  'transmute rock to lave':'transmute rock to lava',
  'yare of air':'yari of air',
  'protection from chaos/evil/good/law':'protection from evil', # not really an error, but easier to handle Grimweird
  'power word (any)':'Power Word Petrify', # sqlite> select distinct level,dnd_spell.name from dnd_spell inner join dnd_spellclasslevel on dnd_spell.id=dnd_spellclasslevel.spell_id inner join dnd_characterclass on character_class_id=dnd_characterclass.id where dnd_spell.name like "power word%" order by level;
  'symbol (any)':'Symbol of Weakness', # sqlite> select distinct level,dnd_spell.name from dnd_spell inner join dnd_spellclasslevel on dnd_spell.id=dnd_spellclasslevel.spell_id inner join dnd_characterclass on character_class_id=dnd_characterclass.id where dnd_spell.name like "symbol%" order by level;
  'locate person':'Locate Creature',
  'turn wood':'Repel Wood',
  'slow poison':'Delay Poison',
  'great sh9out':'Great Shout',
  'wood shape)':'Wood Shape',
  'invisibility\\':'invisibility',
  'se invisibility':'see invisibility',
  'domination':"Psionic Dominate",
  'mind rap (5 rds*)':"Mind Trap (5rounds)",
  'psychic thrust (6d6*)':"Psychic Crush (6d6*)", # http://www.d20srd.org/srd/psionic/monsters/neothelid.htm
  'painful touch':"Painful Strike",
  'blindness':'Blindness/Deafness (blindness only)',
  'remove blindness':'Remove Blindness/Deafness (blindness only)',
  'create grater undead':'Create Greater Undead',
  'animated dead':'Animate Dead',
  'discern lie':'Discern Lies',
  'ghost sounds':'Ghost Sound',
  'chain lighting (dmg increased as though empowered)':'Chain Lightning (Empowered)',
  'call lighting':'Call Lightning',
  'call lighting storm':'Call Lightning Storm',
  "Snilloc's snowball storm":"Snilloc's Snowball Swarm",
  "word of law":"Dictum",
  'purify food or drink':'Purify Food and Drink',
  'create food and drink':'Create Food and Water',
  'pyrotechnic':'Pyrotechnics',
  'firestorm':'Fire Storm',
  'brainlock (any nonmindless*)':'Brain Lock (any nonmindless*)',
  'stonetell':'Stone Tell',
  'cloud kill':'Cloudkill',
  'acid cloud':'Acid Fog',
  'cloud of fish':'Fish Cloud',
  'flaming shroud':'Shroud of Flame',
  'mind store':'Mind Seed',
  'far punch':'Far Hand',
  'combat prescience':'Prescience, Offensive',
  'shield of prudence':'Thought Shield', # Shield of Prudence is NOT Thought Shield, Hamaguan has both http://archive.wizards.com/default.asp?x=dnd/psb/20030523c
  'concussive detonation (9d6*)':'Concussion Blast (9d6*)', # NOT, Hammerfish has both 
  'vigilance':'Vigilant Slumber',
  'dowsing':'Locate Water',

  'chameleon power':"Chameleon",
  'endurance':"Bear's Endurance",
  "Nystul's undetectable aura":"Nystul's Magic Aura",
  'wave of fire':'Wall of Fire',
  'cause disease':'Contagion',
  'cloud trapeze (self +50lbs only)':'Wind Walk (self +50lbs only)',
  'half undead':'Halt Undead',
  'door':'Doom', # http://www.enworld.org/forum/showthread.php?491761-Frostburn-Frost-Giant-Spiritspeaker-s-spell-like-ability
  'crown of brilliant':'Crown of Brilliance',
  'deep  slumber':'Deep Slumber',
                  }
typeNameErrors = {'Exraplanar':'Extraplanar', 'Extraplanr':'Extraplanar','Extaplanar':'Extraplanar'}
def fix_subtype(subtype):
  subtype = subtype.strip()
  if subtype in typeNameErrors:
    subtype = typeNameErrors[subtype]
  return subtype

MINOR_WORDS = ['of', 'the', 'and', 'to', 'from', 'I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX']
def sensible_title(string):
  # capitalize() turns VII into Vii
  # for Blindness/Deafness, should standardize dnd_spell
  if string is None: return None
  elif string == 'blindness/deafness' or string == 'Blindness/Deafness': return 'Blindness/Deafness'
  elif string == 'remove blindness/deafness': return 'Remove Blindness/deafness'
  #print(re.split('([ -])', string))
  try:
    words = re.split('([ -])', string)
  except TypeError as E:
    print(type(string), string)
    raise
  return ''.join( (word if word in MINOR_WORDS else word.capitalize() ) for word in words)

def spell_name_to_id(curs, spellName, allowNoneResult=False):
  # more specific than id_from_name
  if spellName == '':
    raise ValueError(spellName)
  curs.execute('''SELECT published,dnd_spell.id from dnd_spell INNER JOIN dnd_rulebook on dnd_spell.rulebook_id=dnd_rulebook.id WHERE dnd_spell.name = ? ORDER BY published;''', (sensible_title(spellName),) )
  results = curs.fetchall()
  #print('spell_name_to_id', spellName, results)
  if len(results) == 0:
    words = spellName.split()
    if words[0].lower() in ('lesser', 'greater', 'mass', 'psionic'):
      return spell_name_to_id(curs, ' '.join(words[1:]) + ', ' + words[0], allowNoneResult) # MMV lashemoi really does have a special ability called Lesser Strength from Pain that isn't a spell
    else:
      if allowNoneResult:
        return None
      raise KeyError(spellName + " not found in spell database")
  spell_id = results[-1][1]
  assert spell_id is not None
  return spell_id

EDITION_PRIORITY = ('Core (3.5)', 'Supplementals (3.5)', 'Forgotten Realms (3.5)', 'Eberron (3.5)', 'Core (3.0)', 'Oriental Adventures (3.0)', 'Supplementals (3.0)', 'Forgotten Realms (3.0)', None)
def id_from_name(curs, tableName, name, allowExtraOnLeft=False, allowExtraOnRight=False, useEdition=False, additionalCriteria=tuple() ):
  """
  This function uses SQL's like instead of exact equality, to enable tricks like passing "Colossal%" to match Colossal+.
  """
  if re.match("\w+$", tableName) is None:
    raise TypeError(tableName + " does not look like a valid SQL table name.")
  if tableName == 'dnd_monstersubtype':
    name = sensible_title(name) # Living Construct wants title(), but title() turns Tanar'ri into Tanar'Ri
  if allowExtraOnLeft:
    name = '%' + name
    #curs.execute('''SELECT id from {} WHERE name like "%{}";'''.format(tableName, name) ) # unsafe, but how to do it properly?
  if allowExtraOnRight:
    name = name + '%'
  #print('''SELECT id from {} WHERE name like ?;'''.format(tableName), name)
  if allowExtraOnLeft or allowExtraOnRight:
    operator = 'LIKE'
  else:
    operator = '='
  if useEdition:
    SQLcmd = '''SELECT {0}.id, dnd_dndedition.name from {0} INNER JOIN dnd_rulebook on dnd_rulebook.id=rulebook_id LEFT JOIN dnd_dndedition on dnd_edition_id=dnd_dndedition.id WHERE {0}.name {1} ?;'''.format(tableName, operator)
    # print('useEdition', SQLcmd.replace('?', name) )
    curs.execute(SQLcmd, (name,) )
  else:
    cmd = '''SELECT id from {} WHERE name {} ?'''.format(tableName, operator)
    cmd += ''.join(' AND {}=?'.format(crit[0]) for crit in additionalCriteria)
    cmd += ';'
    curs.execute(cmd, (name,) + tuple(crit[1] for crit in additionalCriteria) )
  results = curs.fetchall()
  # rowcount does not work for SELECT statements because we cannot determine the number of rows a query produced until all rows were fetched.
  # Even after all rows have been fetched, rowcount is still -1.
  if results == []:
    return None
  #if result is None:
    #return None
    #raise IndexError("{} does not appear in the SQL table {}.".format(name, tableName) )
  if useEdition:
    for edition in EDITION_PRIORITY:
      for (ret,editionName) in results:
        if editionName == edition:
          return ret
    if len(results) != 1:
      raise RuntimeError("{} results in {} for {}: {}".format(len(results), tableName, name, results) )
    return results[0][0]
  else:
    if len(results) != 1:
      raise RuntimeError("{} results in {} for {}: {}".format(len(results), tableName, name, results) )
    return results[0][0]


def iterate_over_types(curs):
  curs.execute('''SELECT id,name FROM dnd_monstertype;''')
  return curs.fetchall()


darkvisionRE = re.compile(r"DV\d{2,3}")
damageReductionRE = re.compile(r"DR (\d{1,2})/((?:[\w\-\+]|\s(?!and [D\d]))+)(?: and (?:DR )?(\d{1,2})/([\w\s\-\+]+))?")
# match a only with negative lookahead assertion that not followed by nd?
# match \s only with negative lookahead assertion that not followed by and
spellResistanceRE = re.compile(r"SR ?(\d{1,2})")
spellsRE = re.compile(r"[Ss]pells \((\w\w\w)(\d{1,2})\)")
fastHealingRE = re.compile(r"FH(\d{1,2})")

def insert_monster_casts_spells(curs, monster_id, attack):
  """
sqlite> select distinct dnd_special_ability.name from dnd_monster inner join dnd_monstertype on dnd_monster.type_id=dnd_monstertype.id inner join monster_has_special_ability on dnd_monster.id=monster_has_special_ability.monster_id inner join dnd_special_ability on monster_has_special_ability.special_ability_id=dnd_special_ability.id inner join dnd_rulebook on rulebook_id=dnd_rulebook.id where dnd_special_ability.name like "%spells%" order by dnd_special_ability.name, dnd_monstertype.name, hit_dice;
Spells (adept3)
Spells (air shugenja5)
Spells (beguiler1)
Spells (brd7)
Spells (clr/sor/wiz6) drider
Spells (clr29 and wiz29)
Spells (evoker15)
Spells (mys3)
Spells (sor lvl = HD+4) sylph, that's wrong, only caster level increases
Spells (wu-jen10)
spells (assassin5)
spells (bkgd2) blackguard
spells (clr 20 and sor20)
spells (duskblade3)
spells (evoker5)
spells (maho-tsukai3)
spells (mys4)
spells (mystic6)
spells (necromancer13)
spells (pal11)
spells (pal4)
spells (pal6)
spells (rng4)
spells (rng6)
spells (shugenja3)
spells (sor lvl = HD+2)
spells (sor117)
spells (sor14 plus druid spells known)
spells (sor18; 1d20+22 to overcome SR)
spells (sor20 and clr20
  """
  matchObj = spellsRE.match(attack)
  if matchObj is None:
    return None
  classAbbrev = matchObj.group(1)
  if classAbbrev not in CLASS_ABBREVIATIONS:
    return None
  className = CLASS_ABBREVIATIONS[classAbbrev]
  curs.execute('''SELECT dnd_characterclass_with_spells.id FROM dnd_characterclass_with_spells WHERE dnd_characterclass_with_spells.name = ?;''', (className,))
  results = curs.fetchall()
  #if len(results) == 0:
  #  return None  # skip clr for now
  #  raise Exception(attack)
  if len(results) != 1 or len(results[0]) != 1:
    raise Exception(className)
    # curs.execute('''SELECT name FROM dnd_characterclass WHERE name LIKE ?;''', (classAbbrev + '%',))
    # raise Exception(curs.fetchall())
  class_id = results[0][0]
  level = int(matchObj.group(2))
  curs.execute('''INSERT INTO monster_casts_spells (monster_id, character_class_id, level) VALUES (?,?,?);''', (monster_id, class_id, level))

def insert_damage_reduction_by_value(curs, monster_id, value, bypass):
  value = int(value)
  if bypass == '-':
    bypass_id = None
  else:
    curs.execute('''INSERT OR IGNORE INTO damage_reduction (bypass) VALUES (?);''', (bypass,) )
    #bypass_id = curs.lastrowid # does this return None if IGNORED?
    curs.execute('''SELECT id FROM damage_reduction WHERE bypass=?;''', (bypass,) )
    #bypass_id = id_from_name(curs, 'damage_reduction', bypass)
    bypass_id = curs.fetchone()[0]
  curs.execute('''INSERT INTO monster_has_damage_reduction (monster_id, reduction, bypass_id) VALUES (?,?,?);''', (monster_id, value, bypass_id) )
def insert_damage_reduction(curs, monster_id, quality):
  matchObj = damageReductionRE.match(quality)
  if matchObj is None:
    return None
  insert_damage_reduction_by_value(curs, monster_id, matchObj.group(1), matchObj.group(2))
  if matchObj.group(3) is not None:
    assert matchObj.group(4) is not None
    insert_damage_reduction_by_value(curs, monster_id, matchObj.group(3), matchObj.group(4))
  return matchObj

def insert_fast_healing(curs, monster_id, quality):
  matchObj = fastHealingRE.match(quality)
  if matchObj is None:
    return None
  assert matchObj.group(1) is not None
  curs.execute('''INSERT INTO monster_has_fast_healing (monster_id, healing) VALUES (?,?);''', (monster_id, int(matchObj.group(1))) )
  return matchObj

def insert_spell_resistance(curs, monster_id, quality):
  matchObj = spellResistanceRE.match(quality)
  if matchObj is None:
    return None
  assert matchObj.group(1) is not None
  curs.execute('''INSERT INTO monster_has_spell_resistance (monster_id, resistance) VALUES (?,?);''', (monster_id, int(matchObj.group(1))) )
  return matchObj

casterLevelValueREstring = '(?:(?:\d{1,2})|(?:HD)|(?:lvl))\s*(?:\([\w\s]+\d?\))?'
casterLevelTagREstring = '(?:C|M|T)L=?\s?'
casterLevelREstring = casterLevelTagREstring + '(' + casterLevelValueREstring + ')'
casterLevelREstringNoncapturing = casterLevelTagREstring + '(?:' + casterLevelValueREstring + ')'
casterLevelRE = re.compile(casterLevelREstring)
matchObj = casterLevelRE.match('CL12')
assert matchObj is not None
assert matchObj.group(0) == 'CL12'
assert matchObj.group(1) == '12' # keep these, if fiddle with caster level tag RE sometimes the | catches everything to the right or somesuch
matchObj = casterLevelRE.match('CLHD (min3)')
assert matchObj is not None
assert matchObj.group(0) == 'CLHD (min3)'
matchObj = casterLevelRE.match('CL10 (with Graystaff only), at will')
assert matchObj is not None
if matchObj.group(0) != 'CL10 (with Graystaff only)':
  raise RuntimeError(matchObj.group(0) )
#if matchObj.group(1) != '10': # 10 (with Graystaff only) dunno if that's okay
#  raise RuntimeError(matchObj.group(1) )

# sometimes due to typo see 0 instead of right-paren
movementModeRE = re.compile(r'([bcfs])(\d{1,3})(?: ?\(([a-v]{2,4})[\)0])?') # prf, gd, avg, pr, clu

# put (?:) on all substrings just in case
dieRollREstring = r'(?:\d{1,2}d\d{1,2})'
#dieRollRE = re.compile(dieRollREstring)
dieRollOrConstantREstring = r'(?:' + dieRollREstring + '|\d)'
# When you use |, Python re will always try to match the first listed regex, and only if that fails will it match the others.
# So when we had \d|, it was always stopping at the single digit.
# If we want to capture number of dice and die size, we'll have to deal with the fact that sometimes we have constants (just one number instead of two).
damageBonusREstring = r'(?:[+-]\d{1,2})'
damageBonusCapturingREstring = r'([+-]\d{1,2})' # capturing instead of noncapturing parens
criticalRangeREstring = r'(?:/1[5-9]\-20)'
criticalMultiplierREstring = r'(?:/x[2-4])'
criticalHitCommentREstring = r'(?:\, vorpal)'
damageTypeREstring = r'( (?!plus)[A-Za-z\- ]+)' # 1d4 Wis drain, 1 fire, +1 merciful greataxe +8 (3d6+5/x3 or 4d6+5/x3 non-lethal)
nonNumericalDamageREstring = r"([A-Za-zç' ]+)" # could be eg entangle or Kyuss's gift
numericalDamageREstring = (r'(?:' + dieRollOrConstantREstring + damageBonusREstring + '?' + criticalRangeREstring + '?' + criticalMultiplierREstring + '?' +
                           criticalHitCommentREstring + '?' +
                           ')')
numericalDamageRE = re.compile(r'(?:' + '(' + dieRollOrConstantREstring + ')' + damageBonusCapturingREstring + '?' + criticalRangeREstring + '?' +
                           criticalMultiplierREstring + '?' +
                           criticalHitCommentREstring + '?' +
                           ')')
assert len(numericalDamageRE.match("1d4+2").group(0)) == 5
assert len(numericalDamageRE.match('2d10+11/19-20').group(0)) == 13
singleDamageREstring = r'(?:' + numericalDamageREstring + damageTypeREstring + '?' + '|' + nonNumericalDamageREstring + ')'
#singleDamageRE = re.compile(singleDamageREstring)
damageREstring = r'(' + singleDamageREstring + r'(?: (?:(?:plus)|(?:or)|(?:and)) ' + singleDamageREstring + r')*)' # capturing instead of noncapturing parens
# We cannot use (?p<damage> to name this because we are going to include this multiple times; instead we use SINGLE_ATTACK_DAMAGE_GROUP_NUMBER.
# Inevitable, Marut Slam +22 (2d6+12 plus 3d6 sonic or 3d6 electricity)
assert re.match(damageREstring, "2d6+12 plus 3d6 sonic or 3d6 electricity")
assert re.match(damageREstring, "1d3 plus 1 Con")
#re.compile(damageREstring)
parenthesizedDamageREstring = r'(?: +\(' + damageREstring + r'\))'
#parenthesizedDamageRE = re.compile(parenthesizedDamageREstring)
# Gargantuan +3 adamantine warhammer +37 (4d6+27/x3) or +3 javelin +22 (2d6+19) or slam +34 (1d8+16)
numberOfAttacksREstring = r'(\d{1,3} )' # Hydra, 10-Headed 10 bites +14 (1d10+5) # Hecatoncheires 100 greatswords +73
# when spaces are needed to be sure a bit is separate, I arbitrarily assign the spaces to the end of each REstring
weaponSizeREstring = r'(?:(?:(?:Tiny)|(?:Small)|(?:Large)|(?:Huge)|(?:Gargantuan)) )'
weaponEnchantmentREstring = r'(?:\+\d{1,2} )' # Gloom +10 keen dagger of human dread +54 (1d4+21/15-20)
possiblyHyphenatedWordREstring = r"(?:[A-Za-z]+(?:\-[a-z]+)?(?:'s)?)" # add 's to make possessive
attackNameREstring = r'(' + possiblyHyphenatedWordREstring + '(?: ' + possiblyHyphenatedWordREstring + ')*' + ')' # capturing instead of noncapturing parens
compositeBowStrBonusREstring = r'(?: \(\+\d{1,2} Str bonus\))'
# \+|\- doesn't work, possibly because needed parentheses around (A|B)
attackBonusREstring = r'([+-]\d{1,3})' # capturing instead of noncapturing parens
attackRollTypeREstring = r'(?:(?: ranged)?(?: touch)?)' # r'(?: (?:ranged )?touch)'
singleAttackModeREstring = r'(?:' + numberOfAttacksREstring + '?' + weaponSizeREstring + '?' + weaponEnchantmentREstring + '?' + attackNameREstring + compositeBowStrBonusREstring + '?' + ' ' + attackBonusREstring + attackRollTypeREstring + '?' + parenthesizedDamageREstring + '?' + ')'
attacksREstring = '(' + singleAttackModeREstring + ')' + r'(?: or (' + singleAttackModeREstring + r'))*'
#noAttacksREstring = r'(?:\-)' # not much point in allowing this, since cannot extract an attack bonus or anything, will still have to specifically check for the case of no attacks
#print('attacksREstring =', attacksREstring)
singleAttackModeRE = re.compile(singleAttackModeREstring)
attacksRE = re.compile(attacksREstring)
#attackWithoutParensRE = re.compile(attackNameREstring + compositeBowStrBonusREstring + '?' + ' ' + attackBonusREstring + attackRollTypeREstring + '? ' + damageREstring)

assert re.match('(\w)(?: or (\w))*', 'a or b or c or d').groups() == ('a', 'd')
# This doesn't seem like it can possibly be intended behavior.
# If normally-capturing parentheses don't capture when they're wrapped in repetition qualifiers (*, +, ?, {m,n}, etc), then, fine. That would certainly keep things simple and easy to use, even if it would prevent some useful features.
# But why would it capture *one* instance of the repetition, yet not the others? And even if it was going to do that, why the last rather than the first?

SINGLE_ATTACK_DAMAGE_GROUP_NUMBER = 4
# Consider Bite +7 (2d6+7 plus 1d4 acid).
# There are two instances of damage, with different group numbers.
# We could isolate the damage string inside the parentheses and then finditer with singleDamageREstring, but singleDamageREstring will match almost everything, including e.g. "acid" by itself.
# We could iterate over the groups of the damage match object and individually check each one for matching singleDamageREstring, but again, singleDamageREstring will match almost anything, including "acid" by itself, which will get misread as a non-numerical damage instead of a damage type.
SINGLE_ATTACK_DAMAGE_TYPE_GROUP_NUMBER = 5
SINGLE_ATTACK_NON_NUMERICAL_DAMAGE_GROUP_NUMBER = 8

assert attacksRE.match('Short sword +2 (1d6/19-20) or touch +2 (1d4 Int)')
matches = list(singleAttackModeRE.finditer('Short sword +2 (1d6/19-20) or touch +2 (1d4 Int)'))
assert len(matches) == 2
assert matches[1].group(2) == 'or touch' # probably not worth fixing now, but note
assert matches[0].group(SINGLE_ATTACK_DAMAGE_GROUP_NUMBER) == '1d6/19-20'
assert matches[1].group(SINGLE_ATTACK_DAMAGE_GROUP_NUMBER) == '1d4 Int'
assert matches[1].group(SINGLE_ATTACK_DAMAGE_TYPE_GROUP_NUMBER) == ' Int'
matchObj = singleAttackModeRE.match('rock +9 ranged (2d6+8)')
assert matchObj.group(SINGLE_ATTACK_DAMAGE_GROUP_NUMBER) == '2d6+8'

assert singleAttackModeRE.match('+2 composite longbow (+5 Str bonus) +28 (2d6+7/x3 plus slaying)').group(SINGLE_ATTACK_DAMAGE_GROUP_NUMBER) == '2d6+7/x3 plus slaying'
print(singleAttackModeRE.match('+2 composite longbow (+5 Str bonus) +28 (2d6+7/x3 plus slaying)').groups())
assert singleAttackModeRE.match('+2 composite longbow (+5 Str bonus) +28 (2d6+7/x3 plus slaying)').group(SINGLE_ATTACK_NON_NUMERICAL_DAMAGE_GROUP_NUMBER) == 'slaying'
assert singleAttackModeRE.match('Gargantuan +3 adamantine warhammer +37 (4d6+27/x3)').group(SINGLE_ATTACK_DAMAGE_GROUP_NUMBER) == '4d6+27/x3'
assert re.match('(?:Large) ', 'Large ').group(0) == 'Large '
assert re.match(r'(?:(?:Large) )', 'Large ').group(0) == 'Large '
assert re.match(r'(?:(?:Gargantuan) )', 'Gargantuan ').group(0) == 'Gargantuan '
assert re.match(weaponSizeREstring + '?', 'Gargantuan').group(0) == '' # cannot match because it didn't find the trailing space
assert re.match(weaponSizeREstring + '?', 'Gargantuan ').group(0) == 'Gargantuan '
assert re.match(weaponSizeREstring + '?', 'Large').group(0) == ''
assert re.match(weaponSizeREstring + '?', 'Large ').group(0) == 'Large '
assert re.match(weaponSizeREstring + '?' + weaponEnchantmentREstring + '?', 'Large +3 ').group(0) == 'Large +3 '
assert singleAttackModeRE.match('Large +3 adamantine warhammer +37 (4d6+27/x3)').group(SINGLE_ATTACK_DAMAGE_GROUP_NUMBER) == '4d6+27/x3'
assert singleAttackModeRE.match('2 composite longbows +14 (2d6+6/x3)').group(SINGLE_ATTACK_DAMAGE_GROUP_NUMBER) == '2d6+6/x3'
assert singleAttackModeRE.match('2 composite longbows (+5 Str bonus) +14 (2d6+6/x3)').group(SINGLE_ATTACK_DAMAGE_GROUP_NUMBER) == '2d6+6/x3'
assert singleAttackModeRE.match('2 +1 composite longbows (+5 Str bonus) +14 (2d6+6/x3)').group(SINGLE_ATTACK_DAMAGE_GROUP_NUMBER) == '2d6+6/x3'
assert singleAttackModeRE.match('2 Large composite longbows (+5 Str bonus) +14 (2d6+6/x3)').group(SINGLE_ATTACK_DAMAGE_GROUP_NUMBER) == '2d6+6/x3'
print(singleAttackModeRE.match('Large +1 longbow +14 (2d6)').groups())
assert re.match('a?a?a', 'aaa').group(0) == 'aaa' # re tries to match as much as possible
assert singleAttackModeRE.match('Gargantuan +1 longbow +14 (2d6)').group(SINGLE_ATTACK_DAMAGE_GROUP_NUMBER) == '2d6'
assert singleAttackModeRE.match('Large +3 adamantine warhammer +14 (2d6)').group(SINGLE_ATTACK_DAMAGE_GROUP_NUMBER) == '2d6'
assert singleAttackModeRE.match('Large +1 longbow +14 (2d6)').group(SINGLE_ATTACK_DAMAGE_GROUP_NUMBER) == '2d6'
assert singleAttackModeRE.match('2 Large +1 composite longbows +14 (2d6)').group(SINGLE_ATTACK_DAMAGE_GROUP_NUMBER) == '2d6'
assert singleAttackModeRE.match('2 Large +1 composite longbows (+5 Str bonus) +14 (2d6+6/x3)').group(SINGLE_ATTACK_DAMAGE_GROUP_NUMBER) == '2d6+6/x3'
assert singleAttackModeRE.match('+5 holy flaming longsword +44 (2d8+18/17-20 plus 2d6 holy and 1d6 fire)').group(SINGLE_ATTACK_DAMAGE_GROUP_NUMBER) == '2d8+18/17-20 plus 2d6 holy and 1d6 fire'

naturalWeaponREstring = r'([Bb]ite|[Cc]law|[Tt]alon|[Gg]ore|[Ss]lap|[Ss]lam|[Ss]ting|[Tt]entacle)s?' # capturing instead of noncapturing parens
naturalWeaponRE = re.compile(naturalWeaponREstring)
naturalWeaponAttackREstring = naturalWeaponREstring + ' ' + attackBonusREstring + attackRollTypeREstring + '?' + parenthesizedDamageREstring
naturalWeaponAttackRE = re.compile(naturalWeaponAttackREstring)
assert naturalWeaponAttackRE.match("claws +3 (1d3 plus 1 Con)")

parentheticalREstring = r'\([^)]*\)' # ( followed by any characters other than ), then )
parentheticalREstringWithGroup = r'\(([^)]*)\)' # for some unknown reason, including the group in the main causes noUnparenthesizedCommasRE to only match the interiors of parentheses
parentheticalRE = re.compile(parentheticalREstringWithGroup)
singleCharOrParenthesizedREstring = '[^,(]|' + parentheticalREstring
noUnparenthesizedCommasRE = re.compile('(?:' + singleCharOrParenthesizedREstring + ')+')
frequencyREstring = r"((?:[aA]t will(?: as a \w+ action:)?(?:\s*\(every \drds\))?)|(?:\d/day)|(?:\d/hour)|(?:\d/week)|(?:\d/tenday)|(?:\d/month)|(?:\d/year)|(?:\d/century))"
freqChangeRE = re.compile(r'(?:;|,|^)?(?:In addition,)?(?:continually active,)?\s*(?:(' + casterLevelREstringNoncapturing + "),)?\s*" + frequencyREstring + '\s*-\s*')
# planetar: CL17, at will-continual flame, dispel magic, holy smite, invisibility (self only), lesser restoration, remove curse, remove disease, remove fear, speak with dead, 3/day - blade barrier, flame strike, power word stun, raise dead, waves of fatigue; 1/day - earthquake, greater restoration, mass charm monster, waves of exhaustion. 
# emerald dragon: at will - object reading 3/day - greater invisibility
# Bheur: CL10 (with Graystaff only), at will - hold person, solid fog; 3/day - cone of cold, ice storm, wall of ice; CL10 (not Graystaff dependent), at will - chill metal, ray of frost, Snilloc's snowball storm; 1/tenday - control weather
matchObj = freqChangeRE.match("CL10 (with Graystaff only), at will - hold person; CL10 (not Graystaff dependent), at will - chill metal")
assert matchObj is not None
assert freqChangeRE.split("CL10 (with Graystaff only), at will - hold person; CL10 (not Graystaff dependent), at will - chill metal") == ['', 'CL10 (with Graystaff only)', 'at will', 'hold person', 'CL10 (not Graystaff dependent)', 'at will', 'chill metal']
assert freqChangeRE.split("CL4, At will - dancing lights, detect evil, detect magic, faerie fire; 3/day magic missile, sleep") == ['', 'CL4', 'At will', 'dancing lights, detect evil, detect magic, faerie fire; 3/day magic missile, sleep']



eightByteIntMax = 9223372036854775807 # max can store in 
assert eightByteIntMax + 1 == 2**63


def parse_comma_separated_spells(commaSeparatedSpells):
  # comma errors:
  commaSeparatedSpells = commaSeparatedSpells.replace('faerie, fire', 'faerie fire').replace('greater teleport, (self', 'greater teleport (self').replace('confusion dimension door', 'confusion, dimension door').replace('aura sight (chaos/law only) (', 'aura sight (chaos/law only, ').replace('energy arc (fire only) (', 'energy arc (fire only, ').replace('energy missile (fire only) (', 'energy missile (fire only, ').replace('energy burst (fire only) (', 'energy burst (fire only, ').replace('energy ray (electricity only) (', 'energy ray (electricity only, ').replace('energy stun (sonic only) (', 'energy stun (sonic only, ').replace('psionic daze (affects any creature type) (', 'psionic daze (affects any creature type, ')
  # Obsidian Dragon has all the energy blah powers http://archive.wizards.com/default.asp?x=dnd/psb/20030124b
  #  it really is Flaming Shroud, I dunno
  spells = [spell.strip() for spell in noUnparenthesizedCommasRE.findall(commaSeparatedSpells)]
  ret = list()
  for spell in spells:
          if spell in spellNameErrors: spell = spellNameErrors[spell]
          matchObj = parentheticalRE.search(spell)
          if matchObj is not None:
            if matchObj.end() != len(spell): raise RuntimeError(spell)
            assert spell[matchObj.start() - 1] == ' '
            parenthetical = matchObj.group(1)
            spell = spell[:matchObj.start()-1]
          else:
            parenthetical = None
          words = spell.split()
          if words[0].lower() in ('quickened', 'heightened', 'maximized', 'empowered', 'violated', 'mortalbane', 'corrupted'):
            if parenthetical is None: parenthetical = words[0]
            else: parenthetical = words[0] + ', ' + parenthetical
            spell = ' '.join(words[1:])
          ret.append( (spell, parenthetical) )
  return ret

def get_powers_from_dnd35_db(powerDBpath=os.path.join('srd35-db-SQLite-v1.3', 'dnd35.db') ):
  """http://www.enworld.org/forum/showthread.php?205248-SRD-3-5-Database-SQLite
  """
  conn = sqlite3.connect(powerDBpath)
  curs = conn.cursor()
  #curs.execute('''CREATE TEMPORARY TABLE disciplines (id INTEGER, name varchar(32) );''')
  curs.execute('''SELECT name, discipline,subdiscipline, descriptor, level, xp_cost, manifesting_time, range, target, area, effect, duration, saving_throw,power_resistance, full_text FROM power;''')
  # xp_cost from this database is the entire text after XP Cost:, which makes a certain amount of sense since it includes things like "250 XP for each point by which the object's hardness is altered."
  # A plain integer clearly wouldn't be very useful here, but I think the other database had the right idea: have a boolean for whether there's an XP cost *at all*, and leave the details in the description.
  def fix_XP(xp_cost):
    if xp_cost == 'None': return False
    else: return True
  return [(name, discipline, fix_XP(xp_cost), manifesting_time, spellrange, target, area, effect, duration, saving_throw,power_resistance, full_text) for (name, discipline,subdiscipline, descriptor, level, xp_cost, manifesting_time, spellrange, target, area, effect, duration, saving_throw,power_resistance, full_text) in curs.fetchall()]
"""
they're all from the SRD:
sqlite> select distinct reference from power;
SRD 3.5 PsionicPowersA-C
SRD 3.5 PsionicPowersD-F
SRD 3.5 PsionicPowersG-P
SRD 3.5 PsionicPowersQ-W
"""



def insert_if_needed(curs, tableName, name, **kwargs):
  """
it is NOT safe to re-insert an entry into a table:
sqlite> create table tbl (id INTEGER PRIMARY KEY NOT NULL, x nchar(1) );
sqlite> insert into tbl (x) values ('a');
sqlite> insert into tbl (x) values ('a');
sqlite> select id,x from tbl;
1|a
2|a
  INSERT OR IGNORE?
  """
  existingID = id_from_name(curs, tableName, name)
  kwargs['name'] = name
  columnNames,values = zip(*kwargs.items() )
  for col in columnNames:
    if re.match("\w+$", col) is None:
      raise TypeError(col)
  if existingID is not None:
    # Here's the problem: we do legitimately have cases where two different abilities share the same name.
    # For example, the aboleth has a slime special attack while the brine naga has a slime special quality.
    if tableName == "dnd_special_ability": return existingID # hack
    kwargs['id'] = existingID
    curs.execute('''SELECT {} FROM {} WHERE id=:id;'''.format(','.join(columnNames), tableName), kwargs)
    for i,val in enumerate(curs.fetchone() ):
      if val != values[i] and type(val) is not str or val.lower() != values[i].lower():
        raise RuntimeError("{} != {} in row {}".format(val, values[i], kwargs) )
    return existingID
  else:
    assert existingID is None
    commandString = '''INSERT INTO {} ({}) VALUES ({});'''.format(tableName, ','.join(columnNames), ','.join([':'+col for col in columnNames]) )
    #print('commandString =', commandString)
    curs.execute(commandString, kwargs)
    return curs.lastrowid



def extract_material_components(curs):
  pass

def recreate_spell_table(curs):
  curs.execute('''CREATE TEMPORARY TABLE spell_backup (
  id INTEGER PRIMARY KEY NOT NULL,
  added datetime NOT NULL,
  rulebook_id int(11) NOT NULL,
  page smallint(5) DEFAULT NULL,
  name varchar(64) NOT NULL,
  slug varchar(64) NOT NULL,
  school_id int(11) NOT NULL,
  sub_school_id int(11) DEFAULT NULL,
  verbal_component tinyint(1) NOT NULL DEFAULT 0,
  somatic_component tinyint(1) NOT NULL DEFAULT 0,
  material_component tinyint(1) NOT NULL DEFAULT 0,
  arcane_focus_component tinyint(1) NOT NULL DEFAULT 0,
  divine_focus_component tinyint(1) NOT NULL DEFAULT 0,
  xp_component tinyint(1) NOT NULL DEFAULT 0,
  corrupt_component tinyint(1) NOT NULL DEFAULT 0,
  corrupt_level smallint(5)  DEFAULT NULL,
  meta_breath_component tinyint(1) NOT NULL DEFAULT 0,
  true_name_component tinyint(1) NOT NULL DEFAULT 0,
  extra_components varchar(256) DEFAULT NULL,
  casting_time varchar(256) DEFAULT NULL,
  range varchar(256) DEFAULT NULL,
  target varchar(256) DEFAULT NULL,
  effect varchar(256) DEFAULT NULL,
  area varchar(256) DEFAULT NULL,
  duration varchar(256) DEFAULT NULL,
  saving_throw varchar(128) DEFAULT NULL,
  spell_resistance varchar(64) DEFAULT NULL,
  description longtext NOT NULL,
  description_html longtext NOT NULL DEFAULT "",
  verified tinyint(1) NOT NULL DEFAULT 0,
  verified_author_id int(11) DEFAULT NULL,
  verified_time datetime DEFAULT NULL);''')
  curs.execute('''INSERT INTO spell_backup SELECT id,added,rulebook_id,page,name,slug, school_id,sub_school_id, verbal_component,somatic_component,material_component,arcane_focus_component,divine_focus_component,xp_component, corrupt_component,corrupt_level,meta_breath_component,true_name_component, extra_components, casting_time, range,target,effect,area, duration, saving_throw,spell_resistance, description,description_html, verified,verified_author_id,verified_time FROM dnd_spell;''')
  curs.execute('''DROP TABLE dnd_spell;''')
  curs.execute('''CREATE TABLE dnd_spell (
  id INTEGER PRIMARY KEY NOT NULL,
  added datetime NOT NULL DEFAULT CURRENT_TIMESTAMP,
  rulebook_id INTEGER NOT NULL,
  page smallint(5) DEFAULT NULL,
  name varchar(64) NOT NULL,
  slug varchar(64) NOT NULL,
  school_id int(11) NOT NULL,
  sub_school_id int(11) DEFAULT NULL,
  verbal_component tinyint(1) NOT NULL DEFAULT 0,
  somatic_component tinyint(1) NOT NULL DEFAULT 0,
  material_component tinyint(1) NOT NULL DEFAULT 0,
  arcane_focus_component tinyint(1) NOT NULL DEFAULT 0,
  divine_focus_component tinyint(1) NOT NULL DEFAULT 0,
  xp_component tinyint(1) NOT NULL DEFAULT 0,
  corrupt_component tinyint(1) NOT NULL DEFAULT 0,
  corrupt_level smallint(5)  DEFAULT NULL,
  meta_breath_component tinyint(1) NOT NULL DEFAULT 0,
  true_name_component tinyint(1) NOT NULL DEFAULT 0,
  extra_components varchar(256) DEFAULT NULL,
  casting_time varchar(256) DEFAULT NULL,
  range varchar(256) DEFAULT NULL,
  target varchar(256) DEFAULT NULL,
  effect varchar(256) DEFAULT NULL,
  area varchar(256) DEFAULT NULL,
  duration varchar(256) DEFAULT NULL,
  saving_throw varchar(128) DEFAULT NULL,
  spell_resistance varchar(64) DEFAULT NULL,
  description longtext NOT NULL DEFAULT "description missing",
  description_html longtext NOT NULL DEFAULT "",
  verified tinyint(1) NOT NULL DEFAULT 0,
  verified_author_id int(11) DEFAULT NULL,
  verified_time datetime DEFAULT NULL,
UNIQUE(rulebook_id, name),
FOREIGN KEY(verified_author_id) REFERENCES auth_user(id),
  CONSTRAINT "rulebook_id_refs_id_514d0131604a89b" FOREIGN KEY ("rulebook_id") REFERENCES "dnd_rulebook" ("id"),
  CONSTRAINT "school_id_refs_id_5015d8d2133c7ac3" FOREIGN KEY ("school_id") REFERENCES "dnd_spellschool" ("id"),
  CONSTRAINT "sub_school_id_refs_id_75647c3f68dd90be" FOREIGN KEY ("sub_school_id") REFERENCES "dnd_spellsubschool" ("id")
  );''')
  curs.execute('''INSERT INTO dnd_spell SELECT * FROM spell_backup;''')
  curs.execute('''DROP TABLE spell_backup;''')
  curs.execute('''CREATE INDEX index_dnd_spell_name ON dnd_spell (name);''')
  # creating an index on dnd_spell.name actually caused spell_name_to_id cost to INCREASE from 241sec to 259sec...
  # but that was because spell_name_to_id was using the LIKE operator, and apparently SQLite doesn't use the index properly with the LIKE operator even when you don't have a leading %.
  # Altering spell_name_to_id to use = instead of LIKE took some doing, and without an index it didn't make for any significant speedup...
  # but creating and index *and* using = instead of LIKE caused spell_name_to_id runtime to drop like a rock, from 241sec to 1sec.

  # Changing the schema of dnd_spell enables sanitizing, but does not by itself sanitize, so we still have all those varchar '0's in xp_component etc
  # Since casting times and ranges tend to be standardized, it would make more sense to have an intersection table spell_range and an intersection table spell_casting_time...
  # select count(*) from (select distinct casting_time from dnd_spell); returns 60, and a lot of those are probably misspellings or things like "1 action" for 1 standard action
  # For that matter, there are only 375 distinct areas: select count(*) from (select distinct area from dnd_spell);

  curs.execute('''CREATE TABLE dnd_range (
  id INTEGER PRIMARY KEY NOT NULL,
  name varchar(32),
  feetCL0 INTEGER,
  feet_per_level INTEGER
  );''')
  curs.execute('''INSERT INTO dnd_range (name,feetCL0,feet_per_level) VALUES (?,?,?)''',
               ('Close', 25, fraction_to_negative('1/2') ) )
  curs.execute('''INSERT INTO dnd_range (name,feetCL0,feet_per_level) VALUES (?,?,?)''',
               ('Unlimited', eightByteIntMax, eightByteIntMax) )

def insert_psionic_powers(curs):
  curs.execute('''INSERT INTO dnd_rulebook (dnd_edition_id, name, description, year, official_url, slug, published) VALUES (1, "Revised (v.3.5) System Reference Document", "The System Reference Document is a comprehensive toolbox consisting of rules, races, classes, feats, skills, various systems, spells, magic items, and monsters compatible with the d20 System version of Dungeons & Dragons and various other roleplaying games from Wizards of the Coast. You may consider this material Open Game Content under the Open Game License, and may use, modify, and distribute it.", 2004, "http://www.wizards.com/default.asp?x=d20/article/srd35", "system-reference-document", ?);''',
    (datetime.date(2004, 5, 21),) )
  srd_id = curs.lastrowid
  assert srd_id == id_from_name(curs, 'dnd_rulebook', "Revised (v.3.5) System Reference Document")
  curs.execute('''CREATE TEMPORARY TABLE school_backup (id int, name varchar(32), slug varchar(32) );''')
  curs.execute('''INSERT INTO school_backup SELECT id, name, slug FROM dnd_spellschool;''')
  curs.execute('''DROP TABLE dnd_spellschool;''')
  curs.execute('''CREATE TABLE dnd_spellschool (
  id INTEGER PRIMARY KEY NOT NULL,
  name varchar(32) NOT NULL,
  slug varchar(32) NOT NULL,
UNIQUE(name)
  );''')
  curs.execute('''CREATE UNIQUE INDEX dnd_spellschool_name ON dnd_spellschool(name);''')
  curs.execute('''INSERT INTO dnd_spellschool (name, slug) SELECT name, slug FROM school_backup;''')
  curs.execute('''DROP TABLE school_backup;''')
  curs.execute('''INSERT INTO dnd_spellschool (name, slug) VALUES (?, ?)''', ('Psychometabolism', 'psychometabolism') )
  psychometabolismID = curs.lastrowid
  curs.execute('''INSERT INTO dnd_spellschool (name, slug) VALUES (?, ?)''', ('Psychoportation', 'psychoportation') )
  psychoportationID = curs.lastrowid
  curs.execute('''INSERT INTO dnd_spellschool (name, slug) VALUES (?, ?)''', ('Psychokinesis', 'psychokinesis') )
  psychokinesisID = curs.lastrowid
  curs.execute('''INSERT INTO dnd_spellschool (name, slug) VALUES (?, ?)''', ('Metacreativity', 'metacreativity') )
  metacreativityID = curs.lastrowid
  curs.execute('''INSERT INTO dnd_spellschool (name, slug) VALUES (?, ?)''', ('Clairsentience', 'clairsentience') )
  clairsentienceID = curs.lastrowid
  curs.execute('''INSERT INTO dnd_spellschool (name, slug) VALUES (?, ?)''', ('Telepathy', 'telepathy') )
  telepathyID = curs.lastrowid
  
  recreate_spell_table(curs)

  # It will be better to get the SRD psionic powers from dnd35.db, but that power table makes everything varchar even level, so will need processing
  for (name, discipline, xp_cost, manifesting_time, spellrange, target, area, effect, duration, saving_throw,power_resistance, full_text) in get_powers_from_dnd35_db():
    slug = name.lower().replace(',','').replace(' ','-')
    curs.execute('''SELECT id FROM dnd_spellschool WHERE name like ?''', (discipline,) )
    schoolID = curs.fetchone()[0]
    curs.execute('''INSERT INTO dnd_spell (rulebook_id, name, slug, school_id, xp_component, casting_time, range, target, effect, area, duration, saving_throw, spell_resistance, description, description_html) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)''',
                 (srd_id, name, slug, schoolID, xp_cost, manifesting_time, spellrange, target, effect, area, duration, saving_throw, power_resistance, full_text, full_text) )

  # still have problems with things like defensive precognition/precognition, defensive  
  def insertPower(name, schoolID, url):
    # https://stackoverflow.com/questions/418898/sqlite-upsert-not-insert-or-replace
    curs.execute('''INSERT OR IGNORE INTO dnd_spell (rulebook_id, name, slug, school_id, description_html) VALUES (?, ?, ?, ?, ?)''',
                 (srd_id, name, name.lower().replace(' ', '-'), schoolID, url) )
  insertPower('Defensive Precognition', clairsentienceID, "http://www.d20srd.org/srd/psionic/powers/precognitionDefensive.htm")
  insertPower('Offensive Precognition', clairsentienceID, "http://www.d20srd.org/srd/psionic/powers/precognitionOffensive.htm")
  insertPower('Offensive Prescience', clairsentienceID, "http://www.d20srd.org/srd/psionic/powers/precognitionOffensive.htm")

  curs.execute('''INSERT OR IGNORE INTO dnd_spell (added, rulebook_id, name, school_id, description, description_html, slug) VALUES (?, ?, ?, ?, ?, ?, ?)''',
               (datetime.datetime.now(), srd_id, "Ego Whip", telepathyID, "Your rapid mental lashings assault the ego of your enemy, debilitating its confidence. The target takes 1d4 points of Charisma damage, or half that amount (minimum 1 point) on a successful save. A target that fails its save is also dazed for 1 round.", "http://www.d20srd.org/srd/psionic/powers/egoWhip.htm", 'ego-whip') )
  curs.execute('''INSERT OR IGNORE INTO dnd_spell (added, rulebook_id, name, slug, school_id, description, description_html) VALUES (?, ?, ?, ?, ?, ?, ?)''',
               (datetime.datetime.now(), srd_id, "Id Insinuation", 'id-insinuation', telepathyID, "As the confusion spell, except as noted here.", "http://www.d20srd.org/srd/psionic/powers/idInsinuation.htm") )
  curs.execute('''INSERT OR IGNORE INTO dnd_spell (added, rulebook_id, name, slug, school_id, description, description_html) VALUES (?, ?, ?, ?, ?, ?, ?)''',
               (datetime.datetime.now(), srd_id, "Disable", 'disable', telepathyID, "You broadcast a mental compulsion that convinces one or more creatures of 4 Hit Dice or less that they are disabled. Creatures with the fewest HD are affected first. Among creatures with equal Hit Dice, those who are closest to the power's point of origin are affected first. Hit Dice that are not sufficient to affect a creature are wasted. Creatures that are rendered helpless or are destroyed when they reach 0 hit points cannot be affected.", "http://www.d20srd.org/srd/psionic/powers/disable.htm") )
  curs.execute('''INSERT OR IGNORE INTO dnd_spell (added, rulebook_id, name, slug, school_id, description, description_html) VALUES (?, ?, ?, ?, ?, ?, ?)''',
               (datetime.datetime.now(), srd_id, "False Sensory Input", 'false-sensory-input', telepathyID, "You have a limited ability to falsify one of the subject's senses. The subject thinks she sees, hears, smells, tastes, or feels something other than what her senses actually report. You can't create a sensation where none exists, nor make the subject completely oblivious to a sensation, but you can replace the specifics of one sensation with different specifics. For instance, you could make a human look like a dwarf (or one human look like another specific human), a closed door look like it is open, a vat of acid smell like rose water, a parrot look like a bookend, stale rations taste like fresh fruit, a light pat feel like a dagger thrust, a scream sound like the howling wind, and so on.", "http://www.d20srd.org/srd/psionic/powers/falseSensoryInput.htm") )
  curs.execute('''INSERT OR IGNORE INTO dnd_spell (added, rulebook_id, name, slug, school_id, description_html) VALUES (?, ?, ?, ?, ?, ?)''',
               (datetime.datetime.now(), srd_id, "Mental Disruption", 'mental-disruption', telepathyID, "http://www.d20srd.org/srd/psionic/powers/mentalDisruption.htm") )
  curs.execute('''INSERT OR IGNORE INTO dnd_spell (added, rulebook_id, name, slug, school_id, description_html) VALUES (?, ?, ?, ?, ?, ?)''',
               (datetime.datetime.now(), srd_id, "Mindlink", 'mindlink', telepathyID, "http://www.d20srd.org/srd/psionic/powers/mindlink.htm") )
  curs.execute('''INSERT OR IGNORE INTO dnd_spell (added, rulebook_id, name, slug, school_id, description_html) VALUES (?, ?, ?, ?, ?, ?)''',
               (datetime.datetime.now(), srd_id, "Psionic Dominate", 'dominate-psionic', telepathyID, "http://www.d20srd.org/srd/psionic/powers/dominatePsionic.htm") )
  #curs.execute('''INSERT INTO dnd_spell (rulebook_id, name, slug, school_id, description_html) VALUES (?, ?, ?, ?, ?)''',
  #             (srd_id, "Thought Shield", 'thought-shield', telepathyID, "http://www.d20srd.org/srd/psionic/powers/thoughtShield.htm") )
  curs.execute('''INSERT OR IGNORE INTO dnd_spell (rulebook_id, name, slug, school_id, description_html) VALUES (?, ?, ?, ?, ?)''',
               (srd_id, "Psionic Modify Memory", 'modify-memory-psionic', telepathyID, "http://www.d20srd.org/srd/psionic/powers/modifyMemoryPsionic.htm") )
  curs.execute('''INSERT OR IGNORE INTO dnd_spell (rulebook_id, name, slug, school_id, description_html) VALUES (?, ?, ?, ?, ?)''',
               (srd_id, "Cloud Mind", 'cloud-mind', telepathyID, "http://www.d20srd.org/srd/psionic/powers/cloudMind.htm") )
  #curs.execute('''INSERT INTO dnd_spell (rulebook_id, name, slug, school_id, description_html) VALUES (?, ?, ?, ?, ?)''',
  #             (srd_id, "Read Thoughts", 'read-thoughts', telepathyID, "http://www.d20srd.org/srd/psionic/powers/readThoughts.htm") )
  insertPower('Detect Hostile Intent', telepathyID, "http://www.d20srd.org/srd/psionic/powers/detectHostileIntent.htm")
  curs.execute('''INSERT OR IGNORE INTO dnd_spell (rulebook_id, name, slug, school_id, description_html) VALUES (?, ?, ?, ?, ?)''',
               (srd_id, "Remote Viewing", 'remote-viewing', clairsentienceID, "http://www.d20srd.org/srd/psionic/powers/remoteViewing.htm") )
  curs.execute('''INSERT OR IGNORE INTO dnd_spell (rulebook_id, name, slug, school_id, description_html) VALUES (?, ?, ?, ?, ?)''',
               (srd_id, "Aura Sight", 'aura-sight', clairsentienceID, "http://www.d20srd.org/srd/psionic/powers/auraSight.htm") )
  curs.execute('''INSERT OR IGNORE INTO dnd_spell (rulebook_id, name, slug, school_id, description_html) VALUES (?, ?, ?, ?, ?)''',
               (srd_id, "Detect Psionics", 'detect-psionics', clairsentienceID, "http://www.d20srd.org/srd/psionic/powers/detectPsionics.htm") )
  insertPower('Mental Barrier', clairsentienceID, "http://www.d20srd.org/srd/psionic/powers/mentalBarrier.htm")
  curs.execute('''INSERT OR IGNORE INTO dnd_spell (rulebook_id, name, slug, school_id, description_html) VALUES (?, ?, ?, ?, ?)''',
               (srd_id, "Psionic Plane Shift", 'plane-shift-psionic', psychoportationID, "http://www.d20srd.org/srd/psionic/powers/planeShiftPsionic.htm") )
  curs.execute('''INSERT OR IGNORE INTO dnd_spell (rulebook_id, name, slug, school_id, description_html) VALUES (?, ?, ?, ?, ?)''',
               (srd_id, "Psionic Daze", 'daze-psionic', telepathyID, "http://www.d20srd.org/srd/psionic/powers/dazePsionic.htm") )
  curs.execute('''INSERT OR IGNORE INTO dnd_spell (rulebook_id, name, slug, school_id, description_html) VALUES (?, ?, ?, ?, ?)''',
               (srd_id, "Psionic Charm", 'charm-psionic', telepathyID, "http://www.d20srd.org/srd/psionic/powers/charmPsionic.htm") )
  curs.execute('''INSERT OR IGNORE INTO dnd_spell (rulebook_id, name, slug, school_id, description_html) VALUES (?, ?, ?, ?, ?)''',
               (srd_id, "Psionic Dominate", 'dominate-psionic', telepathyID, "http://www.d20srd.org/srd/psionic/powers/dominatePsionic.htm") )
  curs.execute('''INSERT OR IGNORE INTO dnd_spell (rulebook_id, name, slug, school_id, description_html) VALUES (?, ?, ?, ?, ?)''',
               (srd_id, "Psionic Suggestion", 'suggestion-psionic', telepathyID, "http://www.d20srd.org/srd/psionic/powers/suggestionPsionic.htm") )
  insertPower('Far Hand', psychokinesisID, "http://www.d20srd.org/srd/psionic/powers/farHand.htm")
  insertPower('Concealing Amorpha', metacreativityID, "http://www.d20srd.org/srd/psionic/powers/concealingAmorpha.htm")
  insertPower('Entangling Ectoplasm', metacreativityID, "http://www.d20srd.org/srd/psionic/powers/entanglingEctoplasm.htm")
  #insertPower('Telekinetic Thrust', psychokinesisID, "http://www.d20srd.org/srd/psionic/powers/telekineticThrust.htm")
  insertPower('Catfall', psychoportationID, "http://www.d20srd.org/srd/psionic/powers/catfall.htm")
  #insertPower('Concussion Blast', psychokinesisID, "http://www.d20srd.org/srd/psionic/powers/concussionBlast.htm")
  insertPower('Inertial Armor', psychokinesisID, "http://www.d20srd.org/srd/psionic/powers/inertialArmor.htm")
  insertPower('Aversion', telepathyID, "http://www.d20srd.org/srd/psionic/powers/aversion.htm")
  #insertPower('Mind Thrust', telepathyID, "http://www.d20srd.org/srd/psionic/powers/mindThrust.htm")
  insertPower('Exhalation of the Black Dragon', psychometabolismID, "http://www.d20srd.org/srd/psionic/powers/exhalationoftheBlackDragon.htm")
  insertPower('Body Purification', psychometabolismID, "http://www.d20srd.org/srd/psionic/powers/bodyPurification.htm")
  insertPower('Reach', psychometabolismID, "http://archive.wizards.com/default.asp?x=dnd/psm/20010928a")

  insertPower('Corrupt Water', 7, "http://www.d20srd.org/srd/monsters/dragonTrue.htm#blackDragon")
  insertPower('Charm Reptiles', 4, "http://www.d20srd.org/srd/monsters/dragonTrue.htm#blackDragon")
  insertPower('Destroy Water', 4, "http://www.d20srd.org/srd/monsters/dragonTrue.htm#blueDragon")
  insertPower('Luck Bonus', 7, "http://www.d20srd.org/srd/monsters/dragonTrue.htm#goldDragon")
  insertPower('Detect Gems', 3, "http://www.d20srd.org/srd/monsters/dragonTrue.htm#goldDragon")
  insertPower('Create Wine', 2, "http://www.d20srd.org/srd/monsters/genie.htm#djinni")
  insertPower('Dreamscape', 2, "http://www.d20srd.org/srd/epic/spells/dreamscape.htm")
  insertPower('Nailed to the Sky', 2, "http://www.d20srd.org/srd/epic/spells/nailedToTheSky.htm")
  insertPower('Hellball', 5, "http://www.d20srd.org/srd/epic/spells/hellball.htm")
  insertPower('Enslave', 4, "http://www.d20srd.org/srd/epic/spells/enslave.htm")
  insertPower('Safe Time', 2, "http://www.d20srd.org/srd/epic/spells/safeTime.htm")
  insertPower('Time Duplicate', 2, "http://www.d20srd.org/srd/epic/spells/timeDuplicate.htm")
  insertPower('Contingent Recall and Resurrection', 2, "http://www.d20srd.org/srd/epic/monsters/hoaryHunter.htm")
  insertPower('Contingent Resurrection', 2, "http://www.d20srd.org/srd/epic/spells/contingentResurrection.htm")
  insertPower('Ruin', 7, "http://www.d20srd.org/srd/epic/spells/ruin.htm")
  insertPower('Peripety', 1, "http://www.d20srd.org/srd/epic/spells/peripety.htm")
  insertPower('Spell Worm', 4, "http://www.d20srd.org/srd/epic/spells/spellWorm.htm")
  insertPower('Demise Unseen', 6, "http://www.d20srd.org/srd/epic/spells/demiseUnseen.htm")

  # Complete Psionic:
  insertPower('Energy Arc', psychokinesisID, "A cone of the chosen type of energy shoots from your fingertips. Any creature in the area of effect takes 1d4 points of damage.")
  insertPower('Primal Fear', telepathyID, "shaken for 1 round. This effect doesn't stack with other fear effects.")

  # Frostburn and Sandstorn:
  insertPower('Energy Emanation', psychokinesisID, "1d6 points of energy damage to all creatures within the area every round. Creatures in the area must make a new Fortitude save each round.")
  insertPower('Energy Flash', psychokinesisID, "5d6 points of damage to the creature touched, doing cold, electricity, fire, or sonic damage. In addition to the energy damage, the target is dazed for 1 round on a failed Fortitude save (the same save that determines full or half damage).")
  insertPower('Fish Cloud', 2, "When submerged in water, a vodyanoi can summon a huge school of magic fish to provide concealment (similar to the fog cloud spell). This school of fish swims around the point the vodyanoi designates in a 20-foot radius. This cloud of fish obscures all sight, including darkvision, beyond 5 feet. A creature within 5 feet has concealment. Creatures farther away have total concealment. A strong current disperses the cloud of fish in 4 rounds. A very strong current disperses the cloud of fish in 1 round. The fish created by this spell are formed by magic; they are not real animal, and objects and energies pass through them as  though they were not there. A vodyanoi can summon a fish cloud three times per day. The fish cloud remains for 40 minutes or until dispersed or dispelled.")
  insertPower('Global Warming', 5, "You increase the temperature of the region, drying up water and baking the soil within a 100-mile-radius area.")
  insertPower('Inconstant Location', psychoportationID, "At the beginning of your turn, as a swift action, you can teleport yourself to any other space to which you have line of sight, so long as that space is no farther than you could move in normal move action.")

  # Tome of Magic:
  insertPower('Black Candle', 5, "This mystery functions like the spell light or the spell darkness.")
  insertPower('Dusk and Dawn', 5, "You make a dark area lighter or a light area darker, blanketing the affected area in shadowy illumination.")
  insertPower('Pass Into Shadow', 2, "This mystery functions like the spell plane shift, except that your destination or origination must be the Plane of Shadow.")
  insertPower('Bolster', 7, "You grant the subject 5 temporary hit points for each of its Hit Dice (maximum 75).")
  insertPower('Sight Eclipsed', 8, "While this mystery is in effect, you can attempt Hide checks even while being observed, just as if you had cover or concealment for the purpose of this determination.")
  insertPower('Umbral Body', 7, "You gain the incorporeal subtype (see page 164) and all advantages and traits associated with it.")
  insertPower('Potent Word of Nurturing', 12, "You grant a creature fast healing 10.")
  insertPower('Critical Word of Nurturing', 12, "you utter the reverse form of the life-giving words. You deal 8d6 points of damage to the subject.")
  insertPower('Incarnation of Angels', 7, "The target gains the celestial creature template (MM 31). The target gains the fiendish creature template (MM 107).")
  insertPower("Archer's Eye", 7, "Your target's ranged attacks ignore penalties for concealment because her aim sharpens to focus on the unconcealed parts of her foe.")
  insertPower("Shockwave", 5, "A violent shock wave travels through the air, and creatures in the area must make Fortitude saves or be knocked prone and take 1d4 points of nonlethal damage.")
  insertPower("Preternatural Clarity", 3, "+5 insight bonus on any single attack roll, opposed ability or skill check, or saving throw. When your target uses the insight bonus, those within 10 feet of her can hear an echo of your original utterance, even if you're no longer present. Activating the effect is an immediate action. The target can choose to apply the bonus after she has rerolled the d20, but before the Dungeon Master reveals the result of the check.")
  insertPower("Morale Boost", 3, "This utterance functions as the remove fear spell (PH 271). Your target becomes frightened (DMG 301) by the susurrant Truespeech being whispered in its ear.")
  insertPower("Hidden Truth", 3, "You grant the target a +10 bonus on a single Knowledge check and enable her to use the skill, even if untrained. If the target has bardic knowledge, lore, or a similar class feature, this bonus can apply to that check instead. Your target gains a +10 bonus on a single Bluff check made before the duration of the utterance expires.")

  # Kingdoms of Kalamar Player's Guide
  insertPower("Disinter", 7, "Kalamar: brings to the surface any item that you buried")
  insertPower("Silken Grasp", 7, "Kalamar: +20 grapple initial hold")
  insertPower("Cannibalize", 6, "http://www.d20pfsrd.com/magic/3rd-party-spells/kobold-press-open-design/cannibalize/")
  insertPower('Caustic Bile', psychometabolismID, "???")
  insertPower('Insatiable Hunger', psychometabolismID, "???")

  insertPower('Regenerate Worldskin 50', 2, "As a standard action, Ragnorra can initiate regeneration in any worldskin feature (see page 102). This ability has unlimited range. The growth regains 50 hit points per round.")
  insertPower('Regenerate Worldskin', 2, "As a standard action, Ragnorra can initiate regeneration in any worldskin feature (see page 102). This ability has unlimited range. The growth regains 50 hit points per round.")
  insertPower('Skincasting', 2, "As a standard action, Ragnorra can activate the ability of any worldskin feature (see page 102) within 1,000 feet.")
  insertPower('Mass Aversion', telepathyID, "An anathema creates a compulsion effect targeting all enemies within 30 feet. The targets must succeed on a Will save (DC 27) or gain an aversion to snakes for 10 minutes. Affected subjects must stay at least 20 feet from any snake, yuan-ti, or ti-khana creature (described earlier in this book), whether alive or dead; if already within 20 feet, they move away. A subject can overcome the compulsion by succeeding on another Will save (DC 27), but still suffers from deep anxiety. This causes a -4 reduction to Dexterity until the effect wears off or the subject is no longer within 20 feet of a snake, yuan-ti, or ti-khana creature. This ability is otherwise similar to antipathy as cast by a 16th-level sorcerer.")



def cache_IMarvinTPA(cacheDir='IMarvinTPAcache'):
  if cacheDir is None: cacheDir='IMarvinTPAcache'
  if os.path.exists(cacheDir):
    print('IMarvinTPAcache already cached')
    return
  print('creating folder', cacheDir)
  os.mkdir(cacheDir)
  os.mkdir(os.path.join(cacheDir, 'images') )
  logging.basicConfig()
  logging.getLogger().setLevel(logging.DEBUG)
  requests_log = logging.getLogger("requests.packages.urllib3")
  requests_log.setLevel(logging.DEBUG)
  requests_log.propagate = True
  start = time.time()
  for i in range(1, 979):
    #i = 1000
    # Monster 978 does not exist. monster 977 does
    #print('cache_IMarvinTPA', i)
    url = r'http://www.imarvintpa.com/dndlive/TokMonsters.php?ID=' + str(i)
    rptok = os.path.join(cacheDir, '{0:03d}'.format(i) + '.rptok')
    print('cache_IMarvinTPA', i, (time.time() - start)/128, url, 'about to urlretrieve')
    if False:
      try:
        urllib.request.urlretrieve(url, filename=rptok)
      except urllib.error.HTTPError as E:
        if str(E) != "HTTP Error 404: Not Found":
          raise
        # if we get a 404, we're done (sometimes it does that instead of the text thing, maybe it's a user-agent issue)
        break
    elif False:
      # It hangs exactly 128seconds when I ping from Python, not at all when I use Chrome. How does it know?
      headers = {
                 'Host':'www.imarvintpa.com',
                 'Connection':'keep-alive',
                 'User-Agent':r'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/61.0.3163.100 Safari/537.36',
                 'Upgrade-Insecure-Requests':1,
                 'Accept':r'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8',
                 'Accept-Encoding':'gzip, deflate',
                 'Accept-Language':r'en-US,en;q=0.8',
                 'Cookie':'PHPSESSID=91398-wkyL4qShHt6vZJ30',
                 }
      print('cache_IMarvinTPA', i-1, (time.time() - start)/128, 'about to requests.get')
      response = requests.get(url, headers=headers)
      print('cache_IMarvinTPA', i, (time.time() - start)/128, 'done requests.get')
      if response.status_code != requests.codes.ok:
        if response.status_code != 404:
          response.raise_for_status()
        else: # 404
          break
      open(rptok, 'wb').write(response.content)
    else:
      # For some reason, urllib and requests are both delayed exactly 128seconds, but curl isn't.
      curl = pycurl.Curl()
      curl.setopt(curl.URL, url)
      with open(rptok, 'wb') as fileObj:
        curl.setopt(curl.WRITEDATA, fileObj)
        curl.perform()
        curl.close()
    print('cache_IMarvinTPA downloaded', url, i, (time.time() - start)/128)
    # the token files are ZIP files http://lmwcs.com/rptools/wiki/Token
    try:
      zipfile.ZipFile(rptok)
    except zipfile.BadZipFile:
      #if not zipfile.is_zipfile(rptok): # for some reason, sometimes says not is_zipfile but can still open it
      # then it should be a text file saying Monster 1000 does not exist.
      assert os.path.getsize(rptok) < 1000
      text = open(rptok, 'r').read()
      if text != "Monster {} does not exist.\n".format(i):
        print('rptok =', rptok, 'text =', text, open(rptok, 'rb').read() )
        raise Exception(text)
      else: # no monster with that number, so we've reached the end
        print('cache_IMarvinTPA has reached the end of the monsters')
        break
    #with zipfile.ZipFile(rptok) as archive:
      #archive.extract(os.path.join('assets', ) )
    #break

accidentallyEscapedSlash = re.compile(r'<i>(\w+)<\\/i>')
assert accidentallyEscapedSlash.sub(r'<i>\1</i>', " <i>spikes<\/i> +30 ") == " <i>spikes</i> +30 "
def parse_IMarvinTPA(curs, cacheDir='IMarvinTPAcache'):
  """
  The big problem is matching up monsters whose names are not listed exactly the same.
  Unfortunately, neither source has page numbers.
http://monsterfinder.dndrunde.de/ has pages for the following books:
Book of Exalted Deeds 
Book of Vile Darkness 
Fiend Folio 
Libris Mortis 
Lords of Madness 
Monster Manual I (3.5) (without entries from SRD)
Monster Manual II 
Monster Manual III 
Monster Manual IV 
Monsters of Faerun 
Planar Handbook 
System Reference Document (with most from MM)
  """
  if cacheDir is None: cacheDir='IMarvinTPAcache'

  curs.execute('''CREATE TABLE monster_family (
  id INTEGER PRIMARY KEY NOT NULL,
  name TEXT NOT NULL
  );''')
  curs.execute('''CREATE TEMPORARY TABLE feat_backup (
  id INTEGER PRIMARY KEY NOT NULL,
  rulebook_id INTEGER NOT NULL,
  page smallint(2) DEFAULT NULL,
  name varchar(64) NOT NULL,
  description longtext NOT NULL,
  benefit longtext NOT NULL,
  special longtext NOT NULL,
  normal longtext NOT NULL,
  slug varchar(64) NOT NULL,
  description_html longtext NOT NULL,
  benefit_html longtext NOT NULL,
  special_html longtext NOT NULL,
  normal_html longtext NOT NULL,
FOREIGN KEY(rulebook_id) REFERENCES dnd_rulebook(id)
  );''')
  curs.execute('''INSERT INTO feat_backup (rulebook_id, page, name, description, benefit, special, normal, slug, description_html, benefit_html, special_html, normal_html) SELECT rulebook_id, page, name, description, benefit, special, normal, slug, description_html, benefit_html, special_html, normal_html FROM dnd_feat;''')
  curs.execute('''DROP TABLE dnd_feat;''')
  curs.execute('''CREATE TABLE dnd_feat (
  id INTEGER PRIMARY KEY NOT NULL,
  rulebook_id INTEGER NOT NULL,
  page smallint(2) DEFAULT NULL,
  name varchar(64) NOT NULL,
  description longtext NOT NULL,
  benefit longtext NOT NULL,
  special longtext NOT NULL,
  normal longtext NOT NULL,
  slug varchar(64) NOT NULL,
  description_html longtext NOT NULL,
  benefit_html longtext NOT NULL,
  special_html longtext NOT NULL,
  normal_html longtext NOT NULL,
UNIQUE(rulebook_id, name)
FOREIGN KEY(rulebook_id) REFERENCES dnd_rulebook(id)
  );''')
  curs.execute('''INSERT INTO dnd_feat SELECT * FROM feat_backup;''')
  curs.execute('''CREATE INDEX index_dnd_feat_name ON dnd_feat (name);''')
  curs.execute('''DROP TABLE dnd_monsterhasfeat;''')
  curs.execute('''CREATE TABLE monster_has_feat (
  monster_id INTEGER NOT NULL,
  feat_id INTEGER NOT NULL,
FOREIGN KEY(monster_id) REFERENCES dnd_monster(id),
FOREIGN KEY(feat_id) REFERENCES dnd_feat(id)
  );''')
  
  """sqlite> select dnd_rulebook.name, dnd_edition_id, dnd_dndedition.name, published, year from dnd_feat inner join dnd_rulebook on rulebook_id=dnd_rulebook.id inner join dnd_dndedition on dnd_edition_id=dnd_dndedition.id where dnd_feat.name="Improved Sunder";
Deities and Demigods|7|Supplementals (3.0)|2002-02-01|
Enemies and Allies|7|Supplementals (3.0)|2001-10-01|
Player's Handbook v.3.5|1|Core (3.5)|2000-08-01| yeah, that was actually published in 2003
Sword and Fist: A Guidebook to Monks and Fighters|7|Supplementals (3.0)|2001-01-01|
  """

  for i in range(1, 1000):
    rptok = os.path.join(cacheDir, '{0:03d}'.format(i) + '.rptok')
    if not os.path.exists(rptok):
      break
    with zipfile.ZipFile(rptok) as archive:
      #fileObj = archive.open('content.xml')
      #content = fileObj.read()
      content = archive.read('content.xml')
      content = content.decode('us-ascii')
      #print(content)
      content = accidentallyEscapedSlash.sub(r'<i>\1</i>', content)
      root = defusedxml.ElementTree.fromstring(content)
      assert root.tag == 'net.rptools.maptool.model.Token'
      # family only appears as eg "family":"Abomination"
      race = root.find(r"./propertyMap/store/entry[string='Race']/net.rptools.CaseInsensitiveHashMap_-KeyValue[key='Race']/value[@class='string']")
      raceDict = ast.literal_eval(race.text)
      #name = raceDict['name']
      #typeName = raceDict['type']
      #family = raceDict['family']
      #subtypes = raceDict['subtype']
      if raceDict['name'] == 'Bralani':
        raceDict['family'] = 'Eladrin'
      if raceDict['family'] is not None and raceDict['family'] != '':
        curs.execute('''INSERT INTO monster_family (name) VALUES (?);''', (raceDict['family'],) )
      if any(raceDict['name'].startswith(ageCategory) for ageCategory in ('Wyrmling', 'Very young', 'Young', 'Juvenile', 'Young Adult', 'Adult', 'Mature adult', 'Old', 'Very old', 'Ancient', 'Wyrm', 'Great wyrm') ):
        continue

      monsterID = id_from_name(curs, 'dnd_monster', raceDict['name'])
      if monsterID is None:
        # 'family': 'Dragon, Epic', 'name': 'Wyrmling Force Dragon'
        # Dragon, Epic, Force, Wyrmling
        familyFirst = raceDict['family'] + ', ' + ' '.join(word for word in raceDict['name'].split() if word not in [w.rstrip(',') for w in raceDict['family'].split()])
        monsterID = id_from_name(curs, 'dnd_monster', familyFirst)
        if monsterID is None:
          words = raceDict['name'].split()
          maybeDescriptor = words[0]
          assert maybeDescriptor[-1] != ','
          if len(words) > 1 and maybeDescriptor not in raceDict['family']:
            descriptorLast = raceDict['family'] + ', ' + ' '.join(word for word in words if word not in [w.rstrip(',') for w in raceDict['family'].split()] and word!=maybeDescriptor) + ', ' + maybeDescriptor
            print('descriptorLast =', descriptorLast)
            monsterID = id_from_name(curs, 'dnd_monster', descriptorLast)
      if monsterID is None:
        words = raceDict['name'].split()
        if len(words) > 1:
          maybeDescriptor = words[0]
          assert maybeDescriptor[-1] != ','
          descriptorLast = ' '.join(word for word in words[1:]) + ', ' + maybeDescriptor
          print('descriptorLast =', descriptorLast)
          # there are two elder treants with different stats, one from Monster Mayhem
          monsterID = id_from_name(curs, 'dnd_monster', descriptorLast, useEdition=True)
      if monsterID is None:
        words = raceDict['name'].split()
        if len(words) == 2:
          # Behemoth Eagle to Behemoth, Eagle
          monsterID = id_from_name(curs, 'dnd_monster', words[0] + ', ' + words[1])
      if monsterID is None:
        if raceDict['name'] == 'Devastation Spider': continue
        monsterID = id_from_name(curs, 'dnd_monster', raceDict['name'], allowExtraOnRight=True) # Devastation Spider, Web-spinner
        # really, when there are multiple variations we want to assign feats to each...in that particular case, hmm...I guess a monster can't have multiple land_speeds...
      if monsterID is None:
        if raceDict['name'] == 'Demilich': # http://www.d20srd.org/srd/epic/monsters/demilich.htm
          continue # skip templates for now
        if raceDict['name'][:13] == 'Pseudonatural': continue
        continue
        raise Exception(raceDict)

      # <fe>fast healing</fe> 15, <i>spell resistance</i> 34, <fe>damage reduction</fe> 10/chaotic and epic and adamantine
      # no idea what's going on with <fe>
      #print(root)
      #for child in root.find('propertyMap').find('store'):
      #  print(child)
      for feat in root.findall(r"./propertyMap/store/entry[string='RealFeats']/net.rptools.CaseInsensitiveHashMap_-KeyValue[key='RealFeats']/value[@class='string']/fe"):
        featName = feat.text
        featID = id_from_name(curs, 'dnd_feat', featName, useEdition=True)
        curs.execute('''INSERT INTO monster_has_feat (monster_id, feat_id) VALUES (?,?);''', (monsterID, featID) )
      #with archive.open('content.xml') as xml:
        #defusedxml.sax.parse(xml)
        #defusedxml.ElementTree.parse(xml)
        # https://docs.python.org/3/library/xml.html
    # IMarvinTPA's XML files do not have sources, but the HTML at http://www.imarvintpa.com/dndlive/Monsters.php?ID=1 do

# IMarvinTPA family same as what comes before the comma? sometimes
"""
sqlite> select name from dnd_monster where name like "%, %";
Dragon, Planar, Beast, Wyrmling
Dragon, Planar, Concordant, Wyrmling
Golem, Black Ice
Eel, Moray
Vulture, Giant
Demon, Anzu
Living Spell, Living Fireball
Demon, Skurchur
Imp, Choleric
White Hart (White Stag, Elven Deer)
Imix (Prince of evil Fire Creatures, Lord of Hellfire, The Eternal Flame)
Elemental Weird, Air, Lesser
Swarm, Crystal Beetle
"""



nameParentheticalRE = re.compile(r"\(([\w\-\,' ]+)\)")
parenBeforeCommaRE = re.compile(r'\).*,')

AGE_CATEGORIES = ('Wyrmling', 'Very Young', 'Young', 'Juvenile', 'Young Adult', 'Adult', 'Mature Adult', 'Old', 'Very Old', 'Ancient', 'Wyrm', 'Great Wyrm', 'Elder') # Elder sometimes means other things
MONSTERS_WITH_AGES = set( ('Dragon', 'Arrowhawk', 'Tojanida', 'Urkhan Worm (Tractor Worm)') )
MONSTERS_WITH_LINEAR_ORDERING = set()
# Salamanders reproduce asexually, each producing a single larva every ten years and incubating the young in fire pits until they reach maturity. Flamebrothers and average salamanders are actually different species, while nobles rise from the ranks of the average.
SIZE_NAMES = ('Fine', 'Diminutive', 'Tiny', 'Small', 'Medium', 'Large', 'Huge', 'Gargantuan', 'Colossal')
appearInNamesMoreThanOnce = [collections.defaultdict(set) for i in range(4) ]
allUnknownTerms = (set(), set(), set(), set() )


class MonsterName(object):
  def __init__(self, name):
    if name[:16] == 'Elemental Primal':
      name = 'Elemental, Primal' + self.name[16:]
    # Lycanthrope, Drow Werebat
    # Lycanthrope, Goblin Rat
    if name == 'Lycanthrope, Goblin Rat, (human form)': name = 'Lycanthrope, Goblin Rat (human form)'
    if name == 'Imp': name = 'Devil, Imp'
    reducedName = name
    delimiters = re.compile('[, ]')
    if ',' in reducedName:
      nameSegments = [segment.strip() for segment in delimiters.split(reducedName) if segment != '']
    else:
      nameSegments = [segment.strip() for segment in reducedName.split(' ')]
    nameSegments = [sensible_title(segment) for segment in nameSegments]
    # Screech Owl or Owl, Screech? DMG uses Screech Owl
    if name == 'Screech Owl':
      nameSegments = [name]
    # nameSegments = [segment for segment in nameSegments if segment not in SIZE_NAMES and segment.lower() != 'medium-size'] # have size separately
    #print('nameSegments =', nameSegments)
    nameSegments[-1] = nameParentheticalRE.sub('', nameSegments[-1]).strip()
    # Parentheticals that appear before the last comma are usually ones we want to keep, like Chromatic (Lost).
    if len(nameSegments) > 5:
      raise Exception(self.name)
    self.orderedInBookBy = list()
    if 'Dire' in nameSegments and 'Lycanthrope' not in nameSegments:
      # Bear, Polar, Dire
      # Dire Rat
      self.orderedInBookBy.append('Dire')
      nameSegments.remove('Dire')
    elif 'Monstrous' in nameSegments:
      nameSegments.remove('Monstrous')
      startsWith = 'Monstrous'
      if len(nameSegments) > 0:
        startsWith += ' ' + nameSegments[0]
        nameSegments = nameSegments[1:]
      self.orderedInBookBy.append(startsWith)
    elif 'Fox' in nameSegments: # and 'Dire' not in nameSegments
      self.orderedInBookBy.append('Fox')
      nameSegments.remove('Fox')
    elif 'Spider' in nameSegments:
      self.orderedInBookBy.append('Spider')
      nameSegments.remove('Spider')
    elif 'Snake' in nameSegments and 'Sea' not in nameSegments: # and 'Dire' not in nameSegments
      self.orderedInBookBy.append('Snake')
      nameSegments.remove('Snake')
    elif 'Lizard' in nameSegments and 'Shocker' not in nameSegments:
      self.orderedInBookBy.append('Lizard')
      nameSegments.remove('Lizard')
    elif 'Shocker' in nameSegments:
      nameSegments.remove('Shocker')
      startsWith = 'Shocker'
      if len(nameSegments) > 0:
        startsWith += ' ' + nameSegments[0]
        nameSegments = nameSegments[1:]
      self.orderedInBookBy.append(startsWith)
    elif 'Blink' in nameSegments:
      nameSegments.remove('Blink')
      startsWith = 'Blink'
      if len(nameSegments) > 0:
        startsWith += ' ' + nameSegments[0]
        nameSegments = nameSegments[1:]
      self.orderedInBookBy.append(startsWith)
    elif 'Gem' in nameSegments:
      nameSegments.remove('Gem')
      startsWith = 'Gem'
      if len(nameSegments) > 0:
        startsWith += ' ' + nameSegments[0]
        nameSegments = nameSegments[1:]
      self.orderedInBookBy.append(startsWith)
    elif 'Hell' in nameSegments:
      nameSegments.remove('Hell')
      startsWith = 'Hell'
      if len(nameSegments) > 0:
        startsWith += ' ' + nameSegments[0]
        nameSegments = nameSegments[1:]
      self.orderedInBookBy.append(startsWith)
    elif 'Winter' in nameSegments:
      nameSegments.remove('Winter')
      startsWith = 'Winter'
      if len(nameSegments) > 0:
        startsWith += ' ' + nameSegments[0]
        nameSegments = nameSegments[1:]
      self.orderedInBookBy.append(startsWith)
    elif 'Elemental' in nameSegments and 'Demon' not in nameSegments:
      self.orderedInBookBy.append('Elemental')
      nameSegments.remove('Elemental')
    elif 'Formian' in nameSegments:
      self.orderedInBookBy.append('Formian')
      nameSegments.remove('Formian')
    elif 'Mephit' in nameSegments:
      self.orderedInBookBy.append('Mephit')
      nameSegments.remove('Mephit')
    elif nameSegments[0] == 'Sewerm':
      self.orderedInBookBy.append('Snake') # ordered under Snake in Serpent Kingdoms
    self.orderedInBookBy.extend(nameSegments)
  def ezkajii(self):
    nameSegments = self.orderedInBookBy
    if 'Dire' in nameSegments:
      nameSegments.remove('Dire')
      nameSegments.append('Dire')
    # Sea Snake really is ordered in Stormwrack as Sea Snake, not under Snake, despite referencing Snake in MM
    elif nameSegments[0] == 'Sea' and nameSegments [1] == 'Snake':
      nameSegments[0:2] = ['Sea Snake']
    elif len(nameSegments) >= 2 and nameSegments[-1] == 'Python' and nameSegments[-2] == 'Tree':
      nameSegments[-2:] = ['Tree Python']
    return ', '.join(nameSegments)


class Monster(object):
  #allEnvs = set()
  @staticmethod
  def splitSpecialAbilities(commaSeparatedSpecialAbilities):
    if commaSeparatedSpecialAbilities == '-':
      return []
    else:
      specialAbs = [ab.strip() for ab in commaSeparatedSpecialAbilities.split(',')]
      return [ab for ab in specialAbs if ab!='']
  def set_default_alignment(self, default_alignment_string):
    if default_alignment_string == 'Any' or default_alignment_string == 'N': default_alignment_string = 'NN'
    self.lawChaosID = default_alignment_string[0]#lawChaosToInt[default_alignment_string[0]]
    self.goodEvilID = goodEvilToInt[default_alignment_string[-1]]

  def __init__(self, xls_row):
    self.name = xls_row[0].value
    if self.name[:16] == 'Elemental Primal':
      self.name = 'Elemental, Primal' + self.name[16:]
    # Lycanthrope, Drow Werebat
    # Lycanthrope, Goblin Rat
    if self.name == 'Lycanthrope, Goblin Rat, (human form)': self.name = 'Lycanthrope, Goblin Rat (human form)'
    elif self.name == 'Quasit': self.name = 'Demon, Quasit'
    reducedName = self.name
    parentheticals = list()
    for matchObj in nameParentheticalRE.finditer(self.name):
      parenthetical = matchObj.group(1)
      #if ',' in parenthetical and parenthetical not in ('Hill Troll, Mountain Troll', 'Cougar, Mountain Lion', 'The Hunger Eternal, The Devourer, the Ebon Maw', 'White Stag, Elven Deer'):
      #  raise Exception(self.name)
      if ',' in parenthetical:
        aliasList = [alias.strip() for alias in parenthetical.split(',')]
        #print(self.name, 'aliasList =', aliasList)
        reducedName = reducedName.replace(matchObj.group(0), '') # eliminates parenthesized commas for splitting later
      #if ',' in self.name[matchObj.end(0):]:
      #  if parenthetical not in ('Lost', 'Starkin'):
      #    print('parenthetical', self.name)
      if parenthetical[-5:] == ' form':
        form = parenthetical[:-5]
      elif parenthetical[:4] == 'with':
        form = parenthetical
      elif parenthetical[-14:] == ' manifestation':
        form = parenthetical[:-14]
      elif parenthetical in ('incorporeal', 'manifested fully', 'unfurled', 'hunched'):
        form = parenthetical
      else:
        parentheticals.append(parenthetical)
        print(self.name, file=open('unknownParentheticals.txt', 'a') )
    if len(parentheticals) > 1:
      raise Exception(self.name, parentheticals)
    #reducedName = nameParentheticalRE.sub('', self.name) # handily eliminates parenthesized commas
    nameSegments = [segment.strip() for segment in reducedName.split(',')]
    nameSegments = [segment for segment in nameSegments if segment not in SIZE_NAMES and segment.lower() != 'medium-size'] # have size separately
    nameSegments[-1] = nameParentheticalRE.sub('', nameSegments[-1]).strip()
    # Parentheticals that appear before the last comma are usually ones we want to keep, like Chromatic (Lost).
    if len(nameSegments) > 5:
      raise Exception(self.name)
    if 'Dire' in nameSegments and 'Lycanthrope' not in nameSegments:
      alphabetizedBy = 'Dire' # Bear, Polar, Dire
    elif 'Legendary' in nameSegments:
      alphabetizedBy = 'Legendary'
    elif 'hydra' in nameSegments[0]:
      alphabetizedBy = 'Hydra' # Cryohydra, 5-Headed
    else:
      alphabetizedBy = nameSegments[0]
    #if not (all(segment not in AGE_CATEGORIES for segment in nameSegments) or (nameSegments[-1] in AGE_CATEGORIES and alphabetizedBy in MONSTERS_WITH_AGES) ):
    #  raise Exception(self.name)
    nameSegments = tuple(nameSegments)
    assert 'Dragon' not in nameSegments or alphabetizedBy == 'Dragon'
    if alphabetizedBy in MONSTERS_WITH_AGES:
      assert nameSegments[-1] in AGE_CATEGORIES
      ageCategory = nameSegments[-1]
      nameSegments = nameSegments[:-1]
    if 'Lycanthrope' in nameSegments and alphabetizedBy != 'Lycanthrope':
      raise Exception(self.name, nameSegments, alphabetizedBy)
    if alphabetizedBy == 'Lycanthrope':
      # only odd one out is Lycanthrope, Wereboar, Dire, Hill Giant
      if len(nameSegments) > 2:
        assert nameSegments == ('Lycanthrope', 'Wereboar', 'Dire', 'Hill Giant')
        nameSegments = ('Lycanthrope', 'Hill Giant Dire Wereboar')
    #for segment in nameSegments[:-1]:
    #  if '(' in segment:
    #    matchObj = nameParentheticalRE.search(segment)
    #    if matchObj is None:
    #      raise Exception(segment, self.name)
    #    parenthetical = matchObj.group(1)
    #    if parenthetical not in ('Lost', 'Starkin'):
    #      raise Exception(nameSegments, parenthetical)
    for i,segment in enumerate(nameSegments):
      appearInNamesMoreThanOnce[i][segment].add(tuple(nameSegments[i+1:]) )
    allUnknownTerms[len(nameSegments) - 1].add(', '.join(nameSegments) )

    self.size = xls_row[1].value
    if self.name == 'Lycanthrope, Wererat (rat form)': self.size = 'Small'
    self.type_name = xls_row[2].value
    if self.name in ('Energon, Xag-Ya','Energon, Xeg-Yi'):
      self.type_name = 'Outsider'
      # The .xls lists it as Elemental, but lists the rulebook as MotP and I checked the Manual of the Planes, it says Outsider.
    if xls_row[3].value == '': self.subtypes = [] # ''.split(',') is not an empty list, but rather a list containing ''
    else: self.subtypes = [fix_subtype(subtype) for subtype in xls_row[3].value.replace(' or ', ', ').replace('[alignment subtype]', 'Good, Evil, Lawful, Chaotic').split(',')]
    for subtype in self.subtypes: assert subtype != ''
    self.HitDice = fraction_to_negative(xls_row[4].value)

    self.landSpeed = None if xls_row[6].value == '' else int(xls_row[6].value)
    swimflyburrowcrawl = [] if xls_row[7].value == '' else xls_row[7].value.split(', ')
    self.movementModes = list()
    self.maneuverability = None
    for mode in swimflyburrowcrawl:
      if self.name == 'Uloriax' and mode == 'f20':
        mode = 'c20' # http://archive.wizards.com/default.asp?x=dnd/fw/20040509a
      elif self.name == 'Thoqqua' and mode == 's10':
        continue # Thoqquas cannot swim
      matchObj = movementModeRE.match(mode)
      if matchObj is None:
        if self.name != "Inevitable, Marut":
          raise Exception(self.name, swimflyburrowcrawl, mode)
        # dunno why Inevitable, Marut has 34, but it does draw my attention to the fact that its speed is wrongly listed as 30feet when it should be 40feet
        # no others have similar notation
      else:
        self.movementModes.append( (matchObj.group(1), int(matchObj.group(2) ) ) )
        if matchObj.group(1) == 'f':
          try:
            if self.name == 'Albatross':
              self.maneuverability = 'avg'
            elif self.name == 'Energon, Xag-Az':
              self.maneuverability = 'prf' # http://archive.wizards.com/default.asp?x=dnd/psb/20021122c
            else:
              self.maneuverability = matchObj.group(3)
            if self.maneuverability not in ('clu', 'pr', 'avg', 'gd', 'prf', 'prft'):
              raise ValueError(self.name, mode, self.maneuverability)
          except IndexError:
            raise Exception(self.name)

    self.ArmorClass = int(xls_row[8].value)
    touchAC = int(xls_row[9].value)
    if self.name == 'Dragon, Chromatic, Black, Young' and touchAC == 11: touchAC = 10 # error in table
    elif self.name == 'Ghast' and touchAC == 12: touchAC = 13 # error in Monster Manual
    elif self.name == 'Vargouille' and touchAC == 11: touchAC = 12 # error in Monster Manual
    elif self.name == 'Swarm, Bat': touchAC = self.ArmorClass # error in Monster Manual
    elif self.name == 'Inevitable, Marut':
      self.ArmorClass = int(xls_row[7].value)
      touchAC = int(xls_row[8].value)
    elif self.name == 'Monstrous Centipede, Colossal':
      # Colossal Monstrous scorpions and spiders both lose 2 Dex from Gargantuan, and the Colossal Monstrous Centipede has initiative, Reflex, and Hide consistent with a decreased Dex.
      self.ArmorClass -= 1
      touchAC -= 1
    self.armorPlusShieldPlusNaturalArmor = self.ArmorClass - touchAC
    if self.name == 'Minotaur':
      flatfootedAC = self.ArmorClass
    else:
      flatfootedAC = int(xls_row[10].value)

    # http://stackoverflow.com/questions/2415398/can-i-set-a-formula-for-a-particular-column-in-sql
    # https://stackoverflow.com/questions/1124695/can-i-create-computed-columns-in-sqlite
    # SQLite doesn't supported computed columns.
    
    attacks = xls_row[12].value # single attack
    #attacks = xls_row[13].value # full attack often "Same as Attack" or some abbreviated form like "4 tentacles"
    if attacks == 'Huge +5 triple flail +51 (1d12+24/19-20 [for each of 1d3 heads]':
      # In this particular case, paren is opened but should not be closed at the end
      attacks = 'Huge +5 triple flail +51 (1d12+24/19-20) [for each of 1d3 heads]'
    elif '(' in attacks and ')' not in attacks: # mismatched parentheses
      if attacks[-1] == '(' or attacks[-1] == '_': # typo, hit the key immediately to left or right of closing paren
        attacks = attacks[:-1] + ')'
      else: # typo, closing paren was left off
        attacks = attacks + ')'
    elif ')' in attacks and '(' not in attacks:
      if attacks.count(')') == 2:
        attacks = attacks.replace(')', '(', 1)
    elif attacks == 'Slam +21 (2d6+12 or gore +21 (4d8+12)':
      attacks = 'Slam +21 (2d6+12) or gore +21 (4d8+12)'
    elif attacks[:28] == 'Short sword +10 1d6+4/19-20)':
      attacks = 'Short sword +10 (1d6+4/19-20)' + attacks[28:]
    elif attacks == 'Tail slap +4 (1d6+1 plus positive energy) or tail touch +4 (positive energy(':
      attacks = 'Tail slap +4 (1d6+1 plus positive energy) or tail touch +4 (positive energy)'
    elif attacks[:43] == 'Wand of Orcus (+6 chaotic unholy greatclub)':
      attacks = '+6 chaotic unholy greatclub' + attacks[43:]
    elif attacks[:42] == 'Ruby Rod of Asmodeus (+6 unholy greatclub)':
      attacks = '+6 unholy greatclub' + attacks[42:]

    if attacks[:6] == 'Bite 1':
      attacks = 'Bite +1' + attacks[6:]
    elif attacks == 'Bite +56 (2d8+18/19-20/plus 1d6)':
      attacks = 'Bite +56 (2d8+18/19-20 plus 1d6)'
    elif attacks[:24] == 'Greataxe +27 (4d6+18x/3)':
      attacks = 'Greataxe +27 (4d6+18/x3)' + attacks[24:]
    elif attacks[-10:] == '(1d8+2x/3)':
      attacks = attacks[:-10] + '(1d8+2/x3)'
    elif attacks[-17:] == 'javelin +4 (d4+1)':
      attacks = attacks[:-17] + 'javelin +4 (1d4+1)'
    elif attacks[:22] == 'Morningstar +3 (16d+1)':
      attacks = 'Morningstar +3 (1d6+1)' + attacks[22:]
    elif attacks[-10:] == '(18d+3/x3)':
      attacks = attacks[:-10] + '(1d8+3/x3)'
    elif attacks[-7:] == ' (1d6+)':
      attacks = attacks[:-7] + ' (1d6)'

    if attacks == 'Fire bolt +24 (8d6 fire/19-20)':
      attacks = 'Fire bolt +24 (8d6/19-20 fire)'
    elif attacks == 'Flaming sword +36 (3d10 fire/17-20+3d10 fire) or slam +36 (3d6+10 plus 3d6 fire)':
      attacks = 'Flaming sword +36 (3d10/17-20 fire plus 3d10 fire) or slam +36 (3d6+10 plus 3d6 fire)'
    elif attacks == 'Masterwork scimitar +14 (1d6+6 plus 1d6 fire/18-20)':
      attacks = 'Masterwork scimitar +14 (1d6+6/18-20 plus 1d6 fire)'
    elif attacks == '+5 flaming burst halberd of speed +33 (2d8+23 plus 1d6 fire/19-20/x3 plus 2d10 fire) or +5 seeking composite longbow (+12 Str bonus) +36 (2d6+17/x3)':
      # Hmm, flaming burst is tricky...
      attacks = '+5 flaming burst halberd of speed +33 (2d8+23 plus 1d6 fire) or +5 seeking composite longbow (+12 Str bonus) +36 (2d6+17/x3)'
    elif attacks == 'Slam +8 (1d8+7) or bolt +5 (3d8+3 plus 1d6 fire/x2 plus 1d10 fire)':
      # Hmm, flaming burst is tricky...
      attacks = 'Slam +8 (1d8+7) or bolt +5 (3d8+3 plus 1d6 fire)'
    elif attacks == 'Masterwork lance +14 (2d6+7/x3) or masterwork heavy flail +14 (2d8+7/19-20) or composite longbow (+5 Str bonus) +9 (1d8+6/x3 plus 1d6 electricity [plus 2d10 electricity on critical hit])':
      # Hmm, shocking burst is tricky...
      attacks = 'Masterwork lance +14 (2d6+7/x3) or masterwork heavy flail +14 (2d8+7/19-20) or composite longbow (+5 Str bonus) +9 (1d8+6/x3 plus 1d6 electricity)'
    elif attacks == '+5 acidic burst bastard sword +41 (2d8+18/17-20 plus 1d6 acid [plus 1d10 acid on critical hit])':
      # Hmm, acidic burst is tricky...
      attacks = '+5 acidic burst bastard sword +41 (2d8+18/17-20 plus 1d6 acid)'
    elif attacks == '+1 icy burst cold iron maul +32 (3d8+35/19-20/x3 plus 1d6 cold [plus 2d10 cold on critical hit]) (adj for Pwr Atk) or rock +22 (2d8+15)':
      # Hmm, acidic burst is tricky...
      attacks = '+1 icy burst cold iron maul +32 (3d8+35/19-20/x3 plus 1d6 cold) or rock +22 (2d8+15)'
    elif attacks == '+5 shocking burst adamantine scourge +37 (1d8+12/19-20 plus 1d6 electricity [plus 1d10 electricity on critical hit])':
      attacks = '+5 shocking burst adamantine scourge +37 (1d8+12/19-20 plus 1d6 electricity)'
    elif attacks == '+3 fleshgrinding vile ranseur +42 (2d4+15/x3 plus 1 vile [plus extra 2 vile on critical hit]) or claw +38 (1d4+8 plus 1 vile)':
      attacks = '+3 fleshgrinding vile ranseur +42 (2d4+15/x3 plus 1 vile) or claw +38 (1d4+8 plus 1 vile)'
    elif attacks == '+5 flaming burst icy burst ranseur +41 (2d6+15/19-20/x3 plus 1d6 cold and 1d6 fire [plus 2d10 cold and 2d10 fire on critical hit])':
      attacks = '+5 flaming burst icy burst ranseur +41 (2d6+15/19-20/x3 plus 1d6 cold and 1d6 fire)'
    elif attacks == '+2 flaming burst ranseur +19 (2d6+5/19-20/x3 plus 1d6 fire [2d10 fire on critical hit[)':
      attacks = '+2 flaming burst ranseur +19 (2d6+5/19-20/x3 plus 1d6 fire)'
    elif attacks == 'Ruby Rod +51 (2d6+27 plus 3d8+15 negative energy [plus 2d6 against good-aligned creatures])':
      attacks = 'Ruby Rod +51 (2d6+27 plus 3d8+15 negative energy plus 2d6 against good creatures)'
    elif attacks == 'Triple flail +31 (3d6+25, plus see text) (adj for Pwr Atk)':
      attacks = 'Triple flail +31 (3d6+25 plus see text) (adj for Pwr Atk)'
    # Crit specifier does often come before added damage, eg:
    # Armblade +20 (1d8+9/19-20 plus 2d6 holy plus 1d6 fire)
    elif attacks == 'Slam +7 (1d6 plus 1d6 energy/19-20)':
      attacks = 'Slam +7 (1d6/19-20 plus 1d6 energy)'
    elif attacks == 'Gargantuan +4 shortspear +50 (2d8+23 plus 1 vile/x3)':
      attacks = 'Gargantuan +4 shortspear +50 (2d8+23/x3 plus 1 vile)'
    elif attacks == '+4 ranseur +50 (2d4+13 plus 1 vile/x3)':
      attacks = '+4 ranseur +50 (2d4+13/x3 plus 1 vile)'
    elif attacks == 'Flame blade +46 (2d8+20 fire/15-20)':
      attacks = 'Flame blade +46 (2d8+20/15-20 fire)'
    elif attacks == '+4 rapier +46 (1d6+16 plus 2d6 precise strike/15-20)':
      attacks = '+4 rapier +46 (1d6+16/15-20 plus 2d6 precise strike)'
    elif attacks == '+5 bloodfeeding flaming greatsword +49 (2d6+20 plus 1d6 fire/19-20)':
      attacks = '+5 bloodfeeding flaming greatsword +49 (2d6+20/19-20 plus 1d6 fire)'
    elif attacks == 'Slam +52 (1d10+15 plus 1 vile plus withering/19-20)':
      attacks = 'Slam +52 (1d10+15/19-20 plus 1 vile plus withering)'
    elif attacks == 'Huge +5 ranseur +50 (2d6+17 plus 1d6 fire or 1d6 cold/x3)':
      attacks = 'Huge +5 ranseur +50 (2d6+17/x3 plus 1d6 fire or 1d6 cold)'
    elif attacks == 'Slam +54 (4d8+24 plus 1d6 Con drain plus energy drain/19-20 plus 1d6 plus death)':
      # Right next to Atropus, Father Llymic Claw +33 (3d6+11/18-20/x3 plus 2d6 cold plus soul chill)
      attacks = 'Slam +54 (4d8+24/19-20 plus 1d6 Con drain plus energy drain plus 1d6 plus death)'
    elif attacks == 'Tail +27 (8d6 acid/19-20) or claw +24 ranged (1d10+11/19-20)':
      attacks = 'Tail +27 (8d6/19-20 acid) or claw +24 ranged (1d10+11/19-20)'
    elif attacks == 'Touch +49 (2d6 Con drain/19-20)':
      attacks = 'Touch +49 (2d6/19-20 Con drain)'
    elif attacks == 'Large scythe +13 (2d6+18 plus entropic blade/19-20/x4) (adj for Pwr Atk)':
      attacks = 'Large scythe +13 (2d6+18/19-20/x4 plus entropic blade) (adj for Pwr Atk)'
    elif attacks == 'Tentacle +19 (1d8+9 plus 1d6 electricity/19-20/x3)':
      attacks = 'Tentacle +19 (1d8+9/19-20/x3 plus 1d6 electricity)'
    elif attacks == 'Slam +28 (2d10+11) or +5 flaming burst warhammer +33 (1d8+16 plus 1d6 fire/x3 plus 1d10 fire)':
      attacks = 'Slam +28 (2d10+11) or +5 flaming burst warhammer +33 (1d8+16/x3 plus 1d6 fire)'
    elif attacks == 'Spinning blade +43 (2d6+12/19-20 [+1d6 on critical hit]) or slam +35 (2d6+6) or shocking touch +35 (2d6+6 electricity) or electricity ray +35 (10d6 electricity) or spike +30 (2d6+12)':
      attacks = 'Spinning blade +43 (2d6+12/19-20) or slam +35 (2d6+6) or shocking touch +35 (2d6+6 electricity) or electricity ray +35 (10d6 electricity) or spike +30 (2d6+12)'
    elif attacks == 'Slam +44 (3d6+18/19-20 plus hunefer rot +1d6 on critical hit])':
      attacks = 'Slam +44 (3d6+18/19-20 plus hunefer rot)'
    elif attacks == 'Claw +38 (3d8+21/19-20 plus blazefire [+1d6 on critical hit])':
      attacks = 'Claw +38 (3d8+21/19-20 plus blazefire)'
    elif ' [+1d6 on critical hit]' in attacks:
      attacks = attacks.replace(' [+1d6 on critical hit]', '')
    elif attacks == 'Slam +50 (6d8+19/19-20/+1d6)':
      attacks = 'Slam +50 (6d8+19/19-20)'
    elif attacks == '+5 warhammer +87 (4d8+30/19-20 [+2d6 on critical hit]) or +5 javelin +70 (2d10+22/19-20)':
      attacks = '+5 warhammer +87 (4d8+30/19-20) or +5 javelin +70 (2d10+22/19-20)'

    elif attacks == '2 Large +1 composite longbows (+5 Str bonus) +14/+14 (2d6+6/x3) or claw +15 (1d6+5)':
      attacks = '2 Large +1 composite longbows (+5 Str bonus) +14 (2d6+6/x3) or claw +15 (1d6+5)'
    elif attacks == 'Horn +16 or +13 ranged (2d6+6 plus poison)':
      attacks = 'Horn +16 (2d6+6 plus poison) or Horn +13 ranged (2d6+6 plus poison)'
    elif attacks == 'Trident +4 or +6 ranged (1d8+1)':
      attacks = 'Trident +4 (1d8+1) or Trident +6 ranged (1d8+1)'
    elif attacks == 'Dagger +2 melee or ranged (1d4+1)':
      attacks = 'Dagger +2 (1d4+1) or Dagger +2 ranged (1d4+1)'
    elif attacks == 'Claw +8 (2d4+5) or lajatang +7/+7 (2d6+5/2d6+2)':
      attacks = 'Claw +8 (2d4+5) or lajatang +7 (2d6+5)'

    attacks = attacks.replace(' (mimic)', '')
    attacks = attacks.replace(', plus ', ' plus ')
    attacks = attacks.replace('Adamantine +1 throwing returning maul', '+1 Adamantine throwing returning maul')
    attacks = attacks.replace('adamantine +1 throwing returning maul', '+1 adamantine throwing returning maul')

    self.special_damage = set()

    if attacks == '-':
      pass
    elif attacks[:5] == 'Swarm' or attacks[:3] == 'Mob':
      # has no attack bonus
      pass
    elif attacks == 'Vine 1d2+1 plus poison':
      pass
    elif attacks == '12 slams +9 (1d4+4) or Huge weapon +9/+4 (dmg +4) and 11 Huge light weapons +9 (dmg +2), or Huge weapon +7/+2 (dmg +4) and 11 Huge weapons (1 or more non-light weapons) +7 (dmg +2)':
      pass
    elif attacks == 'Claw +13 (2d4+8) or weapon +13 (dmg +8) or weapon +11 (dmg +8) or stone +5 (1d6+8)':
      pass
    elif ' (dmg +' in attacks or ' (dmg+' in attacks: # don't deal with these "arbitrary weapon" entries yet
      pass
    else:
      #matchObj = dieRollRE.search(attacks)
      #if matchObj is not None:
      #  print(self.name, attacks, 'dieRollRE', matchObj.group(0) )
      #matchObj = numericalDamageRE.search(attacks)
      #if matchObj is not None:
      #  print(self.name, attacks, 'numericalDamageRE', matchObj.group(0) )
      #matchObj = re.search(r'\(' + singleDamageREstring + '\)', attacks)
      #if matchObj is not None:
      #  print(self.name, attacks, 'singleDamageREstring', matchObj.group(0) )
      #matchObj = parenthesizedDamageRE.search(attacks)
      #if matchObj is not None:
      #  print(self.name, attacks, 'parenthesizedDamageRE', matchObj.group(0) )
      # attacks.split(' or ') doesn't work because of Slam +22 (2d6+12 plus 3d6 sonic or 3d6 electricity
      matchObj = attacksRE.match(attacks)
      #if matchObj is None:
        #matchObj = attackWithoutParensRE.match(attacks)
      if matchObj is None:
        print('no attacks match for', self.name, 'attacks', attacks)
        raise Exception(self.name, attacks)
      # for attackMatchObj in itertools.chain(singleAttackModeRE.finditer(self.attacks), singleAttackModeRE.finditer(self.fullAttack)):
      # for attackMatchObj in singleAttackModeRE.finditer(self.attacks):
      for g in attacksRE.match(attacks).groups():
        # interior groups will also be there
        if g is None: continue # some groups will not match
        attackMatchObj = singleAttackModeRE.match(g)
        if attackMatchObj is not None:
          damage = attackMatchObj.group(SINGLE_ATTACK_DAMAGE_GROUP_NUMBER)
          if damage is None:
            if attackMatchObj.group(2) not in ('web', 'Energy touch'):
              raise Exception(self.name, attacks, attackMatchObj, attackMatchObj.group(0), attacksRE.match(attacks).groups())
            else:
              self.special_damage.add(attackMatchObj.group(2))
          else:
            if not numericalDamageRE.match(damage) or numericalDamageRE.match(damage).group(0) != damage: # anything other than straight HP damage, including Int damage
              self.special_damage.add(damage)
      firstAttack = matchObj.group(1)
      firstAttackMatchObj = singleAttackModeRE.match(firstAttack)
      numberOfAttacks = None if firstAttackMatchObj.group(1) is None else int(firstAttackMatchObj.group(1) )
      # a few monsters can make multiple attacks with a single attack action
      #if numberOfAttacks is not None:
      #  print(self.name, self.HitDice, self.type_name, numberOfAttacks, attacks) # only with more than hydra is darktentacles 9HD
      attackName = firstAttackMatchObj.group(2)
      #else:
      #  print(self.name, 'attacks', attacks, matchObj.group(1), matchObj.group(2) )
    #if (self.type_name == 'Humanoid' or self.type_name == 'Animal') and ('Slam' in attacks or 'slam' in attacks):
    #  print(self.type_name, self.name, 'has a slam attack', attacks)
    #self.attacks = attacks if xls_row[13].value == 'Same as Attack' else xls_row[13].value
    self.attacks = attacks
    self.fullAttack = xls_row[13].value

    self.strength = integer_or_non(xls_row[20].value)
    if self.name == 'Dragon, Chromatic, Black, Adult': self.strength = 23 # error in table
    self.dexterity = integer_or_non(xls_row[21].value)
    if self.name == 'Monstrous Scorpion, Colossal': self.dexterity = 8 # error in table
    self.constitution = integer_or_non(xls_row[22].value)
    self.intelligence = integer_or_non(xls_row[23].value)
    self.wisdom = int(xls_row[24].value)
    self.charisma = int(xls_row[25].value)

    listedFortitude = int(xls_row[17].value)
    #if self.type_name in fortitude_divisors:
    #  predictedFortitude = (self.HitDice if self.HitDice > 0 else 1)//fortitude_divisors[self.type_name] + fortitude_additions[self.type_name] + (0 if self.constitution is None else (self.constitution - 10)//2)
    #  if listedFortitude != predictedFortitude and listedFortitude != predictedFortitude + 2:
    #    print(self.name, 'has a listed Fortitude of', listedFortitude, 'with Con', self.constitution, 'and', self.HitDice, 'HD, which would predict', predictedFortitude)

    # Use values for hit points and Fortitude as error-checking.
    # The redundant information is helpful here, since when HP and Fortitude outvote HD we can use them to correct HD.
    averageHitPoints = int(xls_row[5].value)
    if self.name == 'Ooze, Gray':
      self.HitDice == 3
      averageHitPoints = 31
    if self.name == 'Ooze, Ochre Jelly':
      self.constitution = 22
    if (self.name.startswith('Ooze, ') and self.HitDice > averageHitPoints/(self.constitution//2 + 0.5)
        and self.HitDice > 3*listedFortitude - 3*self.constitution//2 + 17):
      # Oozes have d10 hit dice, so average (5.5 + Con) HP per Hit Die.
      # HP = HD*(5.5 + constitution//2 - 5) = HD*(constitution//2 + 0.5)
      # HD = HP/(constitution//2 + 0.5)
      # Fort = HD//3 + constituion//2 - 5 = HD//2 + constituion//2 - 5
      # HD//3 = Fort - constitution//2 + 5
      # HD = 3*Fort - 3*constitution//2 + 15 + maybe 2
      self.HitDice = int(math.ceil(averageHitPoints/(self.constitution//2 + 0.5) ) )
      if abs(self.HitDice//3 + self.constitution//2 - 5 - listedFortitude) > 0 and self.name != 'Ooze, Flotsam':
        raise Exception(self.name, self.HitDice, self.constitution, listedFortitude)
        # n the case the the Flotsam Ooze, the error is in the Fiend Folio, so I don't know what the error is.
    if self.name == 'Ooze, Ochre Jelly': assert self.HitDice == 6

    dexterityModifier = self.dexterity//2 - 5 if self.dexterity is not None else 0
    # A creature with no Dexterity score can't move. If it can perform actions (such as casting spells), it applies its Intelligence modifier to initiative checks instead of a Dexterity modifier. The creature automatically fails Reflex saves and Dexterity checks.
    # The modifier for a nonability is +0.
    # An inanimate object has not only a Dexterity of 0 (-5 penalty to AC), but also an additional -2 penalty to its AC.
    # But an immobile creature like a Formian Queen has no penalty to AC. But a Shrieker fungus, despite having no Dex, has a -5 Dex penalty to AC. Similarly an udoroot has no Dex and a -5 Dex penalty to AC.
    # The formian queen, despite having no Dex penalty to AC, does have a Dex penalty to initiative. So does a shrieker.
    # An udoroot, despite having a Dex penalty to AC, has no Dex penalty to initiative. Neither does it apply its Intelligence modifier to initiative; it has initiative of +0, as if they tried to apply the rule the no Dex means no Dex modifier in that case but not in the case of AC.
    #if self.dexterity == 0: print('0 Dex:', self.name)
    #if self.dexterity is None: print('has no Dex:', self.name)
    #if AC - flatfootedAC != max(0, dexterityModifier):
    #  print('flat-footed', flatfootedAC, 'does not match:', AC, '-', max(0, dexterityModifier), self.name)
    # A girallon has a listed flat-footed AC of 15 with an AC of 16 (-1 size, +3 Dex, +4 natural). Not just in the table, in the actual Monster Manual.


    self.environment = xls_row[26].value
    #Monster.allEnvs.add(environment)
    #if environment == "Any":
    #  climate,geography = "Any","Any"
    #elif environment == "Underground":
    #  climate,geography = "Any","Underground"
    #elif environment == "Any land and underground":
    #  climate,geography = "Any","Any"
    #elif environment in ("Baator","Acheron","Celestia","Elemental Plane of Air", "Demiplane of Ectoplasm"):
    #  climate,geography = "Any",environment
    #else:
    #  if ' ' not in environment: raise ValueError(environment)
    #  try:
    #    climate,geography = environment.split()
    #  except ValueError:
    #    raise ValueError(environment)
    default_alignment_string = xls_row[28].value # can be something like "NG or NE"
    if len(default_alignment_string) < 2: raise RuntimeError(default_alignment_string)
    self.set_default_alignment(default_alignment_string)

    if xls_row[29].value == '-': self.challenge_rating = 0
    else: self.challenge_rating = fraction_to_negative(xls_row[29].value)
    if self.challenge_rating < self.HitDice/3 and self.type_name == 'Dragon':
      # suspect that missing a digit like Dragon, Chromatic, Blue, Young Adult
      self.challenge_rating += 10
    elif self.name == 'Ki-rin': self.challenge_rating = 18 # error in table
    elif self.name == 'Mind Shard of Pandorym': self.challenge_rating = 25 # error in table
    elif self.name == 'Abomination, Xixecal': self.challenge_rating = 36 # error in table
    elif self.name == 'Demon, Abyssal Ravager': self.challenge_rating = 5 # error in table

    self.rulebook_abbrev = xls_row[30].value
    if self.name == 'Demon, Alkilith': self.rulebook_abbrev = 'FF'
    elif self.name == 'Ogre, Yrasda': self.rulebook_abbrev = 'SpcSor'

    self.specialAttacks = Monster.splitSpecialAbilities(xls_row[14].value)
    #print('self.specialAttacks =', self.specialAttacks)
    self.specialQualities = Monster.splitSpecialAbilities(xls_row[16].value)

    self.SpellLikeAbilities = list()
    SLAstring = xls_row[15].value
    if 'SLAs' in self.specialAttacks or 'PLAs' in self.specialAttacks:
      if SLAstring == '':
        if self.name=="Devil, Ice (Gelugon)":
          SLAstring = "CL13, At will - cone of cold (DC 20), fly, ice storm (DC 19), greater teleport (self plus 50 pounds of objects only), persistent image (DC 20), unholy aura (DC 23), wall of ice"
        elif self.name=="Eladrin, Coure":
          SLAstring = "CL4, At will - dancing lights, detect evil, detect magic, faerie fire; 3/day - magic missile, sleep"
        elif self.name[:15]=='Eladrin, Ghaele':
          SLAstring = "CL12, At will - aid, charm monster (DC 17), color spray (DC 14), comprehend languages, continual flame, cure light wounds (DC 14), dancing lights, detect evil, detect thoughts (DC 15), disguise self, dispel magic, hold monster (DC 18), greater invisibility (self only), major image (DC 16), see invisibility, greater teleport (self plus 50 pounds of objects only); 1/day - chain lightning (DC 19), prismatic spray (DC 20), wall of force"
        elif self.name == 'Galeb Duhr':
          SLAstring = "CL20, At will - animate objects (stone only), stone shape; 1/day - move earth, passwall, transmute rock to mud, wall of stone"
        elif self.name == 'Half-Fiend, Durzagon':
          SLAstring = "CL10, 3/day - darkness, 1/day - desecrate, enlarge person (self only), invisibility (self only), unholy blight"
        elif self.name == 'Demon, Blood Fiend':
          SLAstring = "CL18, At will - detect good, detect magic, greater teleport (self plus maximum load of objects only); 3/day - chaos hammer, darkness, unholy blight; 1/day - blasphemy, desecrate"
        elif self.name == 'Gargoyle, Crystal': # http://archive.wizards.com/default.asp?x=dnd/psb/20021025c
          SLAstring = "ML3, 1/day - charm person, inflict pain, color spray"
        elif self.name == 'Demon, Malrauthin':
          SLAstring = "CL16, 1/day - horrid wilting"
        # Dragon, Ectoplasmic, Very Young has Ectoplasmic Cocoon PLA but no manifester level, which makes no sense since duration 1round/level http://archive.wizards.com/default.asp?x=dnd/psb/20040123a
        # dunno who the specific golem named Ruby is, but standard Golem, Gemstone, Ruby has no SLAs
        elif self.name in ('Dragon, Ectoplasmic, Very Young', 'Dragon, Ferrous, Tungsten, Young', 'Ruby (Ruby Gemstone Golem)', 'Demon, Malrauthin'):
          pass        
        else:
          raise Exception(self.name, self.rulebook_abbrev, 'has SLAs but they are not apparent')
      if self.name == 'Githyanki':
        SLAstring = "ML=HD/2, 3/day - far hand; 3/day - concealing amorpha; 3/day - dimension door; 3/day - telekinetic thrust; 1/day - plane shift"
      elif self.name == 'Githzerai':
        SLAstring = "ML=HD/2, 3/day - catfall, concussion blast, inertial armor; 1/day - plane shift"
      elif self.name == 'Tayfolk, Tayling':
        SLAstring = "CL=lvl, 3/day - cure minor wounds, mage hand; 3/day - cat's grace; 3/day - haste; 3/day - polymorph"
      elif self.name == "Planetouched, D'hin'ni":
        SLAstring = "CL1, at will - prestidigitation; 1/day - gust of wind, whispering wind, wind wall"
      elif self.name == 'Gingwatzim, Naranzim':
        SLAstring = "CL5, at will - color spray, ghost sound, invisibility (self only), Nystul's magic aura, silent image, ventriloquism; 3/day - blur (self only), hypnotic pattern, minor image, mirror image, misdirection; 1/day - displacement, invisibility sphere, major image (self only); 1/day - greater invisibility, phantasmal killer, rainbow pattern; 1/day - dream, nightmare, persistent image"
      elif self.name == "Titan":
        SLAstring = "CL20, at will - chain lightning, charm monster, cure critical wounds, fire storm, greater dispel magic, hold monster, invisibility, invisibility purge, levitate, persistent image; 3/day - etherealness, word of chaos, summon nature's ally IX; 1/day - gate, maze, meteor swarm; at will - daylight, holy smite, remove curse; 1/day - greater restoration; at will - bestow curse, deeper darkness, unholy blight; 1/day - Bigby's crushing hand"
      elif self.name == "Titan, Stormbringer":
        SLAstring = "CL20, at will - chain lightning, charm monster, control weather, cure critical wounds, fire storm, gaseous form, greater dispel magic, hold monster, invisibility, invisibility purge, water breathing, wind wall; 3/day - whirlwind, wind walk, word of chaos; 1/day - gate, maze, storm of vengeance; at will - daylight, holy smite, remove curse; 1/day - greater restoration; at will - bestow curse, deeper darkness, unholy blight; 1/day - Bigby's crushing hand"
      elif self.name == "Phoenix":
        SLAstring = "CL20, at will - blindness, blink, blur, color spray, cure light wounds, dancing lights, death ward, find the path, find traps, fire seeds, heal, invisibility, misdirection, neutralize poison, produce flame, remove fear, remove curse, see invisibility; 1/day - incendiary cloud, reincarnate, pyrotechnics, summon nature's ally IX, veil, wall of fire; CL40, at will - dismissal, dispel evil, dispel magic"
      elif self.name == "Grisgol" or self.name == "Nature Spirit, Large":
        SLAstring = '' # varies by construction, skip
      elif self.name == "Thaumavore":
        SLAstring = "CLHD, 10/day - comprehend languages, protection from evil, ray of enfeeblement, sleep, invisibility, touch of idiocy, blink, deep slumber, confusion, dimension door, symbol of sleep, antimagic field, plane shift"
      elif self.name == "Ssvaklor":
        SLAstring = "ML7, 1/day - aversion (duration 10 hours*), control light, entangling ectoplasm, id insinuation (up to 4 creatures*), psionic freedom of movement"
      elif self.name[:9] == "Zeitgeist":
        SLAstring = "CL20, at will - animate objects, call lighting storm, calm emotions, confusion, contagion, fear, impeding stones, make whole, move earth, produce flame, remove disease, repel metal or stone, repel wood, soften earth and stone, spike stones, stone shape, wall of iron, wall of stone, wood shape, zone of peace"
      elif self.name == "Gibbering Orb":
        SLAstring = "CL27, at will - wish (takes pit fiend SLA and makes it at will)"
      elif self.name == "Bacchae":
        SLAstring = "CL7, 3/day - charm person, Tasha's hideous laughter; 1/day - crushing despair, fear, good hope, rage" # hack; emotion is replaced with four different 3.5 spells, may as well list them all since they're all available just not on the same day
      elif self.name == "Kelpie":
        SLAstring = "CL7, at will - detect thoughts; 3/day - charm person, crushing despair, fear, good hope, rage"
      elif self.name == "Archon, Word":
        SLAstring = "TL10, 10/day - potent word of nurturing, incarnation of angels, archer's eye; 10/day - shockwave"
      elif self.name == "Devil, Logokron":
        SLAstring = "CL15, at will - major image, greater teleport (self +50lbs only), critical word of nurturing (reverse only), preternatural clarity, morale boost, hidden truth"
      elif self.name == "Demon, Carnevus":
        SLAstring = "CL8 (choose one from each set per individual) 3/day - charm person, disguise self, magic missile, sleep; 3/day - invisibility, Melf's acid arrow, spider climb, web; 3/day - fireball, hold person, lightning bolt, vampiric touch; 3/day - Evard's black tentacles, lesser globe of invulnerability, ice storm, shadow conjuration"
      elif self.name == "Demon, Turagathshnee":
        SLAstring = "CL13, at will - blasphemy, deeper darkness, desecrate, detect good, detect magic, greater teleport (self +50lbs only); 3/day - slow consumption; 1/day - enervation"
      elif self.name == "Shirokinukatsukami":
        SLAstring = "CL14, at will - astral projection, dream, dream sight, gaseous form, invisibility, magic circle against evil, greater teleport (self +50lbs only); 3/day - cloud trapeze (self +50lbs only), dispel evil, dominate monster; 1/day - heal, raise dead; at will - detect evil, detect thoughts, discern shapechanger"
      elif ', Emprix' in self.name:
        SLAstring = "CL12, at will - aid, continual flame, cure light wounds, daylight, detect magic, dispel magic, flare, holy aura, remove disease, remove fear; CL17, at will - charm person, charm monster, confusion, crushing despair, daze, good hope, rage, suggestion; 3/day - mass charm; 1/month - geas/quest; CL15, at will - detect chaos, detect evil, see invisibility, true seeing"
      elif self.name == "Dread, Greater":
        SLAstring = "CL20, at will - detect magic, blasphemy, deeper darkness, desecrate, detect good, detect law, greater dispel magic, fear, pyrotechnics, read magic, suggestion, symbol of death, symbol of fear, symbol of insanity, symbol of fear, symbol of persuasion, symbol of sleep, symbol of stunning, symbol of weakness, telekinesis, greater teleport (self +50lbs only), tongues (self only), unhallow, unholy aura, unholy blight, wall of fire; 3/day - magic missile; 1/day - fire storm, implosion"
      elif self.name == 'Devil, Advespa':
        SLAstring = "CL4, 3/day - disguise self, command, produce flame, pyrotechnics" # general 3.5 update booklet specifies change self is now disguise self http://archive.wizards.com/default.asp?x=dnd/dnd/20030718a
      elif self.name == 'Devil, Imp':
        SLAstring = "CL6, at will - detect good, detect magic, invisibility (self only); 1/day - suggestion; 1/week - commune (six questions)"
      elif self.name == 'Guardinal, Lupinal':
        SLAstring = "CL8, at will - disguise self, blink, blur, darkness, ethereal jaunt; 3/day - cone of cold, cure light wounds, fly, magic missile"  # general 3.5 update booklet specifies change self is now disguise self http://archive.wizards.com/default.asp?x=dnd/dnd/20030718a
      elif self.name == 'Eladrin, Firre (fire pillar form)':
        SLAstring = "CL10, at will - detect thoughts, fireball, greater invisibility, persistent image, see invisibility, wall of fire; 1/day - prismatic spray" # Errata: Remove polymorph from spell-like abilities.
      elif self.name == 'Archon, Warden':
        SLAstring = "CL11, at will - aid, continual flame, detect scrying, detect thoughts, locate creature, scrying, see invisibility, true strike; 3/day - shield of the archons, true seeing"
    if self.name == "Pech":
        SLAstring = "CL4, 4/day - stone shape, stone tell; 1/day - wall of stone (requires four pechs), stone to flesh (requires eight pechs)"
    if SLAstring != '':
      SLAstring = SLAstring.strip().replace('; 1 day - ', '; 1/day - ').replace('; 1/day 0 ', '; 1/day - ').replace('; 1d3/day - ', '; 1/day - ')
      #SLAstring = SLAstring.replace('LotEM: 4th', '10/day').replace('LofPM - 1st', '10/day')
      if SLAstring[-1] == '.': SLAstring = SLAstring[:-1]
      SLAstring = SLAstring.replace(', ;', ';')
      # If no caster level is specified, the caster level is equal to the creature's Hit Dice.
      SLAstring = SLAstring.replace('CL unknown', 'CLHD')
      #print('split this:', SLAstring)
      split = freqChangeRE.split(SLAstring)
      if split[0] == '': split = split[1:]
      #if self.name == 'Eladrin, Coure': raise Exception(self.name, SLAstring, split)
      for new_level_or_new_freq_or_comma_separated_spells in split:
        #print('REsplit freqGroup =', new_level_or_new_freq_or_comma_separated_spells)
        if new_level_or_new_freq_or_comma_separated_spells is None: continue # capturing the group for caster level means ALWAYS capture that group even though it's optional, so when frequency is changed and caster level is not changed we'll get a None, annoying
        if new_level_or_new_freq_or_comma_separated_spells[:7].lower() == 'at will':
          currentUsesPerDay = 127
          continue
        if new_level_or_new_freq_or_comma_separated_spells[-5:] == '/hour':
          assert int(new_level_or_new_freq_or_comma_separated_spells[:-5]) == 1
          currentUsesPerDay = 24
          continue
        if new_level_or_new_freq_or_comma_separated_spells[-4:] == '/day':
          currentUsesPerDay = int(new_level_or_new_freq_or_comma_separated_spells[:-4])
          continue
        if new_level_or_new_freq_or_comma_separated_spells[-5:] == '/week':
          assert int(new_level_or_new_freq_or_comma_separated_spells[:-5]) == 1
          currentUsesPerDay = -7
          continue
        if new_level_or_new_freq_or_comma_separated_spells[-7:] == '/tenday':
          assert int(new_level_or_new_freq_or_comma_separated_spells[:-7]) == 1
          currentUsesPerDay = -10
          continue
        if new_level_or_new_freq_or_comma_separated_spells[-6:] == '/month':
          assert int(new_level_or_new_freq_or_comma_separated_spells[:-6]) == 1
          currentUsesPerDay = -30
          continue
        if new_level_or_new_freq_or_comma_separated_spells[-5:] == '/year':
          assert int(new_level_or_new_freq_or_comma_separated_spells[:-5]) == 1
          currentUsesPerDay = -127#-365
          continue
        if new_level_or_new_freq_or_comma_separated_spells[-8:] == '/century':
          assert int(new_level_or_new_freq_or_comma_separated_spells[:-8]) == 1
          currentUsesPerDay = -127#-36500
          continue
        matchObj = casterLevelRE.match(new_level_or_new_freq_or_comma_separated_spells)
        if matchObj is not None: # this is a new caster level
          assert matchObj.group(1) is not None
          currentCL = matchObj.group(1)
          continue
        # else match Obj is None
        for spell,parenthetical in parse_comma_separated_spells(new_level_or_new_freq_or_comma_separated_spells):
          if self.name == 'Varrangoin, Arcanist' and spell == 'polymorph':
            # Errata: replace polymorph other with baleful polymorph.
            spell = 'Baleful Polymorph'
          self.SpellLikeAbilities.append( (spell, currentCL, currentUsesPerDay, parenthetical) )
      """
      matchObj = casterLevelRE.match(SLAstring)
      assert matchObj is not None
      CorM = matchObj.group(1)
      assert CorM in ('C', 'M')
      defaultCL = matchObj.group(2)
      print('The default caster level for', self.name, 'is', defaultCL, 'but this can be overridden by individual clauses.')
      #frequencies = freqGroupRE.findall(SLAstring)
      #for i,matchObj in enumerate(frequencies): print('matchObj = ', matchObj)
      #freqGroups = [SLAstring[frequencies[i].end():frequencies[i+1].start()] for i,matchObj in enumerate(frequencies)]
      SLAfreqs = SLAstring.split('; ')
      for freqGroup in SLAfreqs:
        matchObj = casterLevelRE.match(freqGroup)
        if matchObj is not None:
          assert freqGroup[matchObj.end():matchObj.end()+2] == ', '
          thisClauseCL = matchObj.group(2)
          freqGroup = freqGroup[matchObj.end()+2:]
        else:
          thisClauseCL = defaultCL
        print('freqGroup =', freqGroup)
        freq, commaSeparatedSpells = freqGroup.split('-')
        freq = freq.strip()
        if freq == 'at will':
          usesPerDay = 127 # max can store in single-byte signed int
        else:
          assert freq[-4:] == '/day'
          usesPerDay = int(freq[:-4])
        spells = [spell.strip() for spell in noUnparenthesizedCommasRE.findall(commaSeparatedSpells)]
        #print('spells =', spells)
        for spell in spells:
          matchObj = parentheticalRE.search(spell)
          if matchObj is not None:
            assert matchObj.end() == len(spell)
            assert spell[matchObj.start()-1] == ' '
            parenthetical = matchObj.group(1)
            spell = spell[:matchObj.start()-1]
          else:
            parenthetical = None
          self.SpellLikeAbilities.append( (spell, thisClauseCL, usesPerDay, parenthetical) )
        """
        #quoteParens = commaSeparatedSpells.replace('(', '"(').replace(')', ')"')
        # Quoting parentheses does not work. I think csv only permits quotes at the beginning of a field, which makes csv near useless.
        #print('quoteParens =', quoteParens)
        # csv.reader(['goodbye, "this, is row 3", foo3'], skipinitialspace=True).__next__()
        # csv.reader(['disable "(30ft cone, 12HD*)", false sensory input "(five targets*)", mental disruption "(20ft radius*)", mindlink "(unwilling, 9 targets*)"'], skipinitialspace=True).__next__()
        #onlyOneRow = csv.reader([quoteParens], skipinitialspace=True).__next__() # must use skipinitialspace for quoting to work http://stackoverflow.com/questions/6879596/why-is-the-python-csv-reader-ignoring-double-quoted-fields
        #onlyOneRow = [spell.replace('"(', '(').replace(')"', ')') for spell in onlyOneRow]
        #print('onlyOneRow=', onlyOneRow)
  
  statblockMovementModeRE = re.compile('''(burrow|climb|fly|swim) (\d\d) ft\.(?: \((clumsy|poor|average|good|perfect)\))?''')
  assert statblockMovementModeRE.search(", fly 60 ft. (average)")
  assert len(list(statblockMovementModeRE.finditer(", fly 60 ft. (average)"))) == 1
  statblockRE = re.compile(r"([\w\, ]+): CR (\d|\d/\d{1,2}); (Fine|Diminutive|Tiny|Small|Medium|Large|Huge|Gargantuan|Colossal) ([\w ]+); HD (\d|\d/\d)d\d(?:(?:\-|\+)\d{1,2})?; hp \d{1,3}; Init \+\d; Spd (\d\d) ft\.((?:\, (?:burrow|climb|fly|swim) \d\d ft\.(?: \((?:clumsy|poor|average|good|perfect)\))?)*); AC (\d\d), touch (\d\d), flat-footed \d\d; Base Atk \+\d; Grp (?:\-|\+)\d{1,2}; Atk (?:\-|\+\d melee \([\d\-\+\, \w]+\)); Full Atk (?:\-|\+\d melee \([\d\-\+\, \w]+\)(?: and \+\d melee \([\d\-\+\, \w]+\))?); Space/Reach (?:\d|\d/\d) ft\./\d ft\.; SA ([\w\-]+); SQ ([\w \,\-]+); AL (\w{1,2}); SV Fort \+\d, Ref \+\d, Will \+\d; Str (\d{1,2})\, Dex (\d{1,2})\, Con (\d{1,2})\, Int (\d{1,2})\, Wis (\d{1,2})\, Cha (\d{1,2})\.\s+Skills and Feats: [\w \+\-\d\,]+; [\w \,]+\.")
  assert statblockRE.match("Owl, Medium: CR 1; Medium animal; HD 2d8; hp 13; Init +1; Spd 10 ft., fly 60 ft. (average); AC 14, touch 11, flat-footed 13; Base Atk +1; Grp +3; Atk +2 melee (1d4+2, talons); Full Atk +2 melee (1d4+2, talons) and +0 melee (1d6+1, bite); Space/Reach 5 ft./5 ft.; SA -; SQ -; AL N; SV Fort +4, Ref +4, Will +2; Str 14, Dex 13, Con 12, Int 2, Wis 14, Cha 4. Skills and Feats: Listen +14, Move Silently +19, Spot +14; Multiattack.") is not None
  assert statblockRE.match("Raven, Small: CR 1/2; Small animal; HD 1d8; hp 5; Init +1; Spd 10 ft., fly 40 ft. (average); AC 13, touch 12, flat-footed 12; Base Atk +0; Grp -7; Atk +2 melee (1d3-3, talons); Full Atk +2 melee (1d3-3, talons); Space/Reach 5 ft./5 ft.; SA -; SQ -; AL N; SV Fort +3, Ref +3, Will +2; Str 5, Dex 13, Con 12, Int 2, Wis 14, Cha 6. Skills and Feats: Listen +6, Spot +6; Weapon Finesse.") is not None
  @staticmethod
  def from_statblock(statblock, rulebook_abbrev):
    matchObj = Monster.statblockRE.match(statblock)
    if matchObj is None: raise ValueError(statblock)
    ret = object.__new__(Monster)
    ret.name = matchObj.group(1)
    ret.challenge_rating = fraction_to_negative(matchObj.group(2) )
    ret.size = matchObj.group(3)
    ret.type_name = sensible_title(matchObj.group(4) )
    ret.subtypes = []
    ret.HitDice = fraction_to_negative(matchObj.group(5) )
    ret.landSpeed = int(matchObj.group(6) )
    movementModes = matchObj.group(7)
    ret.movementModes = []
    for movementModeMatchObj in Monster.statblockMovementModeRE.finditer(movementModes):
      ret.movementModes.append( (movementModeMatchObj.group(1)[0], int(movementModeMatchObj.group(2))) )
    ret.maneuverability = None
    ret.ArmorClass = int(matchObj.group(8))
    ret.armorPlusShieldPlusNaturalArmor = ret.ArmorClass - int(matchObj.group(9))
    ret.attacks = ''
    ret.fullAttack = ''
    ret.special_damage = set()
    ret.specialAttacks = Monster.splitSpecialAbilities(matchObj.group(10) )
    ret.specialQualities = Monster.splitSpecialAbilities(matchObj.group(11) )
    ret.set_default_alignment(matchObj.group(12) )
    ret.strength     = integer_or_non(matchObj.group(13) )
    ret.dexterity    = integer_or_non(matchObj.group(14) )
    ret.constitution = integer_or_non(matchObj.group(15) )
    ret.intelligence = integer_or_non(matchObj.group(16) )
    ret.wisdom       = integer_or_non(matchObj.group(17) )
    ret.charisma     = integer_or_non(matchObj.group(18) )
    ret.rulebook_abbrev = rulebook_abbrev
    ret.environment = 'environment not given'
    ret.SpellLikeAbilities = []
    return ret

  def check_if_special_ability_is_spell(self, curs, monster_id, ability):
    if ability not in ('Poison', 'poison', 'death throes', 'Enslave', 'Camouflage', 'camouflage', 'Curse of lycanthropy', 'Darkvision', 'Low-Light Vision', 'scent', 'Spell Resistance', 'web', 'freeze'):
      # There's little real harm in mistakenly listing a monster as having an SLA if it has an almost-identical ability;
      # The only reason we want to eliminate things like scent and darkvision is because there's a large difference between being able to give those to others and having them only for yourself.
      spell_id = spell_name_to_id(curs, ability, True)
      if spell_id is not None:
        #print(self.name, 'special attack', attack)
        curs.execute('''INSERT INTO monster_has_spell_like_ability (monster_id, spell_id, caster_level, caster_level_scales_with_HD, uses_per_day, parenthetical) VALUES (?, ?, ?, ?, ?, ?)''', (monster_id, spell_id, self.HitDice, True, 127, 'listed as special ability, guessed to be SLA') )

  @staticmethod
  def sizeModifierToAttacksAndAC(size_id):
    return 0 if size_id == 5 else int(math.copysign(2**(abs(size_id - 5) - 1), 5 - size_id))

  def insert_into(self, curs):
    #print('rulebook_abbrev =', self.rulebook_abbrev)
    rulebook_abbrevs = [s.strip() for s in self.rulebook_abbrev.split(',')] # comma-separated list of rulebooks
    if len(rulebook_abbrevs) > 1:
      #print(self.name, 'has multiple rulebooks:', rulebook_abbrevs)
      rulebook_abbrevs = [next(a for a in rulebook_priority if a in rulebook_abbrevs)]
      #print(rulebook_abbrevs)
    for rulebook_abbrev in rulebook_abbrevs:
      if rulebook_abbrev in ('BoK','BoKR'): rulebook_abbrev = 'BoKr'
      elif re.match(r'\d\d\d', rulebook_abbrev): rulebook_abbrev = r'\d\d\d'
      elif re.match(r'A\d\d', rulebook_abbrev): rulebook_abbrev = r'A\d\d'
      curs.execute('''SELECT rulebook_id FROM rulebook_abbrev WHERE abbr=?;''',
                   (rulebook_abbrev,) )
      result = curs.fetchone()
      if result is None:
        raise Exception('no rulebook found:', self.name, rulebook_abbrev)
      rulebook_id = result[0]
      # if the same monster is listed in multiple books with the same stats...what?
      #print(rulebook_abbrev, 'rulebook_id from dnd_rulebook =', rulebook_id)
    self.size = self.size.rstrip('+.')
    # Colossal+ Although there is no size category larger than Colossal, the oldest epic dragons deal more damage with their attacks than other Colossal dragons, as shown on the Epic Dragon Face and Reach and Epic Dragon Attacks tables below.
    # So Colossal+ can be folded into Colossal.
    #curs.execute('''SELECT id from dnd_racesize WHERE name like ?''', (self.size + '%',) )
    #size_id = curs.fetchone()[0]
    size_id = id_from_name(curs, 'dnd_racesize', self.size, allowExtraOnRight=True)
    assert size_id is not None
    type_id = id_from_name(curs, 'dnd_monstertype', self.type_name)
    assert type_id is not None

    sizeModifierToAC = Monster.sizeModifierToAttacksAndAC(size_id)
    if self.dexterity is None:
      dexterityModifierToAC = 0 if self.name == 'Formian, Queen' else -5
    else:
      dexterityModifierToAC = self.dexterity//2 - 5
    if self.name == 'Titan':
      # This is actually correct; the titan is wearing half-plate armor, which has a max Dex bonus of +0.
      dexterityModifierToAC = min(0, dexterityModifierToAC)
    if 'Incorporeal' in self.subtypes or 'unearthly grace' in self.specialQualities:
      charismaBonusToAC = max(1, self.charisma//2 - 5)
    else:
      charismaBonusToAC = 0
    #if self.ArmorClass != 10 + self.armorPlusShieldPlusNaturalArmor + sizeModifierToAC + dexterityModifierToAC + charismaBonusToAC:
    #  print('unaccounted-for AC:' + self.name + ''.join(self.subtypes) + 'AC {} != 10 + {}armor + {}size + {}Dex + {}Cha'.format(self.ArmorClass, self.armorPlusShieldPlusNaturalArmor, sizeModifierToAC, dexterityModifierToAC, charismaBonusToAC) )

    curs.execute('''INSERT INTO dnd_monster
                 (name, rulebook_id, size_id, type_id, hit_dice, land_speed, strength, dexterity, constitution, intelligence, wisdom, charisma, natural_armor_bonus, challenge_rating)
                 VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                 (self.name, rulebook_id, size_id, type_id, self.HitDice, self.landSpeed, self.strength, self.dexterity, self.constitution, self.intelligence, self.wisdom, self.charisma, self.armorPlusShieldPlusNaturalArmor, self.challenge_rating) )
    monster_id = curs.lastrowid

    for subtype in self.subtypes:
      if subtype[:9] == 'Augmented': subtype = 'Augmented'
      if subtype == self.name: continue
      subtype_id = id_from_name(curs, 'dnd_monstersubtype', subtype)
      if subtype_id is None: raise IndexError(self.name + subtype)
      curs.execute('''INSERT INTO monster_has_subtype (monster_id, subtype_id) VALUES (?, ?);''', (monster_id, subtype_id) )
    
    curs.executemany('''INSERT INTO monster_movement_mode (monster_id,abbrev,speed) VALUES (?,?,?);''', [(monster_id,a,s) for (a,s) in self.movementModes])
    if self.maneuverability == 'clu':
      curs.execute('''INSERT INTO monster_maneuverability (monster_id,maneuverability) VALUES (?,?);''', (monster_id, 1) )
    elif self.maneuverability == 'pr':
      curs.execute('''INSERT INTO monster_maneuverability (monster_id,maneuverability) VALUES (?,?);''', (monster_id, 2) )
    elif self.maneuverability == 'avg':
      curs.execute('''INSERT INTO monster_maneuverability (monster_id,maneuverability) VALUES (?,?);''', (monster_id, 3) )
    elif self.maneuverability == 'gd':
      curs.execute('''INSERT INTO monster_maneuverability (monster_id,maneuverability) VALUES (?,?);''', (monster_id, 4) )
    elif self.maneuverability == 'prf' or self.maneuverability == 'prft':
      curs.execute('''INSERT INTO monster_maneuverability (monster_id,maneuverability) VALUES (?,?);''', (monster_id, 5) )

    for naturalWeaponMatchObj in naturalWeaponAttackRE.finditer(self.attacks):
      weapon_id = id_from_name(curs, 'dnd_weapon', naturalWeaponMatchObj.group(1).lower())
      # check for None?
      # Frostburn urskan will match claw twice: Steelclaw +11 (1d12+7) or claw +11 (1d8+7)
      # We could insist on the natural weapon name being either at the start of the string or having space before it...but honestly,
      # if it has a prefix it's probably pretty much the same thing, so we don't want to not match that in general.
      # We can just allow multi-match.
      # A camel's bite is treated as a secondary natural attack and adds only half the camel's Strength bonus to damage.
      attackBonus = int(naturalWeaponMatchObj.group(2))
      if self.name == 'Owlbear': attackBonus = 9 # error in table
      elif self.name == 'Demon, Sorrowsworn': attackBonus = 23 # error in table, but even getting 23 from MM3 not clear what's going on
      elif self.name == 'Corpse Gatherer': attackBonus = 24 # per 3.5 update
      elif self.name == 'Bat, Guard': attackBonus = 5 # error in table
      elif self.name == 'Demon, Carnage': attackBonus = 8 # error in table
      # Abyssal Drake appears to be an error in the Draconomicon
      curs.execute('''SELECT base_attack_per_4HD FROM dnd_monstertype WHERE id=?;''', (type_id,) )
      BAB = curs.fetchone()[0]*max(0, self.HitDice)//4
      strengthModifier = 0 if self.strength is None else self.strength//2 - 5
      attackBonusWithoutBaseAndSize = attackBonus - BAB - sizeModifierToAC
      initialDamageMatchObj = numericalDamageRE.match(naturalWeaponMatchObj.group(3))
      # might not be able to find numerical damage, e.g. Tentacle +3 (paralysis) paralysis
      if initialDamageMatchObj is None:
        initialDamage = '0+0'
      else:
        initialDamage = initialDamageMatchObj.group(0)
      #if '+' not in initialDamage and '-' not in initialDamage: initialDamage += '+0'
      #print(initialDamage)
      if initialDamageMatchObj is None or initialDamageMatchObj.group(1) is None: # if no damage roll, then it's zero
        numberOfWeaponDamageDice = 0
        weaponDamageDieSize = 0
      else:
        damageRollAsSingletonOrPair = [int(n) for n in initialDamageMatchObj.group(1).split('d')]
        assert len(damageRollAsSingletonOrPair) in (1,2)
        if len(damageRollAsSingletonOrPair) == 1:
          numberOfWeaponDamageDice = 1
        else:
          numberOfWeaponDamageDice = damageRollAsSingletonOrPair[0]
        weaponDamageDieSize = damageRollAsSingletonOrPair[-1]
      if initialDamageMatchObj is None or initialDamageMatchObj.group(2) is None: # if no damage modifier, then it's +0
        damageModifier = 0
      else:
        damageModifier = int(initialDamageMatchObj.group(2)) #int(re.split(r'[+-]', initialDamage)[1])
      if damageModifier != 0:
        pass
      if attackBonusWithoutBaseAndSize > strengthModifier:
        if attackBonusWithoutBaseAndSize == self.dexterity//2 - 5:
          # suspect Weapon Finesse
          pass
        elif attackBonusWithoutBaseAndSize == strengthModifier + 1:
          # suspect Weapon Focus feat
          pass
        elif attackBonusWithoutBaseAndSize == self.dexterity//2 - 5 + 1:
          # suspect Weapon Finesse *and* Weapon Focus, as greater fire elementals have
          pass
        # Enhanced Bite (Ex): A spellgaunt's fangs function as +4 weapons.
        # A blighted bloodfire ooze has the traits of an outsider but the features (including BAB) of an ooze.
        #else:
          # what happened to Bodak? it was showing up before...ah, Bodak *looks* like a Weapon Finesse case though in fact it's Weapon Focus (slam)
          # Nightshades all make liberal use of Haste, which they have at-will, so it looks like Haste might already be included in their attack bonus?
          # Pseudodragon has Weapon Finesse so should have a total attack bonus of +6, but only has +4...no explanation yet
        #  print('unaccounted-for attack bonus of size {}:'.format(attackBonusWithoutBaseAndSize - strengthModifier) + rulebook_abbrev + self.name + ''.join(self.subtypes) + '+{} != {}BAB + {}size + {}Str/Dex'.format(attackBonus, BAB, sizeModifierToAC, max(strengthModifier, self.dexterity//2 - 5)) )
      elif attackBonusWithoutBaseAndSize < strengthModifier:
        # MM3 Charnel Hound has -5 penalty for Power Attack included, ditto MM3 Drowned, MM3 Necronaut, MM3 Odopi
        if strengthModifier - attackBonusWithoutBaseAndSize in (damageModifier - strengthModifier, damageModifier - strengthModifier*3//2):
          # suspect Power Attack
          pass
        elif strengthModifier - attackBonusWithoutBaseAndSize == 5 and damageModifier == strengthModifier//2:
          # suspect secondary natural weapon
          pass
        # Powerful Slam (Ex) A goristro’s slam attacks are treated as if they were two-handed weapons for purposes of applying modifiers to damage with Power Attack and from its Strength bonus.
        #else:
        #  print('unaccounted-for attack penalty of size {}:'.format(attackBonusWithoutBaseAndSize - strengthModifier) + rulebook_abbrev + self.name + ''.join(self.subtypes) + '+{} != {}BAB + {}size + {}Str'.format(attackBonus, BAB, sizeModifierToAC, strengthModifier) )
        #  print(initialDamage, damageModifier, strengthModifier)
      curs.execute('''INSERT OR IGNORE INTO monster_has_natural_weapon (monster_id, weapon_id, number_of_damage_dice, damage_die_size) VALUES (?, ?, ?, ?);''', (monster_id, weapon_id, numberOfWeaponDamageDice, weaponDamageDieSize) )
    for naturalWeaponMatchObj in naturalWeaponRE.finditer(self.fullAttack):
      # Secondary natural weapons, such as a rulkar madclaw's claws, appear only under full attack (this is the only time attack bonus and damage will be listed in the full attack column)
      weapon_id = id_from_name(curs, 'dnd_weapon', naturalWeaponMatchObj.group(1).lower())
      curs.execute('''INSERT OR IGNORE INTO monster_has_natural_weapon (monster_id, weapon_id, number_of_damage_dice, damage_die_size) VALUES (?, ?, ?, ?);''', (monster_id, weapon_id, 0, 0) )
    # When a creature has more than one natural weapon, one of them (or sometimes a pair or set of them) is the primary weapon. All the creature's remaining natural weapons are secondary.
    # The primary weapon is given in the creature's Attack entry, and the primary weapon or weapons is given first in the creature's Full Attack entry.
    # An attack with a primary natural weapon uses the creature's full attack bonus. Attacks with secondary natural weapons are less effective and are made with a -5 penalty on the attack roll, no matter how many there are. (Creatures with the Multiattack feat take only a -2 penalty on secondary attacks.) This penalty applies even when the creature makes a single attack with the secondary weapon as part of the attack action or as an attack of opportunity.
    # This suggests that the primary natural weapon should go in the dnd_monster table itself, while secondaries should go in the auxiliary table.
    # Almost every monster will have one primary natural weapon, right?
    for dmg in self.special_damage:
      curs.execute('''INSERT OR IGNORE INTO monster_deals_special_damage (monster_id, damage) VALUES (?, ?);''', (monster_id, dmg) )

    curs.execute('''INSERT INTO monster_has_alignment (monster_id, good_evil, law_chaos) VALUES (?,?,?);''', (monster_id, self.goodEvilID, self.lawChaosID) )
    #curs.execute('''SELECT id from dnd_plane WHERE name like "%{}";'''.format(self.environment) )
    #result = curs.fetchone()
    planeAndLocationRE = re.compile(r"([\w\s]+)\(([\w\s\']+)\)")
    matchObj = planeAndLocationRE.match(self.environment)
    if matchObj is not None:
      parentPlane = matchObj.group(1).strip()
      locationWithin = matchObj.group(2)
      #print(self.environment, parentPlane, locationWithin)
      if locationWithin == "Ice Wastes": locationWithin = "Iron Wastes"
      elif locationWithin in ('Brine Flats', 'Screaming Jungle'):
        # these are locations within the Gaping Maw, but for the moment we just want to store the layer
        locationWithin = 'Gaping Maw'
      elif locationWithin == 'Empyrea':
        # Empyrea, the City of Tempered Souls, on Celestia's fifth layer Mertion
        locationWithin = 'Mertion'
      elif locationWithin == "Stygia":
        if parentPlane == 'The Abyss':
          parentPlane = 'Baator'
      numberedLayerRE = re.compile(r"((?:\d\d\d)|(?:Six))(?:rd)|(?:th) Layer")
      matchObj = numberedLayerRE.match(locationWithin)
      if matchObj is not None:
        if matchObj.group(1) == 'Six':
          layerNumber = 6
        else:
          layerNumber = int(matchObj.group(1) )
        curs.execute('''INSERT INTO monster_on_plane SELECT ?,id from dnd_plane WHERE dnd_plane.layer_number=?''', (monster_id, layerNumber) )
      for name in (parentPlane, locationWithin):
        if name == "Azzagrat":
          # not sure how to handle cases that return multiple results for the same name, since Azzagrat will return three results
          curs.execute('''INSERT INTO monster_on_plane SELECT ?,id from dnd_plane WHERE dnd_plane.name=?''', (monster_id, locationWithin) )
        else:
          plane_id = id_from_name(curs, 'dnd_plane', name, True)
          if plane_id is not None:
            #print(name, 'is a plane')
            curs.execute('''INSERT INTO monster_on_plane (monster_id, plane_id) VALUES (?,?);''', (monster_id, plane_id) )
    else:
        plane_id = id_from_name(curs, 'dnd_plane', self.environment, True)
        if plane_id is not None:
          curs.execute('''INSERT INTO monster_on_plane (monster_id, plane_id) VALUES (?,?);''', (monster_id, plane_id) )
    if 'Cold' in self.environment or 'cold' in self.environment:
      curs.execute('''INSERT INTO monster_in_climate (monster_id, climate) VALUES (?,-1);''', (monster_id,) )
    elif 'Temperate' in self.environment or 'temperate' in self.environment:
      curs.execute('''INSERT INTO monster_in_climate (monster_id, climate) VALUES (?,0);''', (monster_id,) )
    elif 'Warm' in self.environment or 'warm' in self.environment:
      curs.execute('''INSERT INTO monster_in_climate (monster_id, climate) VALUES (?,1);''', (monster_id,) )
    for name in TERRAIN_NAMES:
      if name in self.environment:
        curs.execute('''INSERT INTO monster_on_terrain (monster_id, terrain_id) VALUES (?,?);''', (monster_id, id_from_name(curs, 'dnd_terrain', name) ) )

    for attack in self.specialAttacks:
      #if attack == '-': continue
      ability_id = insert_if_needed(curs, 'dnd_special_ability', attack, special_attack=1)
      curs.execute('''INSERT INTO monster_has_special_ability (monster_id, special_ability_id) VALUES (?, ?);''', (monster_id, ability_id) )
      self.check_if_special_ability_is_spell(curs, monster_id, attack)
      if insert_monster_casts_spells(curs, monster_id, attack):
        attack = "Spells"
      #if attack not in ('Poison', 'poison', 'death throes', 'Enslave', 'Camouflage', 'camouflage', 'Curse of lycanthropy'):
      #  spell_id = spell_name_to_id(curs, attack, True)
      #  if spell_id is not None:
          #print(self.name, 'special attack', attack)
      #    curs.execute('''INSERT INTO monster_has_spell_like_ability (monster_id, spell_id, caster_level, caster_level_scales_with_HD, uses_per_day, parenthetical) VALUES (?, ?, ?, ?, ?, ?)''', (monster_id, spell_id, self.HitDice, True, 127, 'listed as special attack, guessed to be SLA') )
    for quality in self.specialQualities:
      #if quality == '': continue
      if quality == "LLV": quality = "Low-Light Vision"
      elif darkvisionRE.match(quality): quality = "Darkvision"
      elif insert_damage_reduction(curs, monster_id, quality):
        quality = "Damage Reduction"
      elif insert_spell_resistance(curs, monster_id, quality): quality = "Spell Resistance"
      elif quality == 'immunity to magic' or quality == 'magic immunity' or quality == 'immunity to cold and magic':
        curs.execute('''INSERT INTO monster_has_spell_resistance (monster_id, resistance) VALUES (?,?);''', (monster_id, 127) )
      elif insert_fast_healing(curs, monster_id, quality): quality = "Fast Healing"
      #if 'sensitiv' in quality: print(self.name, quality)
      ability_id = insert_if_needed(curs, 'dnd_special_ability', quality, special_attack=0)
      curs.execute('''INSERT INTO monster_has_special_ability (monster_id, special_ability_id) VALUES (?, ?);''', (monster_id, ability_id) )
      self.check_if_special_ability_is_spell(curs, monster_id, quality)
      #if quality not in ('Darkvision', 'Low-Light Vision', 'scent', 'Spell Resistance', 'poison', 'Poison', 'web', 'death throes', 'freeze'):
      #  spell_id = spell_name_to_id(curs, quality, True)
      #  if spell_id is not None:
          #print(self.name, 'special quality', quality)
      #    curs.execute('''INSERT INTO monster_has_spell_like_ability (monster_id, spell_id, caster_level, caster_level_scales_with_HD, uses_per_day, parenthetical) VALUES (?, ?, ?, ?, ?, ?)''', (monster_id, spell_id, self.HitDice, True, 127, 'listed as special quality, guessed to be SLA') )
    for (spellName, CL, usesPerDay, parenthetical) in self.SpellLikeAbilities:
      # probably want date of monster rulebook and take latest
      try:
        spell_id = spell_name_to_id(curs, spellName)
      except KeyError as E:
        if 'mephit' in self.name.lower(): continue
        else: raise
      # http://www.d20srd.org/srd/specialAbilities.htm#spellLikeAbilities
      # If no caster level is specified, the caster level is equal to the creature's Hit Dice.
      if CL[:2] == 'HD' or CL[:3]=='lvl':
        caster_level_scales_with_HD = True
        if 'min' in CL:
          assert re.compile("\(min(\d)\)").match(CL[-6:])
          casterLevel = int(CL[-2])
        else:
          casterLevel = self.HitDice
          if '/' in CL:
            assert CL[:4] == 'HD/2'
            casterLevel = self.HitDice/2
      else:
        casterLevel = int(CL[:2])
        caster_level_scales_with_HD = False
      curs.execute('''INSERT INTO monster_has_spell_like_ability (monster_id, spell_id, caster_level, caster_level_scales_with_HD, uses_per_day, parenthetical)
                   VALUES (?, ?, ?, ?, ?, ?)''',
                   (monster_id, spell_id, casterLevel, caster_level_scales_with_HD, usesPerDay, parenthetical) )

# Creatures who could not be played in a default-campaign-setting game without adaptation have been omitted. This largely means outsiders native to planes unique to other settings.

class Template(object):
  #bookREcapturing = re.compile(r" \(([\w\d]{2,4}) (\d{2,3})\)")
  augmentedRE = re.compile(r"; (?:gains? )?[aA]ugmented subtype| with the appropriate augmented subtype")
  rockdeworldRE = re.compile(r"([\w\- \,]+)(?: \((?:(?:([\w\d]{2,5}) (\d{2,3}))|(?:Web))\))? +LA ((?:n/a)|(?:Variable)|(?:varies)|(?:\+|\-)\d)(?: \(minimum 0\))? -{1,3}([\w\d\- \,\.:;\(\)]+) (?:(?:Type Changes)|(?:TC)): +([\w ;\,\[\]\-/]+)(?: \([\.\d\w ]+\))?")
  exceptTypesRE = re.compile(r" (?:except|that is not) an? ([\w\, ]+)")
    # Any living creature except an ooze
    # Any corporeal creature except a construct, undead, or elemental
    # any living, corporeal creature, except a dragon
    # but the word except appears for other reasons as well
    # Any corporeal that is not an oustsider
    # Any corporeal creature that is not a construct
    # Any living creature of good alignment that is not an outsider or an elemental
  changeOnlyTheseTypesRE = re.compile(r"Unchanged;?(?: add (?:Psionic|Dragonblood) subtype)?; ([Aa]nimal)s?(?:(?: and |/| or )(\w+)s?)? becomes? ([Mm]agical [Bb]east)s?")
  # Winged Creature SS 137 type not found: Unchanged; Animals and vermin become magical beasts; Humanoids become monstrous humanoids
  # Beasts or Animals become Magical Beasts otherwise type is unchanged but can ignore that in favor of newer version of each such template

  @staticmethod
  def from_groups(groups):
      name = groups[0]
      assert name != ''
      rulebook_abbrev = groups[1]
      page = groups[2]
      if page is None: assert rulebook_abbrev is None
      if page == '': assert rulebook_abbrev == ''
      if rulebook_abbrev is None or rulebook_abbrev == '':
        assert page is None or page == ''
        rulebook_abbrev = 'Web'
        assert page is None
      else:
        #print('page =', page)
        page = int(page)
      #location = groups[1]
      #if location == '':
      #  rulebook_abbrev = 'Web'
      #else:
      #  matchObj = Template.bookREcapturing.match(location)
      #  assert matchObj is not None
      #  rulebook_abbrev = matchObj.group(1)
      #  page = int(matchObj.group(2) )
      levelAdjustment = groups[3]
      # Note that if the template doesn't mention level adjustment, then the level adjustment should be +0, not -.
      # However, half-golems and captured ones were given LA - in the 3.5 update to MM2.
      if levelAdjustment in ('n/a', 'Variable', 'varies'):
        levelAdjustment = None
      else:
        levelAdjustment = int(levelAdjustment)
      baseCreature = groups[4]
      resultType = groups[5]
      #print('resultType =', resultType)
      # According to Monster Manual rules, there is no point in specifically noting the augmented subtype when given: If a template changes the base creature's type, the creature also acquires the augmented subtype (see page 306) unless the template description indicates otherwise. The augmented subtype is always paired with the creature's original type. For example, a unicorn with the half-celestial template is an outsider with the augmented magical beast subtype. Unless a template indicates otherwise, the new creature has the traits of the new type but the features of the original type. The example half-celestial unicorn has outsider traits (darkvision out to 60 feet, no soul) and magical beast features (10-sided Hit Dice, base attack bonus equal to total Hit Dice, good Fortitude and Reflex saves, and so on).
      if name == 'Vampire': resultType = 'Undead'
      resultType = Template.augmentedRE.sub('', resultType)
      #print(name, rulebook_abbrev, levelAdjustment, baseCreature, resultType)
      ret = Template()
      ret.name = name
      ret.rulebook_abbrev = rulebook_abbrev
      ret.page = page
      ret.levelAdjustment = levelAdjustment
      ret.baseCreature = baseCreature.strip()
      ret.resultType = resultType.strip()
      return ret
  @staticmethod
  def findall(string):
    for groups in Template.rockdeworldRE.findall(string):
      yield Template.from_groups(groups)
  @staticmethod
  def read_file(filename):
    for line in open(filename):
      if line[:6] == '</div>' or line[:11] == '<div class=' or line[:3] == '<b>':
        continue
      matchObj = Template.rockdeworldRE.match(line)
      if matchObj is None:
        raise Exception(line)
      yield Template.from_groups(matchObj.groups() )
  def insert_type_change(self, curs, baseTypeID, outputTypeID):
    if self.resultType is None:
      assert outputTypeID is None
      assert baseTypeID is not None
      outputTypeID = baseTypeID
    #print(self.name, 'outputTypeID =', outputTypeID)
    if outputTypeID is None:
      raise Exception(self.name, self.resultType)
    curs.execute('''INSERT INTO template_type (template_id, base_type, output_type) VALUES (?,?,?);''', (self.template_id, baseTypeID, outputTypeID) )
  def insert_into(self, curs):
    curs.execute('''SELECT rulebook_id from rulebook_abbrev where abbr=?;''',
                 (self.rulebook_abbrev,) )
    result = curs.fetchone()
    if result is None:
      print('no rulebook found:', self.name, self.rulebook_abbrev)
    rulebook_id = result[0]

    self.baseCreature = self.baseCreature.replace('oustsider', 'outsider')
    self.resultType = self.resultType.replace('animl', 'animal')

    curs.execute('''INSERT INTO dnd_template
                 (name, rulebook_id, page, level_adjustment)
                 VALUES (?, ?, ?, ?);''',
                 (self.name, rulebook_id, self.page, self.levelAdjustment) )
    self.template_id = curs.lastrowid
    
    # A character who succeeds at all the saves he or she is required to make takes on the attributes of a half-golem as described below except that the character retains his or her alignment, gains a +4 bonus to Constitution, and does not change type or gain construct traits. As soon as the character fails one of these required saves, he or she becomes a half-golem of neutral evil alignment. The character then has no Constitution score and the character's type changes to construct, granting him or her construct traits.
    # entropic Size and Type: Unless the creature was undead, its type changes to outsider. Size is unchanged. The creature also gains the extraplanar subtype.
    outputTypeID = None
    for (baseTypeID,name) in iterate_over_types(curs):
      if self.resultType is not None and (self.resultType.startswith(name + ';') or self.resultType.startswith(name + ', retains') ):
        outputTypeID = baseTypeID
    if self.resultType == 'Unchanged' or self.resultType == 'none':
      self.resultType = None
    # Type Changes: Unchanged; Animals become Magical Beast [augmented animl]; gain psionic subtype
    elif (self.resultType[:11] == 'Unchanged; ' or self.resultType=='gain Aquatic subtype') and ('ubtype' in self.resultType or 'Add [Shapechanger]' in self.resultType) and 'become' not in self.resultType:
      self.resultType = None # ignore subtypes for now
    elif outputTypeID is None:
      outputTypeID = id_from_name(curs, 'dnd_monstertype', self.resultType)
      # allow outputTypeID to be None for now
      #if outputTypeID is None:
      #  print('template', self.name, self.rulebook_abbrev, self.page, 'type not found:', self.resultType)
      #  raise Exception(self.resultType)

    allTypeNames = set(pair[1] for pair in iterate_over_types(curs) )
    baseTypeNames = set()

    # remember, outsiders are living: An outsider is at least partially composed of the essence (but not necessarily the material) of some plane other than the Material Plane. Unlike most other living creatures, an outsider does not have a dual nature --- its soul and body form one unit.
    # thornier is dealing with Living Constructs, since that's a subtype, but those are few and can be neglected for now
    matchObj = Template.exceptTypesRE.search(self.baseCreature)
    if matchObj is not None: # base has an except-types clause
      exceptTypesList = matchObj.group(1)
      #print(self.name, self.baseCreature, exceptTypesList)
      baseTypeNames = set(name for name in allTypeNames if name.lower() not in exceptTypesList.lower() and ('ny living' not in self.baseCreature or name!='Undead' or name!='Construct') )
      #for (baseTypeID,name) in iterate_over_types(curs):
      #  if name.lower() not in exceptTypesList.lower() and ('ny living' not in self.baseCreature or name!='Undead' or name!='Construct'):
          #if outputTypeID is None: print(self.name)
      #    self.insert_type_change(curs, baseTypeID, outputTypeID)
      #return
    elif self.baseCreature[:16] == 'Any Non-undead, ':
      baseTypeNames = set(name for name in allTypeNames if 'non' + name.lower() not in self.baseCreature.lower() )
      #for (baseTypeID,name) in iterate_over_types(curs):
      #  if 'non' + name.lower() not in self.baseCreature.lower():
      #    self.insert_type_change(curs, baseTypeID, outputTypeID)
      # sqlite> select dnd_template.name,base.name,output.name from template_type inner join dnd_template on template_id=dnd_template.id inner join dnd_monstertype as base on base_type=base.id inner join dnd_monstertype as output on output_type=output.id;
      #return
    # already eliminated the excepts, so if any type name is mentioned, we have a list of allowed type names
    elif any(name.lower() in self.baseCreature.lower() for name in allTypeNames):
      for name in allTypeNames:
        if name.lower() in self.baseCreature.lower() and 'generally ' + name.lower() + ' form' not in self.baseCreature.lower():
          baseTypeNames.add(name)
    elif self.baseCreature == 'Any Lich':
      baseTypeNames = set(['Undead'])
    elif self.baseCreature == 'any naga':
      baseTypeNames = set(['Aberration', 'Undead', 'Monstrous Humanoid', 'Humanoid'])
      # sqlite> select dnd_monstertype.name,dnd_monster.name from dnd_monster inner join dnd_monstertype on type_id=dnd_monstertype.id where dnd_monster.name like "%naga%";
      # Undead|Naga, Bone Humanoid|Naga, Shinomen, Greensnake Monstrous Humanoid|Naga, Shinomen, Chameleon
    else:
      baseTypeNames = allTypeNames
    # living is always used as a requirement
    if 'living' in self.baseCreature and 'living construct' not in self.baseCreature:
      baseTypeNames = baseTypeNames.difference(set(['Undead', 'Construct']) )
      # except living constructs...

    if 'Draconic' in self.name or self.name == "Shadow Creature": assert self.resultType is not None
    if self.resultType is not None:
      matchObj = Template.changeOnlyTheseTypesRE.match(self.resultType)
      if ('Draconic' in self.name or self.name == "Shadow Creature") and matchObj is None:
        raise Exception(self.name, self.resultType)
      if matchObj is not None:
        twoOrThree = matchObj.groups()
        #print(self.name, 'twoOrThree =', twoOrThree)
        assert len(twoOrThree) in (2,3)
        twoOrThree = [sensible_title(typeName) for typeName in twoOrThree] # make sure match for .difference() later
        for baseTypeName in twoOrThree[:-1]:
          if baseTypeName is None: continue
          #print('changeOnlyTheseTypesRE', len(twoOrThree), baseTypeName, twoOrThree[-1])
          #print(self.name, baseTypeName, twoOrThree[-1])
          self.insert_type_change(curs, id_from_name(curs, 'dnd_monstertype', baseTypeName), id_from_name(curs, 'dnd_monstertype', twoOrThree[-1]) )
        for name in baseTypeNames.difference(twoOrThree[:-1]):
          sameID = id_from_name(curs, 'dnd_monstertype', name)
          self.insert_type_change(curs, sameID, sameID)
        return

    if self.resultType == 'Outsider unless the base creature was undead':
      undeadID = id_from_name(curs, 'dnd_monstertype', 'Undead')
      assert undeadID is not None
      self.insert_type_change(curs, undeadID, undeadID)
      for name in baseTypeNames.difference(['Undead']):
        self.insert_type_change(curs, id_from_name(curs, 'dnd_monstertype', name), id_from_name(curs, 'dnd_monstertype', 'Outsider') )
      return
    elif self.resultType == 'Humanoid if medium-size or smaller, or Giant if it is Large or larger':
      for name in baseTypeNames:
        self.insert_type_change(curs, id_from_name(curs, 'dnd_monstertype', name), id_from_name(curs, 'dnd_monstertype', 'Humanoid') )
        self.insert_type_change(curs, id_from_name(curs, 'dnd_monstertype', name), id_from_name(curs, 'dnd_monstertype', 'Giant') )
      return
    
    # if didn't hit any other case:
    #print(self.name, 'default case', self.resultType)
    for name in baseTypeNames:
      self.insert_type_change(curs, id_from_name(curs, 'dnd_monstertype', name), outputTypeID)

assert Template.changeOnlyTheseTypesRE.match(r"Unchanged; Animals become Magical Beast [augmented animl]; gain psionic subtype")
assert Template.changeOnlyTheseTypesRE.match(r"Unchanged; animals or vermin become magical beasts")




standardFamiliarREstring = r'\n(?P<name>[\w ,]+)\t\|<div style="text-align: center;">\t(?P<monsterRulebook>[A-Za-z]+) p(?P<monsterPage>\d+)\t</div>\|<div style="text-align: center;">\t(?P<asFamiliarRulebook>\w+|#323|#341) p(?P<asFamiliarPage>\d+)\t</div>\|{1,2}<div style="text-align: center;">\t(?:(?P<spellcasterLevel>\d+)(?:st|nd|rd|th)|[A-Za-z ]+(?:\+\d)?[A-Za-z ]+)\t</div>\|(?P<restOfLine>.*)'
realNameRE = re.compile(r'<div style="text-align: center;">\tStat: (?P<realName>[A-Za-z]+)')
standardFamiliarRE = re.compile(standardFamiliarREstring)
assert standardFamiliarRE.match('''
Bat	|<div style="text-align: center;">	MMI p268	</div>|<div style="text-align: center;">	PHB p52	</div>||<div style="text-align: center;">	Masters gains a +3 on Listen checks	</div>||<br />''')
assert realNameRE.search('''<div style="text-align: center;">	Masters gains a +1 to AC when prone or behind cover	</div>|<div style="text-align: center;">	Stat: Hedgehog without poison quills	</div>|<br />''')
def make_familiar_table(curs):
  print('make_familiar_table')
  curs.execute('''CREATE TABLE dnd_familiar (
  monster_id INTEGER NOT NULL
  ,alternate_name TEXT DEFAULT NULL
  ,rulebook_id INTEGER NOT NULL
  ,page unsigned smallint(3) NOT NULL
  ,prereq_spellcaster_level unsigned tinyint(2) NOT NULL DEFAULT 0
,FOREIGN KEY(monster_id) REFERENCES dnd_monster(id)
,FOREIGN KEY(rulebook_id) REFERENCES dnd_rulebook(id)
  );''')
  for matchObj in standardFamiliarRE.finditer(open('familiars_daremetoidareyo.txt', 'r').read()):
    #print(matchObj.group(0))
    givenName = matchObj.group('name')
    if (givenName == 'Celestial standard familiar' or givenName == 'Fiendish standard familiar' or
        givenName == 'Hairy spider' or 'Guardian' in givenName or 'Monstrous spider' in givenName or
        givenName == 'Night Hunter Bat' or givenName == 'Lizard, spitting crawler' or
        givenName == 'Tressym'):
      continue
    if givenName == 'Sea snake':
      givenName = 'Sea Snake, Tiny'
    elif givenName == 'Quasit':
      givenName = 'Demon, Quasit'

    monsterRulebookID = get_rulebook_id(curs, matchObj.group('monsterRulebook'))
    if monsterRulebookID is None:
      raise Exception('Rulebook' + matchObj.group('monsterRulebook') + 'not found!')
    asFamiliarRulebookID = get_rulebook_id(curs, matchObj.group('asFamiliarRulebook'))
    if asFamiliarRulebookID is None:
      raise Exception('Rulebook' + matchObj.group('asFamiliarRulebook') + 'not found!')

    name = MonsterName(givenName)
    #print(matchObj.group('name'), matchObj.group('monsterRulebook'), matchObj.group('monsterPage'))
    if givenName == 'Fish Owl':
      searchForName = 'Owl'
    elif givenName == 'Great Horned Owl':
      searchForName = 'Owl'
    elif givenName == 'Parrot':
      searchForName = 'Raven'
    elif givenName == 'Gyrfalcon':
      searchForName = 'Hawk'
    elif givenName == 'Lemming':
      searchForName = 'Rat'
    elif givenName == 'Snowy owl':
      searchForName = 'Owl'
    elif givenName == 'Puffin':
      searchForName = 'Raven'
    elif givenName == 'King Cobra':
      searchForName = 'Snake, Viper, Small'
    else:
      searchForName = name.ezkajii()
      givenName = None
      
    monsterID = id_from_name(curs, 'dnd_monster', searchForName, additionalCriteria=(('rulebook_id', monsterRulebookID),))
    if monsterID is None:
      realNameMatch = realNameRE.search(matchObj.group('restOfLine'))
      #print('real name?', realNameMatch, matchObj.group('restOfLine'))
      if realNameMatch is not None:
        givenName = searchForName
        searchForName = realNameMatch.group('realName')
        monsterID = id_from_name(curs, 'dnd_monster', searchForName, additionalCriteria=(('rulebook_id', monsterRulebookID),))
    if monsterID is None:
      raise Exception(searchForName + ' not found in ' + matchObj.group('monsterRulebook'))

    level = int(matchObj.group('spellcasterLevel')) if matchObj.group('spellcasterLevel') else 0

    curs.execute('''INSERT INTO dnd_familiar (monster_id, alternate_name, rulebook_id, page, prereq_spellcaster_level) VALUES (?,?,?,?,?);''',
                 (monsterID, givenName, asFamiliarRulebookID, int(matchObj.group('asFamiliarPage') ), level) )
  print('done make_familiar_table')





advancementRE = re.compile(r'(\d{1,3})(\+|-\d{1,3})?(?: HD)? \((Diminutive|Tiny|Small|Medium|Large|Huge|Gargantuan|Colossal)\)')
assert advancementRE.match('2-3 HD (Small)')
assert advancementRE.match('19-36 HD (Huge)')

def make_skill_table(conn, curs):
  with open('advancement.sql') as advancementFile:
    curs.executescript(advancementFile.read())
  curs.execute('''CREATE TABLE monster_advancement (monster_id INTEGER NOT NULL, max_HD_this_size tinyint(2) NOT NULL,
FOREIGN KEY(monster_id) REFERENCES dnd_monster(id), UNIQUE(monster_id, max_HD_this_size));''')
  curs.execute('''CREATE TABLE monster_level_adjustment (monster_name TEXT, level_adjustment tinyint(1));''')
  curs.execute('''CREATE TABLE monster_has_skills (monster_name TEXT, skills TEXT, advancement TEXT);''')
  book = xlrd.open_workbook(os.path.join('..', 'Creature Catalog 3.5 noSLAsButLAandAdvancementAndSkills.xls'))
  #print(XLSfilepath, 'has', book.nsheets, 'sheets')
  #print('The sheet names are', book.sheet_names() )
  for sheet in book.sheets():
    names = set()
    for i,xls_row in itertools.dropwhile(lambda pair: pair[0] < 3,
                 enumerate(sheet.get_rows() ) ):
      name = xls_row[0].value
      if name in names:
        # don't deal with monsters appearing in multiple books until have firm foundation to match them up
        continue
      names.add(name)
      HD = xls_row[2].value
      size = xls_row[5].value
      skills = xls_row[55].value
      feats = xls_row[58].value
      advancement = xls_row[64].value
      LA = xls_row[65].value
      source = xls_row[66].value
      if not name:
        if HD or (skills and skills!='None') or advancement:
          raise Exception(name, HD, skills, advancement)
        continue
      if source[:4] == 'DRA ':
        continue
      if name == 'Frost Giant Variants':
        continue
      if type(HD) is str:
        if 'd' not in HD:
          raise Exception(name, HD)
        hit_dice = int(HD[:HD.find('d')])
      else:
        assert type(HD) is float
        if int(HD) != HD:
          assert HD == 0.5 or HD == 0.25 or HD == 0.125
        hit_dice = int(HD)
      size_id = id_from_name(curs, 'dnd_racesize', size.split()[0])
      assert size_id
      curs.execute('''INSERT INTO monster_has_skills (monster_name, skills, advancement) VALUES (?, ?, ?);''', (name, skills, advancement))
      if advancement == 'By character class':
        advancement = ''
      if advancement == 'By character class (Fighter)' or advancement == 'By character class (Rogue)':
        advancement = ''
      if advancement == 'By character class (Favored: Cleric)' or advancement == 'By character class (Favored: Sorcerer)' or advancement == 'By character class (Favored: Swordsage)':
        advancement = ''
      elif advancement == 'By class':
        advancement = ''
      elif advancement == 'see text' or advancement == 'See text':
        advancement = ''
      elif advancement == 'Special':
        advancement = ''
      elif advancement == 'None':
        advancement = None
      elif advancement == '* [Pseudonatural + Entropic]':
        advancement = None
      elif advancement == '2-3 HD (Medium); 4-9 HD (Large aquatic); 10-20 HD (Huge aquatic)':
        advancement = '2-3 HD (Medium); 4-9 HD (Large); 10-20 HD (Huge)'
      elif advancement[2:] == '+ HD':
        advancement = advancement + ' (Colossal)'
      elif name == 'Kython, Slaughterking':
        advancement = '19-30 HD (Large); 31-36 HD (Huge)'
      elif name == 'Drake, Elemental, Fire':
        advancement = '15-19 HD (Large); 20-29 HD (Huge); 30-39 HD (Gargantuan); 40-42 HD (Colossal)'
      elif name == 'Casurua':
        advancement = '21-28 HD (Huge); 29-38 HD (Gargantuan)'
      elif name == 'Golem, Cadaver':
        advancement = '11-20 HD (Large); 21-30 HD (Huge)'
      elif name == 'Archon, Warden':
        advancement = '9-18 HD (Large); 19-24 HD (Huge)'
      handle_one_monster_advancement_line(curs, name, source, hit_dice, size_id, advancement)
      try:
        LA = int(LA)
      except ValueError:
        continue
      curs.execute('''INSERT INTO monster_level_adjustment (monster_name, level_adjustment) VALUES (?, ?)''', (name, LA))
  handle_one_monster_advancement_line(curs, 'Dinosaur, Cryptoclidus', None, 3, 6, '4-6 HD (Large); 7-9 HD (Huge)')
  handle_one_monster_advancement_line(curs, 'Bat, Guard', None, 4, 6, '5-12 HD (Huge)')
  print('About to read Echohawk')
  before = time.time()
  echohawk = pandas.read_excel(os.path.join('..', "Echohawk's Complete D&D Monster Index (2008-12-29) LA and class levels.ods"), engine='odf')
  print('read Echohawk in', int(time.time() - before))
  for index, row in echohawk.iterrows():
    LA = row['LA']
    try:
      LA = int(LA)
    except ValueError:
      continue
    name = row['Monster']
    curs.execute('''INSERT INTO monster_level_adjustment (monster_name, level_adjustment) VALUES (?, ?)''', (name, LA))

def handle_one_monster_advancement_line(curs, name, source, hit_dice, size_id, advancement):
      advancementRanges = [r.strip() for r in advancement.split(';')] if advancement else []
      upperEndpoints = list()
      maxAdvancementHDforSize = int(hit_dice)
      prev_size_id = size_id
      for index,r in enumerate(advancementRanges):
        if r == '3136 HD (Gargantuan)':
          r = '31-36 HD (Gargantuan)'
        elif r == '40-48 HD Colossal)':
          r = '40-48 HD (Colossal)'
        elif r == '7-12 HD Large)':
          r = '7-12 HD (Large)'
        elif r == '11-20 HD (Hudge)':
         r = '11-20 HD (Huge)'
        matchObj = advancementRE.match(r)
        if not matchObj:
          raise Exception(name, r)
        advancementSize = matchObj.group(3)
        new_size_id = id_from_name(curs, 'dnd_racesize', advancementSize)
        assert new_size_id
        minAdvancementHDforSize = int(matchObj.group(1))
        if minAdvancementHDforSize != maxAdvancementHDforSize + 1:
          if name == 'Dusk Giant, Greater' or name == 'Devil, Abishai, Blue' or name == 'Devil, Abishai, Green':
            # This is in HoH itself, and FC2.
            assert index == 0
            if minAdvancementHDforSize != hit_dice:
              raise Exception(name, minAdvancementHDforSize, hit_dice)
            minAdvancementHDforSize = hit_dice + 1
          elif name == 'Naga, Bright' or name == 'Vodyanoi' or name == 'Devil, Abishai, White' or name == 'Spider, Snow, Large':
            assert index == 0
            assert minAdvancementHDforSize == hit_dice + 2
          elif name == 'Landwyrm, Underdark' or name == 'Ashworm':
            assert minAdvancementHDforSize == maxAdvancementHDforSize
          elif name == 'Bleakborn':
            assert index == 0
            assert minAdvancementHDforSize == hit_dice - 2
          elif name == 'Cinderspawn':
            assert index == 1
            assert minAdvancementHDforSize == maxAdvancementHDforSize - 4
          elif name == 'Firefly, Giant':
            assert index == 1
            assert minAdvancementHDforSize == maxAdvancementHDforSize - 1
          else:
            raise Exception(source, name, hit_dice, advancement, maxAdvancementHDforSize, minAdvancementHDforSize)
        if matchObj.group(2) == '+':
          assert advancementSize == 'Colossal'
          maxAdvancementHDforSize = 127
        elif not matchObj.group(2):
          maxAdvancementHDforSize = minAdvancementHDforSize
        else:
          assert matchObj.group(2)[0] == '-'
          maxAdvancementHDforSize = int(matchObj.group(2)[1:])
        if new_size_id != prev_size_id + 1:
          if index == 0:
            assert prev_size_id == size_id
            if name == 'Manta Ray':
              assert new_size_id == size_id - 1
              # According to the Monster Manual, manta rays really are smaller with fewer HD.
              # Rather than deal with this, we just pretend they stay Large.
              new_size_id = size_id
            else:
              assert new_size_id == size_id
          else:
            raise Exception(source, name, size, advancement)
        prev_size_id = new_size_id
        if index == 0 and new_size_id != size_id:
          if new_size_id != size_id + 1:
            raise Exception(source, name, size_id, advancement)
          upperEndpoints.append(hit_dice)
        upperEndpoints.append(maxAdvancementHDforSize)
      monster_id = id_from_name(curs, 'dnd_monster', name, useEdition=True)
      if not monster_id:
        return
      try:
        curs.executemany('INSERT INTO monster_advancement (monster_id, max_HD_this_size) VALUES (?, ?)', ((monster_id, e) for e in upperEndpoints))
      except sqlite3.IntegrityError as ex:
        curs.execute('SELECT max_HD_this_size FROM monster_advancement WHERE monster_id=?', (monster_id,))
        advancementAlreadyThere = curs.fetchall()
        if advancementAlreadyThere != [(e,) for e in upperEndpoints]:
          raise Exception(name, upperEndpoints, advancementAlreadyThere) from ex






ShaxItemRE = re.compile(r'<b>(?P<name>[\w\s]+)</b>[\s\w\d\.\,\(\)]*?<br />\nPrice: (?P<price>\d+) (?P<coin>[CSG])P<br />\nWeight: (?P<weight>\d+\.?\d*|\-+)#?<br />\n\((?P<book>[\w\s&;]+) p\. (?P<page>\d+)\)<br />\n(?P<desc>.+)')
# still need to handle:
assert not ShaxItemRE.match('''<b>Wick, Candle</b> (x5, 2 SP ea.)<br />
Price: 1 GP<br />s
Weight: --<br />
(Arms &amp; Equipment Guide p. 27)<br />
Can be used anywhere you would use Twine (Dungeonscape p. 33), such as tripwires, improvised alarm systems, fishing lines, signal kites, or hang from the ceiling to detect invisible flying creatures. In addition, you can use it as a timing device: it takes 30 seconds (5 rounds) to burn 1 inch. Comes in 50' rolls. <br />''')

assert(ShaxItemRE.match('''<b>Bolt Cutters</b><br />
Price: 6 GP<br />
Weight: 5#<br />
(Arms &amp; Equipment Guide p. 21)<br />
A useful tool, but can be replicated with Shapesand until you can get hold of Stone Dragon Belt/Mountain Hammer, which can cut through just about anything.<br />''').group('weight') == '5')
assert(ShaxItemRE.match('<b>Acidic Fire</b> (x2, 30 GP ea.)<br />\nPrice: 60 GP<br />\nWeight: 0.3#<br />\n(Expedition to Castle Ravenloft p. 208)<br />\nblah').group('name') == 'Acidic Fire')
assert(ShaxItemRE.match("<b>Kyo Crystals</b><br />\nPrice: 50 GP<br />\nWeight: --<br />\n(Expedition to Undermountain p. 217)<br />\nOne-shot items similar to potions, but unlike potions activating them with a standard action does not provoke an AoO. They cost the same as potions (I've listed a cost above for a 1st level spell because I thought they were the most useful), but the spell must be chosen from a limited list: <i>burning hands</i>, <i>cure light wounds</i>, <i>cure moderate wounds</i> (150 GP), <i>cure serious wounds</i> (250 GP), <i>light</i> (25 GP), <i>mage armor</i>, <i>magic missle</i>, <i>magic weapon</i>, <i>mirror image</i> (150 GP), or <i>ray of frost</i> (25 GP). Oddly enough, you don't have to have these spells on your spell list to create these items, so wizards can create <i>cure</i> crystals. Of the available spells, <i>magic missle</i> (autohit force damage) and <i>magic weapon</i> (easier than applying an oil) look like the most useful.").group('weight') == '--')
def readShax(filepath='ShaxItems.txt'):
  print('readShax')
  items = []
  rulebooks = set()
  contents = open(filepath, 'r').read()
  split = contents.split('<br />\n<br />\n')
  #print(split[0], ' = split[0]')
  #print(split[1], ' = split[1]')
  for entry in split:
    matchObj = ShaxItemRE.match(entry)
    if matchObj is not None:
      name = matchObj.group('name')
      price_sp = int(matchObj.group('price') )
      if matchObj.group('coin') == 'G':
        price_sp *= 10
      elif matchObj.group('coin') == 'C':
        price_sp /= 10
      try:
        weight = float(matchObj.group('weight') )
      except ValueError:
        weight = 0
      book = matchObj.group('book').replace('Xendrik', "Xen'drik").replace('&amp;', '&')
      rulebooks.add(book)
      page = int(matchObj.group('page') )
      #print('Shax entry', entry[:20], name, price_sp, weight, book, page)
      assert (name, price_sp, weight, matchObj.group('desc'), page, book) not in items
      items.append( (name, price_sp, weight, matchObj.group('desc'), page, book) )
  print('done readShax')
  return items

parenthesizedReferenceRE = re.compile(r'\((\w{1,5})\sp(\d{1,3})\)')
def readCrystalKeepItems(filepath='IndexMagicItems.docx'):
  # You can open any Word 2007 or later file this way (.doc files from Word 2003 and earlier won't work).
  fileObj = open(filepath, 'rb')
  wordDoc = docx.Document(filepath)
  personalItems = wordDoc.tables[0]
  for row in personalItems.rows:
    parenthesizedReferenceRE.match(row.cells[1].text)
  #for table in wordDoc.tables:
  #  for row in table.rows:
  #      for cell in row.cells:
  #          print(cell.text)
  # https://stackoverflow.com/questions/10366596/how-to-read-contents-of-an-table-in-ms-word-file-using-python
  # blog Reading Table Contents Using Python by etienne
  # The format of a docx file is described at Open Office XML.
  # with zipfile.ZipFile('<path to docx file>') as docx:
  #  tree = xml.etree.ElementTree.XML(docx.read('word/document.xml'))
  # http://etienned.github.io/posts/extract-text-from-word-docx-simply/
  # .docx's are basically zip files with several folders and files within them. 'word' is one of the folders and 'document.xml' is one of the files within that folder, which contains most of the document itself

def deitiesDocToCSV(filepath='indexDeities.docx'):
  import docx
  import pandas
  wordDoc = docx.Document(filepath)
  allDeities = wordDoc.tables[0]
  columnNames = [cell.text for cell in allDeities.rows[0].cells]
  df = pandas.DataFrame.from_records([[cell.text for cell in row.cells] for row in allDeities.rows[1:]], columns=columnNames)
  df.to_csv('indexDeities.csv')
  return df

def readDeities(filepath='indexDeities.docx'):
  '''
  file created by opening the original .doc and Save A Copy as .docx
  '''
  import docx
  import pandas
  if os.path.exists('indexDeities.csv'):
    df = pandas.read_csv('indexDeities.csv')
  else:
    df = deitiesDocToCSV(filepath)
  return df

"""
CREATE INDEX "dnd_item_slug" ON "dnd_item" ("slug");
CREATE INDEX "dnd_item_dnd_item_52094d6e" ON "dnd_item" ("name");
CREATE INDEX "dnd_item_dnd_item_51956a35" ON "dnd_item" ("rulebook_id");
CREATE INDEX "dnd_item_dnd_item_35a44c52" ON "dnd_item" ("body_slot_id");
CREATE INDEX "dnd_item_dnd_item_c181fb11" ON "dnd_item" ("aura_id");
CREATE INDEX "dnd_item_dnd_item_a7ff055e" ON "dnd_item" ("activation_id");
CREATE INDEX "dnd_item_dnd_item_6a812853" ON "dnd_item" ("property_id");
CREATE INDEX "dnd_item_dnd_item_ed720ca8" ON "dnd_item" ("synergy_prerequisite_id");
"""
def make_item_tables(curs):
  # http://www.imarvintpa.com/dndlive/items.php?ID=2128
  curs.execute('''DROP TABLE dnd_itemslot;''')
  curs.execute('''CREATE TABLE dnd_itemslot (
  id INTEGER PRIMARY KEY NOT NULL
  ,name CHAR(9) NOT NULL
  ,abbrev CHAR(2) DEFAULT NULL
  ,examples VARCHAR(128) DEFAULT NULL
  ,affinities VARCHAR(128) DEFAULT NULL
,UNIQUE(name)
  );''')
  curs.execute('''INSERT INTO dnd_itemslot (name, abbrev, examples, affinities) VALUES
    ("held", "-h", "weapons, shields, tools", NULL),
    ("Shield", "-s", "augment crystals", NULL), ("Armor", "-a", "augment crystals", NULL), ("ArmorOrShield", "-as", "augment crystals", NULL), ("Weapom", "-w", "augment crystals", NULL),
    ("Arms", "A", "armbands, bracelets, bracers", "Combat, Allies"),
    ("Body", "B", "armor, robes", "Multiple effects"),
    ("Face", "Fa", "goggles, lenses, masks, spectacles, third eyes", "Vision"),
    ("Feet", "Ft", "boots, sandals, shoes, slippers", "Movement"),
    ("Hands", "Ha", "gauntlets, gloves", "Quickness, Destructive power"),
    ("Head", "Hd", "circlets, crowns, hats, headbands, helmets, phylacteries", "Mental improvement, ranged attacks, Interaction, Morale, alignment"),
    ("Rings", "R", "rings", NULL),
    ("Shoulders", "S", "capes, cloaks, mantles, shawls", "Transformation, protection"),
    ("Throat", "Th", "amulets, badges, brooches, collars, medals, medallions, necklaces, pendants, periapts, scarabs, scarfs, torcs", "Protection, discernment"),
    ("Torso", "To", "shirts, tunics, vests, vestments", "Physical improvement, Class ability improvement"),
    ("Waist", "W", "belts, girdles, sashes", "Physical improvement");''')
  curs.execute('''CREATE INDEX index_dnd_itemslot_name ON dnd_itemslot(name);''')
  curs.execute('''CREATE TABLE armorweapon_property_type (
  id INTEGER PRIMARY KEY NOT NULL
  ,name CHAR(16) NOT NULL
  ,slug CHAR(16) NOT NULL
,UNIQUE(name)
  );''')
  curs.execute('''INSERT INTO armorweapon_property_type (name, slug) SELECT name, slug from dnd_itemproperty;''')
  curs.execute('''CREATE INDEX index_armorweapon_property_type_name ON armorweapon_property_type(name);''')

  curs.execute('''CREATE TABLE wealth_by_level (
  level TINYINT(2) PRIMARY KEY
  ,PCwealth INT(6)
  ,NPCwealth INT(6) NOT NULL
  );''')
  curs.execute('''INSERT INTO wealth_by_level (level, PCwealth, NPCwealth) VALUES (1, 10, 900), (2, 900, 2000), (3, 2700, 2500), (4, 5400, 3300), (5, 9000, 4300), (6, 13000, 5600), (7, 19000, 7200), (8, 27000, 9400), (9, 36000, 12000), (10, 49000, 16000), (11, 66000, 21000), (12, 88000, 27000), (13, 110000, 35000), (14, 150000, 45000), (15, 200000, 59000), (16, 260000, 77000), (17, 340000, 100000), (18, 440000, 130000), (19, 580000, 170000), (20, 760000, 220000);''')
  with open('treasure.sql') as treasureFile:
    curs.executescript(treasureFile.read())
  curs.execute('''CREATE TABLE item_level (
  level TINYINT(2)
  ,market_price_max_gp int(6)
  );''');
  curs.execute('''INSERT INTO item_level VALUES (-2, 50), (1, 150), (2, 400), (3, 800), (4, 1300), (5, 1800), (6, 2300), (7, 3000), (8, 4000), (9, 5000), (10, 6500), (11, 8000), (12, 10000), (13, 13000), (14, 18000), (15, 25000), (16, 35000), (17, 48000), (18, 64000), (19, 80000), (20, 100000), (21, 120000), (22, 140000), (23, 160000), (24, 180000), (25, 200000), (26, 220000), (27, 240000), (28, 260000), (29, 280000), (30, 300000);''')
  # http://www.d20srd.org/srd/spells/detectMagic.htm
  curs.execute('''DROP TABLE dnd_itemauratype;''')
  curs.execute('''CREATE TABLE dnd_itemauratype (
  name CHAR(12) NOT NULL,
  min_caster_level TINYINT(2) DEFAULT NULL,
  max_spell_level TINYINT(1) DEFAULT NULL,
  linger_duration CHAR(14)
  );''')
  curs.execute('''INSERT INTO dnd_itemauratype (name, min_caster_level, max_spell_level, linger_duration) VALUES ("Faint", NULL, 3, "1d6 rounds"), ("Moderate", 6, 6, "1d6 minutes"), ("Strong", 12, 9, "1d6x10 minutes"), ("Overwhelming", 21, NULL, "1d6 days");''')
  
  # It was Shakespeare, in his play Henry IV (1597), who first used "avoirdupois" to mean "heaviness."
  # pound is equal to 16 ounces, the ounce 16 drams, and the dram 27.344 grains.
  # The avoirdupois pound contains 7,000 grains
  # Troy weight, traditional system of weight in the British Isles based on the grain, pennyweight (24 grains), ounce (20 pennyweights), and pound (12 ounces).
  # The troy pound is 5760 grains (12 oz t), while an avoirdupois pound is approximately 21.53% heavier at 7000 grains
  # The standard weight measurements used for all precious metal products are Troy ounces and pounds. All legal tender silver, gold and platinum coins are struck according to these Troy weights, which are not to be confused with the Avoirdupois weights.
  # 1 troy ounce = 480 grains; 1 ounce = 437.5 grains. Basic elementary school math will tell you that 1 Troy ounce contains 2.75 grams more of metal than the standard ounce.
  # Apothecaries' weight, traditional system of weight in the British Isles used for the measuring and dispensing of pharmaceutical items and based on the grain, scruple (20 grains), dram (3 scruples), ounce (8 drams), and pound (12 ounces).
  # The standard coin weighs about a third of an ounce (fifty to the pound). So D&D uses about 16ounces to the pound.
  # The avoirdupois pound contains 7,000 grains. It is equal to about 1.22 apothecaries' or troy pounds (5760grains).

  curs.execute('''CREATE TEMPORARY TABLE item_backup (id int(11), name varchar(64), slug varchar(64), rulebook_id int(11), page smallint(5), price_gp int(10), price_bonus smallint(5), item_level smallint(5), body_slot_id int(11), caster_level smallint(5), aura_id int(11), aura_dc varchar(16), activation_id int(11), weight double, visual_description longtext, description longtext, description_html longtext, type varchar(3), property_id int(11), cost_to_create varchar(128), synergy_prerequisite_id int(11), required_extra varchar(64) );''')
  curs.execute('''INSERT INTO item_backup SELECT id, name, slug, rulebook_id, page, price_gp, price_bonus, item_level, body_slot_id, caster_level, aura_id, aura_dc, activation_id, weight, visual_description, description, description_html, type, property_id, cost_to_create, synergy_prerequisite_id, required_extra FROM dnd_item;''')
  curs.execute('''DROP TABLE dnd_item;''')
  curs.execute('''CREATE TABLE dnd_item (
  id INTEGER PRIMARY KEY NOT NULL
  ,name varchar(64) NOT NULL
  ,slug varchar(64) DEFAULT NULL
  ,rulebook_id INTEGER
  ,page unsigned smallint(3) DEFAULT NULL
  ,market_price_copper_pieces unsigned int(8) DEFAULT NULL
  ,body_slot_id INTEGER DEFAULT NULL
  ,caster_level UNSIGNED TINYINT(2) DEFAULT NULL
  ,activation_id, INTEGER DEFAULT NULL
  ,weight double DEFAULT NULL
  ,visual_description longtext DEFAULT NULL
  ,description longtext NOT NULL
  ,description_html longtext DEFAULT NULL
,FOREIGN KEY(rulebook_id) REFERENCES dnd_rulebook(id)
,FOREIGN KEY(body_slot_id) REFERENCES dnd_itemslot(id)
,FOREIGN KEY(activation_id) REFERENCES dnd_itemactivationtype(id)
,UNIQUE(name, rulebook_id, market_price_copper_pieces)
  );''')
  # The Spellcraft DC to identify the aura can always be rederived: 15 + caster_level/2 as DC.
  # The aura strength can always be rederived from the caster level dnd_itemauratype table.
  # Armor and weapon properties don't really belong with the main table especially when you're looking at
  # the low end of price, since they require at least a +1 to start.
  # A least augment crystal functions whenever attached to an object of at least masterwork quality, even if the object itself has no magical properties.
  curs.execute('''CREATE TABLE dnd_armorweapon_property (
  id INTEGER PRIMARY KEY NOT NULL
  ,name varchar(64) NOT NULL
  ,slug varchar(64) DEFAULT NULL
  ,rulebook_id INTEGER
  ,page unsigned smallint(3) DEFAULT NULL
  ,candidates_id INTEGER NOT NULL
  ,price_in_gp unsigned int(6) DEFAULT NULL
  ,price_as_armorweapon_bonus unsigned tinyint(1) DEFAULT NULL
  ,synergy_prereq INTEGER DEFAULT NULL
  ,caster_level TINYINT(2) DEFAULT NULL
  ,activation_id, INTEGER DEFAULT NULL
  ,visual_description longtext DEFAULT NULL
  ,description longtext NOT NULL
  ,description_html longtext DEFAULT NULL
,UNIQUE(rulebook_id, name)
,FOREIGN KEY(rulebook_id) REFERENCES dnd_rulebook(id)
,FOREIGN KEY(candidates_id) REFERENCES armorweapon_property_type(id)
,FOREIGN KEY(synergy_prereq) REFERENCES dnd_armorweapon_property(id)
,FOREIGN KEY(activation_id) REFERENCES dnd_itemactivationtype(id)
  );''')
  #readCrystalKeepItems()
  deities = readDeities()
  curs.execute('''CREATE TEMPORARY TABLE deities_backup (
  "id" int(11) NOT NULL ,
  "name" varchar(64) NOT NULL,
  "slug" varchar(64) NOT NULL,
  "description" longtext NOT NULL,
  "description_html" longtext NOT NULL,
  "alignment" varchar(2) NOT NULL,
  "favored_weapon_id" int(11) DEFAULT NULL,
  PRIMARY KEY ("id")
  CONSTRAINT "favored_weapon_id_refs_id_b09a3eba" FOREIGN KEY ("favored_weapon_id") REFERENCES "dnd_item" ("id")
  );''')
  curs.execute('''INSERT INTO deities_backup SELECT * FROM dnd_deity;''')
  curs.execute('''DROP TABLE dnd_deity;''')
  curs.execute('''CREATE TABLE dnd_deity (
  id INTEGER PRIMARY KEY NOT NULL
  ,name varchar(64) NOT NULL
  ,slug varchar(64) DEFAULT NULL
  ,description longtext DEFAULT NULL
  ,description_html longtext DEFAULT NULL
  ,portfolio longtext DEFAULT NULL
  ,law_chaos char(1) NOT NULL
  ,good_evil tinyint(1) NOT NULL
  ,favored_weapon_id INTEGER DEFAULT NULL
  ,rulebook_id INTEGER DEFAULT NULL
  ,page unsigned smallint(3) DEFAULT NULL
  ,rank char(1) DEFAULT NULL
,UNIQUE(rulebook_id, name)
,FOREIGN KEY(rulebook_id) REFERENCES dnd_rulebook(id)
,FOREIGN KEY(favored_weapon_id) REFERENCES dnd_item(id)
  );''')
  curs.execute('''INSERT INTO dnd_deity (name, slug, description, description_html, good_evil, law_chaos, favored_weapon_id, rulebook_id, page, rank) SELECT name, slug, description, description_html, CASE substr(alignment, length(alignment), 1) WHEN 'G' THEN 1 WHEN 'N' THEN 0 WHEN 'E' THEN -1 ELSE 0 END, substr(alignment, 1, 1), favored_weapon_id, null, 0, 'D' FROM deities_backup;''')
  curs.execute('''CREATE TEMPORARY TABLE domains_backup (
  "id" int(11) NOT NULL ,
  "name" varchar(64) NOT NULL,
  "slug" varchar(64) NOT NULL,
  PRIMARY KEY ("id")
  );''')
  curs.execute('''INSERT INTO domains_backup SELECT * from dnd_domain;''')
  curs.execute('''DROP TABLE dnd_domain;''')
  curs.execute('''CREATE TABLE dnd_domain (
  id INTEGER PRIMARY KEY NOT NULL
  ,name varchar(64) NOT NULL
  ,slug varchar(64) NOT NULL
  );''')
  curs.execute('''INSERT INTO dnd_domain (id, name, slug) SELECT id, name, slug FROM domains_backup;''')
  curs.execute('''CREATE TABLE deity_has_domain (
  deity_id INTEGER
  ,domain_id INTEGER
,FOREIGN KEY(deity_id) REFERENCES dnd_deity(id)
,FOREIGN KEY(domain_id) REFERENCES dnd_domain(id)
  );''')
  deities.dropna(inplace=True)
  for deityTuple in deities.itertuples():
    deityAlignment = deityTuple.Align
    portfolioMaybePrecededByEpithets = deityTuple[7].split('\n')
    #print(deityTuple[0], deityTuple[1], deityTuple[2], deityTuple[3], deityTuple[4], deityTuple[5], deityTuple[6], deityTuple[7], deityTuple[8])
    if len(portfolioMaybePrecededByEpithets) == 1:
      portfolio = portfolioMaybePrecededByEpithets[0]
    elif len(portfolioMaybePrecededByEpithets) == 2:
      epithets, portfolio = portfolioMaybePrecededByEpithets
    else:
      indexOfLastCloseQuote = max(index for index, line in enumerate(portfolioMaybePrecededByEpithets) if '”' in line)
      epithets = ''.join(portfolioMaybePrecededByEpithets[:indexOfLastCloseQuote+1])
      portfolio = ''.join(portfolioMaybePrecededByEpithets[indexOfLastCloseQuote+1:])
    #print(type(portfolio), portfolio)
    curs.execute('''INSERT INTO dnd_deity (name, description, good_evil, law_chaos, portfolio) VALUES (?, ?, ?, ?, ?)''', (deityTuple.Name.strip(), deityTuple[7] + deityTuple.Symbol, 1 if 'G' in deityAlignment else -1 if 'E' in deityAlignment else 0, deityAlignment[0], portfolio) )
    deity_id = curs.lastrowid
    domains = deityTuple.Domains.replace('Strength Wrath', 'Strength, Wrath').replace('Chaos Luck', 'Chaos, Luck').split(', ')
    for domain in domains:
      if domain == 'Oillusion':
        domain = 'Illusion'
      elif domain == 'Inquinistion':
        domain = 'Inquisition'
      elif domain == 'Portal’':
        domain = 'Portal'
      domain_id = insert_if_needed(curs, 'dnd_domain', domain, slug=domain.lower())
      curs.execute('''INSERT INTO deity_has_domain (deity_id, domain_id) VALUES (?, ?);''', (deity_id, domain_id) )
  items = readShax()
  if True: # the below SQL command will silently drop any non-matched rulebook name
    for item in items:
      curs.execute('''SELECT dnd_rulebook.id FROM dnd_rulebook WHERE dnd_rulebook.name=?;''', item[-1:] )
      #print('trying to fetch dnd_rulebook for', item)
      if curs.fetchone() is None:
        raise Exception(item)
  items = set(items)
  #print([pair[0] for pair in collections.Counter([item[0] for item in items]).items() if pair[1] != 1])
  curs.executemany('''INSERT INTO dnd_item (name, market_price_copper_pieces, weight, description, page, rulebook_id) SELECT ?, ?*10, ?, ?, ?, dnd_rulebook.id FROM dnd_rulebook WHERE dnd_rulebook.name=?;''', items)
"""
I changed the dnd_itemslot table because separating weapons from gauntlets doesn't work for spiked gauntlets.
A normal humanoid creature has twelve body slots, enumerated here with some examples of the kinds of items that might be worn there (for nonhumanoid creatures, see Size and Shape, below).
As a default rule, treat creatures of any shape as having all the normal body slots available. Creatures never gain extra body slots for having extra body parts (for example, a marilith still has only one hands body slot and two rings body slots).
Amorphous Creatures: Creatures without any shape, such as most oozes and the phasm (in its normal form), have no body slots and can't wear magic items at all.
Headless Creatures: Creatures without an identifiable head, such as shambling mounds, lack the face, head, and throat body slots.
Armless Creatures: Creatures without forelimbs, such as snakes, don't have the arms, hands, or rings body slot (but see multilegged creatures, below). A creature with only a single forelimb retains these body slots, and can wear both of a pair on the same limb (such as both gloves on the same hand, and so on).
Multilegged Creatures: Creatures with more than two legs can treat their foremost pair of limbs as their arms (allowing them access to the arms, hands, and rings body slots), even if those limbs are used for locomotion rather than for manipulation.
Fingerless Creatures: Creatures without flexible digits or extremities, such as horses, lack the rings body slot. A creature need not be able to manipulate objects to wear rings: a hell hound can wear a ring on a toe of its forelimb.
Legless Creatures: Creatures without hind limbs, such as lillends, don't have the feet body slot.
Arms: armbands, bracelets, bracers. Combat, Allies
Body: armor, robes. Multiple effects
 One robe or suit of armor on the body (over a vest, vestment, or shirt)
 One belt around the waist (over a robe or suit of armor)
A = Arms; B = Body; Fa = Face; Ft = Feet; Ha = Hands; Hd = Head; R = Ring; S = Shoulders; Th = Throat; To = Torso; W = Waist
the following entries on the Body Slot line.
-: The item functions or can be activated as long as it is carried somewhere on your body (but not if it's stored in an extradimensional or nondimensional storage space, such as a bag of holding). Some rare items in this category might describe a particular manner in which you must carry them for the item to function (such as ioun stones).
- (held): You must hold the item or otherwise manipulate the item with your hand for it to function or be activated. All weapons and shields have this entry, as do many tools. In the case of a shield, simply carrying it isn't enough. you must wear it properly as described on page 125 of the Player's Handbook.
 Technically, however, you can gain the magical benefits of as many shields as you can wear.
- ([armor, shield, or weapon] crystal): Augment crystals are magic items that function only when attached to a suit of armor, shield, weapon, or other appropriate item. Like properties, you can only gain an augment crystal's benefit while you're wearing or holding the item in the appropriate manner.

Properties are part of another item (a weapon, shield, or suit of armor), and they function or can be activated as long as the item is worn or held properly. A shield property offers no benefit if the shield is slung over your shoulder, and a weapon (typically) doesn't offer any benefit if it's sheathed.
Instead of a body slot entry, a property has a property entry, which describes the types of items to which this property can be applied.
Unless noted otherwise in the property's Property entry, each special property in this chapter can be added either to a suit of armor or shield.
Unlike most other items, properties have no weight.
To add a special property to a shield or suit of armor, the shield or armor must already have at least a +1 enhancement bonus.

An item's market price determines its level. Find the market price range on Table 6-3: Item Levels by Price in which this value falls; this tells you its level. For example, a cloak of elvenkind has a market price of 2,500 gp. This falls between 2,301 gp and 3,000 gp, which makes the cloak a 7th-level magic item.
Fast-Play Exception: For magic weapons and armor, it's easiest to ignore the portion of market price derived from the masterwork item itself, as long as that's just a small fraction of the overall price. For example, treat a +1 greatsword (2,350 gp) as a 6th-level item, even though its actual market price is a little bit above that range. Don't abuse this shortcut by claiming that +1 full plate (2,650 gp market price) is only a 4th-level item (801–1,300 gp).
When you need to equip a bunch of NPCs in a hurry, or you just want a playable PC for tonight's game, you might reasonably choose for speed to take precedence over precision. replaces the precision of market price with the abstraction of level. A cloak of resistance +1, pipes of the sewers, and a divine scroll of slay living are all 4th-level magic items, even though their market prices are slightly different.
Be warned that this system consciously trades precision for speed. It allows you to equip a character quickly, but it doesn't necessarily spend every last gold piece available, nor does it exactly replicate what you could purchase with the normal systems available. When creating an important NPC, building a player character's equipment list for a long-term campaign, or designing a treasure hoard (see page 265), consider using the normal rules in place of these.

When detect magic identifies a magic item's school of magic, this information refers to the school of the spell placed within the potion, scroll, or wand, or the prerequisite given for the item. If more than one spell is given as a prerequisite, use the highest-level spell. If no spells are included in the prerequisites, use the following default guidelines.
Armor and protection items	Abjuration
Weapons or offensive items	Evocation
Bonus to ability score, on skill check, etc.	Transmutation
"""

def create_rulebook_table(curs):
  rulebook_max_name_len = max(max(len(n) for n in rulebook_abbreviations.values() ), 128)
  rulebook_max_abbr_len = max(max(len(n) for n in rulebook_abbreviations.keys() ), 7)
  # backup mirrors the schema of the original
  curs.execute('''CREATE TEMPORARY TABLE rulebooks_backup (id int, dnd_edition_id int, name varchar({}), abbr varchar({}), description longtext, year varchar(4), official_url varchar(255), slug varchar(128), image varchar(128), published date);'''.format(rulebook_max_name_len, rulebook_max_abbr_len) )
  curs.execute('''INSERT INTO rulebooks_backup SELECT id, dnd_edition_id, name, abbr, description, year, official_url, slug, image, published FROM dnd_rulebook;''')
  curs.execute('''DROP TABLE dnd_rulebook;''')
  curs.execute('''CREATE TABLE dnd_rulebook (
  id INTEGER PRIMARY KEY NOT NULL,
  dnd_edition_id INTEGER DEFAULT NULL,
  name varchar({}) NOT NULL,
  description longtext DEFAULT NULL,
  official_url varchar(255) DEFAULT NULL,
  slug varchar(128) DEFAULT NULL,
  image varchar(100) DEFAULT NULL,
  published date DEFAULT NULL,
  year smallint(2) DEFAULT NULL,
  FOREIGN KEY(dnd_edition_id) REFERENCES dnd_dndedition(id),
  UNIQUE(name)
  );'''.format(rulebook_max_name_len, rulebook_max_abbr_len) )
  curs.execute('''INSERT INTO dnd_rulebook (dnd_edition_id, name, description, year, official_url, slug, image, published) SELECT dnd_edition_id, name, description, year, official_url, slug, image, published FROM rulebooks_backup;''')
  # Save the abbreviations from the old rulebook table before dropping it.
  curs.execute('''CREATE TABLE rulebook_abbrev (
  abbr CHAR({}),
  rulebook_id INTEGER,
UNIQUE(abbr),
FOREIGN KEY(rulebook_id) REFERENCES dnd_rulebook(id)
  );'''.format(rulebook_max_abbr_len) )
  curs.execute('''INSERT INTO rulebook_abbrev (abbr, rulebook_id) SELECT abbr, dnd_rulebook.id FROM rulebooks_backup INNER JOIN dnd_rulebook ON dnd_rulebook.name=rulebooks_backup.name WHERE abbr!="EE" and abbr!="DD" and dnd_rulebook.name!="Monster Manual";''')
  migrate_rulebook_id(curs)
  curs.execute('''DROP TABLE rulebooks_backup;''')
  # Many of the rulebooks in rulebook_abbreviations are not listed yet, so list them first.
  curs.executemany('''INSERT OR IGNORE INTO dnd_rulebook (name) VALUES (?);''', [(name,) for name in rulebook_abbreviations.values()] )
  # Now that the rulebooks are listed, we set up their abbreviations.
  curs.executemany('''INSERT OR REPLACE INTO rulebook_abbrev (abbr, rulebook_id) SELECT ?, dnd_rulebook.id FROM dnd_rulebook WHERE name=?;''', rulebook_abbreviations.items() )
  curs.execute('''CREATE INDEX index_dnd_rulebook_name ON dnd_rulebook(name);''')
  curs.execute('''CREATE INDEX index_rulebook_abbrev ON rulebook_abbrev(abbr);''')

def migrate_rulebook_ids_in_table(curs, tableName):
  SQLcmd = '''UPDATE {0} SET rulebook_id = (SELECT dnd_rulebook.id FROM dnd_rulebook INNER JOIN rulebooks_backup ON dnd_rulebook.name=rulebooks_backup.name WHERE rulebooks_backup.id=rulebook_id);'''.format(tableName)
  curs.execute(SQLcmd)

def migrate_rulebook_id(curs):
  migrate_rulebook_ids_in_table(curs, 'dnd_feat')
  curs.execute('''UPDATE dnd_characterclassvariant SET rulebook_id = (SELECT dnd_rulebook.id FROM dnd_rulebook INNER JOIN rulebooks_backup ON dnd_rulebook.name=rulebooks_backup.name WHERE rulebooks_backup.id=rulebook_id);''')
  curs.execute('''UPDATE dnd_spell SET rulebook_id = (SELECT dnd_rulebook.id FROM dnd_rulebook INNER JOIN rulebooks_backup ON dnd_rulebook.name=rulebooks_backup.name WHERE rulebooks_backup.id=rulebook_id);''')

def get_rulebook_id(curs, rulebook_abbrev):
  if re.match(r'#?\d\d\d', rulebook_abbrev): rulebook_abbrev = r'\d\d\d'
  elif re.match(r'A\d\d', rulebook_abbrev): rulebook_abbrev = r'A\d\d'
  curs.execute('''SELECT rulebook_id FROM rulebook_abbrev WHERE abbr=?;''',
                   (rulebook_abbrev,) )
  result = curs.fetchone()
  if result is None:
    raise Exception('no rulebook found:', rulebook_abbrev)
  return result[0]







# http://stackoverflow.com/questions/4055564/what-does-the-number-in-parenthesis-really-mean
# The number after INT is in base 10, not base 2. https://blogs.oracle.com/jsmyth/what-does-the-11-mean-in-int11
# https://dba.stackexchange.com/questions/369/int5-vs-smallint5/370
# An int can be between -2147483648 and 2147483647 signed, or 0 and 4294967295 unsigned.
# A smallint is between -32768 and 32767 signed, or 0 and 65535 unsigned.
# TINYINT -128 to 127
# The (5) represents the display width of the field.
# The display width does not constrain the range of values that can be stored in the column.
# You get no performance improvement if you use CHAR against VARCHAR in one field, but the table contains other fields that are VARCHAR.
# Slug is used to make a name that is not acceptable for various reasons - e.g. containing special characters, too long, mixed-case, etc. - appropriate for the target usage. What target usage means is context dependent
# FOREIGN KEY https://www.sqlite.org/foreignkeys.html

# In SQLite, a column with type INTEGER PRIMARY KEY is an alias for the ROWID https://www.sqlite.org/autoinc.html
# On an INSERT, if the ROWID or INTEGER PRIMARY KEY column is not explicitly given a value, then it will be filled automatically with an unused integer
# interestingly, int(11) PRIMARY KEY does not work for auto-increment, so for example you cannot insert into dnd_spellschool
# http://www.sqlabs.com/blog/2010/12/sqlite-and-unique-rowid-something-you-really-need-to-know/
# From the official documentation: "Rowids can change at any time and without notice. If you need to depend on your rowid, make an INTEGER PRIMARY KEY, then it is guaranteed not to change. The VACUUM command may change the ROWIDs of entries in any tables that do not have an explicit INTEGER PRIMARY KEY."

def read_xls(XLSfilepath="Monster Compendium.xls"):
  book = xlrd.open_workbook(XLSfilepath)
  #print(XLSfilepath, 'has', book.nsheets, 'sheets')
  #print('The sheet names are', book.sheet_names() )
  alphabetical = book.sheet_by_index(0)
  templates = book.sheet_by_index(1)
  ODE = book.sheet_by_index(3)
  #print(ODE.row_values(0) )
  #print(alphabetical.row_values(0) )
  assert ODE.row_values(0)[:33] == alphabetical.row_values(0)[:33]
  # BUT SLAs are given in the SLA column in ODE and not in alphabetical
  return alphabetical,ODE

"""
sqlite> .schema dnd_racetype
CREATE TABLE "dnd_racetype" (
  "id" int(11) NOT NULL ,
  "name" varchar(32) NOT NULL,
  "slug" varchar(32) NOT NULL,
  "hit_die_size" smallint(5)  NOT NULL,
  "base_attack_type" varchar(3) NOT NULL,
  "base_fort_save_type" varchar(4) NOT NULL,
  "base_reflex_save_type" varchar(4) NOT NULL,
  "base_will_save_type" varchar(4) NOT NULL,
  PRIMARY KEY ("id")
);
sqlite> select * from dnd_racetype;
1|Aberration|aberration|8|CLR|BAD|BAD|GOOD
2|Animal|animal|8|CLR|GOOD|GOOD|BAD
3|Animal (good will)|animal-good-will|8|CLR|GOOD|GOOD|GOOD
4|Contruct|contruct|10|CLR|BAD|BAD|BAD
5|Dragon|dragon|12|FIG|GOOD|GOOD|GOOD
6|Elemental (Air)|elemental-air|8|CLR|BAD|GOOD|BAD
7|Elemental (Fire)|elemental-fire|8|CLR|GOOD|BAD|GOOD
8|Elemental (Water)|elemental-water|8|CLR|GOOD|BAD|BAD
9|Elemental (Earth)|elemental-earth|8|CLR|GOOD|BAD|BAD
10|Fey|fey|6|WIZ|BAD|GOOD|GOOD
11|Giant|giant|8|CLR|GOOD|BAD|BAD
12|Humanoid (good fort)|humanoid-good-fort|8|CLR|GOOD|BAD|BAD
13|Humanoid (good reflex)|humanoid-good-reflex|8|CLR|BAD|GOOD|BAD
14|Humanoid (good will)|humanoid-good-will|8|CLR|BAD|BAD|GOOD
15|Magical Beast|magical-beast|10|FIG|GOOD|GOOD|BAD
16|Monstrous humanoid|monstrous-humanoid|8|FIG|BAD|GOOD|GOOD
17|Ooze|ooze|10|CLR|BAD|BAD|BAD
18|Outsider|outsider|8|FIG|GOOD|GOOD|GOOD
19|Plant|plant|8|CLR|GOOD|BAD|BAD
20|Undead|undead|12|WIZ|BAD|BAD|GOOD
21|Vermin|vermin|8|CLR|GOOD|BAD|BAD
"""
fortitude_divisors = {'Animal':2, # certain animals have different good saves
      'Construct':3, 'Dragon':2,
      # Good saves depend on the element: Fortitude (earth, water) or Reflex (air, fire).
      'Fey':3, 'Giant':2,
      # Good Reflex saves (usually; a humanoid's good save varies). # Especially if they have class levels.
      'Ooze':3, 'Outsider':2, 'Plant':2, 'Undead':3, 'Vermin':2,
      }
fortitude_additions = {key:(2 if d==2 else 0) for key,d in fortitude_divisors.items()}

def recreate_db_file(DBpath='dnd.sqlite', monsterOnlyDB = 'dnd_monsters.sqlite'):
  if os.path.exists(monsterOnlyDB):
    os.remove(monsterOnlyDB)
  assert not os.path.exists(monsterOnlyDB)
  shutil.copyfile(DBpath, monsterOnlyDB)
  print('creating file', monsterOnlyDB)
  return sqlite3.connect(monsterOnlyDB)

def create_database(XLSfilepath="Monster Compendium.xls", DBpath='dnd.sqlite',
                    marvinCache=None):
  '''The original dnd_monster has only 29 monsters and has such design flaws as the default for attack being greatsword, so start from scratch.
  '''
  monsterOnlyDB = 'dnd_monsters.sqlite'
  conn = recreate_db_file(DBpath, monsterOnlyDB)
  curs = conn.cursor()

  if marvinCache is not None:
    if os.path.exists(marvinCache):
      #os.rmdir(marvinCache)
      print('deleting', marvinCache)
      shutil.rmtree(marvinCache)
  cache_IMarvinTPA(marvinCache)

  #def seven(x): return 7
  #conn.create_function("seven", 1, seven) # not accessible from sqlite3

  print('about to read_xls')
  alphabetical,ODE = read_xls(XLSfilepath)
  print('done reading xls')
  #ipdb.set_trace()
  #print('maxlen among names =', max([str(row[0]) for row in alphabetical.get_rows()], key=len) )

  # see existing tables dnd_racespeed and dnd_racespeedtype:
  # table with a single char for Burrow/Climb/Land/Fly/Swim, each monster_id might have a couple of entries
  # http://www.d20srd.org/srd/specialAbilities.htm#movementModes
  curs.execute('''CREATE TABLE dnd_movement_mode (
  id INTEGER PRIMARY KEY NOT NULL,
  name char(6) NOT NULL,
  abbrev char(1) NOT NULL
  );''')
  curs.execute('''INSERT INTO dnd_movement_mode (name, abbrev) VALUES ("Burrow", 'b'), ("Climb", 'c'), ("Fly", 'f'), ("Swim", 's');''')
  curs.execute('''CREATE TABLE monster_movement_mode (
  monster_id INTEGER NOT NULL,
  abbrev char(1) NOT NULL,
  speed smallint(2) NOT NULL,
FOREIGN KEY(monster_id) REFERENCES dnd_monster(id),
FOREIGN KEY(abbrev) REFERENCES dnd_movement_mode(abbrev)
  );''')
  curs.execute('''CREATE TABLE dnd_maneuverability (
  maneuverability INTEGER PRIMARY KEY NOT NULL,
  name char(7) NOT NULL,
  max_up_degrees tinyint(1) NOT NULL,
  max_up_per_move FLOAT NOT NULL
  );''')
  max_up_degrees = [45, 45, 60, 90, 90]
  max_up_per_move = [0.25, 0.25, 1/(1/math.tan(math.pi*60/180) + 1)/2, 0.5, 1]
  #print(tuple(itertools.chain(*zip(max_up_degrees, max_up_per_move) ) ))
  curs.execute('''INSERT INTO dnd_maneuverability (maneuverability,name,max_up_degrees,max_up_per_move) VALUES (1,"Clumsy",?,?), (2,"Poor",?,?), (3,"Average",?,?), (4,"Good",?,?), (5,"Perfect",?,?);''', tuple(itertools.chain(*zip(max_up_degrees, max_up_per_move) ) ) )
  # max up angle: 60, 45, 45
  # So if maneuverability is less than good, then in addition to upward movement counting double, they must spend an equal amount of movement (for clumsy or poor) or 1/\sqrt{3} as much movement (if average) on horizontal movement. If x = (1 - x)\sqrt{3}, then x = \sqrt{3} / (1 + \sqrt{3}) = 3/(3 + \sqrt3). That multiplies their net upward speed by 1/2 and 3/(3 + \sqrt3) = 5/8 = 10387/2**14 respectively.
  # continued fraction? 1/(1 + \sqrt3): \sqrt3 = 1 + 2/(2 + ?)
  # If x = (1 - x)*math.tan(angle), then (math.tan(angle) + 1)*x = math.tan(angle), so x = math.tan(angle)/(math.tan(angle) + 1).
  curs.execute('''CREATE TABLE monster_maneuverability (
  monster_id INTEGER NOT NULL,
  maneuverability tinyint(1) NOT NULL,
FOREIGN KEY(monster_id) REFERENCES dnd_monster(id),
FOREIGN KEY(maneuverability) REFERENCES dnd_maneuverability(maneuverability)
  );''')

  curs.execute('''CREATE TABLE dnd_weapon (
  id INTEGER PRIMARY KEY NOT NULL
  ,name char(8) NOT NULL
  );''')
  curs.execute('''INSERT INTO dnd_weapon (name) VALUES ("bite"), ("claw"), ("talon"), ("gore"), ("horn"), ("slap"), ("slam"), ("arm"), ("arms"), ("sting"), ("tentacle");''')
  curs.execute('''CREATE INDEX index_dnd_weapon_name ON dnd_weapon(name);''')
  curs.execute('''CREATE TABLE monster_has_natural_weapon (
  monster_id INTEGER NOT NULL,
  weapon_id INTEGER NOT NULL,
  number_of_damage_dice tinyint(1) NOT NULL,
  damage_die_size tinyint(2) NOT NULL,
FOREIGN KEY(monster_id) REFERENCES dnd_monster(id),
FOREIGN KEY(weapon_id) REFERENCES dnd_weapon(id),
UNIQUE(monster_id, weapon_id)
  );''')
  curs.execute('''CREATE TABLE monster_deals_special_damage (
  monster_id INTEGER NOT NULL,
  damage TEXT NOT NULL,
FOREIGN KEY(monster_id) REFERENCES dnd_monster(id)
UNIQUE(monster_id, damage)
  );''')

  create_rulebook_table(curs)

  make_item_tables(curs)
  with open('diseases.sql') as diseaseFile:
    curs.executescript(diseaseFile.read())
  print('done making item tables')

  curs.execute('''CREATE TEMPORARY TABLE types_backup (id int, name varchar(32), slug varchar(32) );''')
  curs.execute('''INSERT INTO types_backup SELECT id, name, slug FROM dnd_monstertype;''')
  curs.execute('''DROP TABLE dnd_monstertype;''')
  curs.execute('''CREATE TABLE dnd_monstertype (
  id INTEGER PRIMARY KEY NOT NULL,
  name varchar(32) NOT NULL,
  slug varchar(32) NOT NULL,
  hit_die tinyint(2) NOT NULL,
  base_attack_per_4HD tinyint(1) NOT NULL,
  CR_per_12HD tinyint(1) NOT NULL,
  breathe tinyint(1) NOT NULL,
  eat tinyint(1) NOT NULL,
  sleep tinyint(1) NOT NULL,
UNIQUE(name),
UNIQUE(slug)
  );''') # add BAB as float, 1 or 0.75 or 0.5
  curs.execute('''CREATE TEMPORARY TABLE types_HD (name varchar(32), hit_die tinyint(2), base_attack_per_4HD tinyint(1), CR_per_12HD tinyint(1), breathe tinyint(1), eat tinyint(1), sleep tinyint(1) );''')
  curs.execute('''INSERT INTO types_HD (name, hit_die, base_attack_per_4HD, CR_per_12HD, breathe, eat, sleep) VALUES
                  ("Aberration", 8, 3, 3, 1, 1, 1), ("Animal", 8, 3, 4, 1, 1, 1), ("Construct", 10, 3, 3, 0, 0, 0), ("Dragon", 12, 4, 6, 1, 1, 1),
                  ("Elemental", 8, 3, 3, 0, 0, 0), ("Fey", 6, 2, 3, 1, 1, 1), ("Giant", 8, 3, 3, 1, 1, 1), ("Humanoid", 8, 3, 3, 1, 1, 1),
                  ("Magical Beast", 10, 4, 4, 1, 1, 1), ("Monstrous Humanoid", 8, 4, 4, 1, 1, 1), ("Ooze", 10, 3, 3, 1, 1, 0),
                  ("Outsider", 8, 4, 6, 1, 0, 0), ("Plant", 8, 3, 3, 1, 1, 0), ("Undead", 12, 2, 3, 0, 0, 0), ("Vermin", 8, 3, 3, 1, 1, 1),
                  ("Deathless", 8, 3, 3, 0, 0, 0);''')
  curs.execute('''INSERT INTO dnd_monstertype (name, slug, hit_die, base_attack_per_4HD, CR_per_12HD, breathe, eat, sleep) SELECT types_backup.name, slug, hit_die, base_attack_per_4HD, CR_per_12HD, breathe, eat, sleep FROM types_backup INNER JOIN types_HD ON types_backup.name=types_HD.name;''')
  curs.execute('''DROP TABLE types_backup;''')
  curs.execute('''DROP TABLE types_HD;''')
  # turns out original dnd_monstertype did not actually contain Animal at all
  curs.execute('''INSERT INTO dnd_monstertype (name,slug,hit_die,base_attack_per_4HD,CR_per_12HD,breathe,eat,sleep) VALUES (?,?,?,?,?,?,?,?);''', ('Animal','animal',8,3,4,1,1,1) )
  curs.execute('''CREATE INDEX index_monstertype_name ON dnd_monstertype(name);''')

  curs.execute('''CREATE TABLE monstertype_save_bonus (
  type_id INTEGER NOT NULL,
  fortitude_per_6HD tinyint(1) NOT NULL,
  reflex_per_6HD tinyint(1) NOT NULL,
  will_per_6HD tinyint(1) NOT NULL,
FOREIGN KEY(type_id) REFERENCES dnd_monstertype(id)
  );''')
  curs.execute('''CREATE TEMPORARY TABLE typename_save_bonus (
  type_name varchar(32) NOT NULL,
  fortitude_per_6HD tinyint(1) NOT NULL,
  reflex_per_6HD tinyint(1) NOT NULL,
  will_per_6HD tinyint(1) NOT NULL
  );''')
  curs.execute('''INSERT INTO typename_save_bonus (type_name, fortitude_per_6HD, reflex_per_6HD, will_per_6HD) VALUES
                  ("Aberration", 2, 2, 3), ("Animal", 3, 3, 2), ("Animal", 3, 3, 3), ("Construct", 2, 2, 2),
                  ("Dragon", 3, 3, 3), ("Elemental", 3, 2, 2), ("Elemental", 2, 3, 2), ("Fey", 2, 3, 3),
                  ("Giant", 3, 2, 2), ("Humanoid", 3, 2, 2), ("Humanoid", 2, 3, 2), ("Humanoid", 2, 2, 3), ("Monstrous Humanoid", 2, 3, 3),
                  ("Magical Beast", 3, 3, 2), ("Ooze", 2, 2, 2), ("Outsider", 3, 3, 3), ("Plant", 3, 2, 2),
                  ("Undead", 2, 2, 3), ("Vermin", 3, 2, 2), ("Deathless", 2, 2, 2);''')
  curs.execute('''INSERT INTO monstertype_save_bonus (type_id, fortitude_per_6HD, reflex_per_6HD, will_per_6HD) SELECT id, fortitude_per_6HD, reflex_per_6HD, will_per_6HD FROM dnd_monstertype INNER JOIN typename_save_bonus ON dnd_monstertype.name=typename_save_bonus.type_name;''')
  curs.execute('''DROP TABLE typename_save_bonus;''')

  curs.execute('''CREATE TEMPORARY TABLE subtypes_backup (id int, name varchar(32), slug varchar(32) );''')
  curs.execute('''INSERT INTO subtypes_backup SELECT id, name, slug FROM dnd_monstersubtype;''')
  curs.execute('''DROP TABLE dnd_monstersubtype;''')
  curs.execute('''CREATE TABLE dnd_monstersubtype (
  id INTEGER PRIMARY KEY NOT NULL,
  name varchar(32) NOT NULL,
  slug varchar(32) NOT NULL,
UNIQUE(name)
  );''')
  curs.execute('''INSERT INTO dnd_monstersubtype SELECT id, name, slug FROM subtypes_backup;''')
  curs.execute('''DROP TABLE subtypes_backup;''')
  curs.executemany('''INSERT INTO dnd_monstersubtype (name,slug) VALUES (?,?);''', [(name, name.lower() ) for name in (
    'Aquatic','Augmented','Living Construct','Cyborg',
    'Catfolk','Tayfolk','Mongrelfolk','Dwarf','Elf','Goblinoid','Gnoll','Gnome','Kenku','Human','Orc','Skulk','Maenad','Xeph','Darfellan','Hadozee',
    'Reptilian','Dragonblood','Psionic','Incarnum','Force','Void','Shapechanger',
    'Spirit','Dream','Tasloi','Swarm','Mob','Symbiont','Wretch')])
  curs.execute('''CREATE INDEX index_monstersubtype_name ON dnd_monstersubtype(name);''')
  
  curs.execute('''DROP TABLE dnd_racesize;''')
  curs.execute('''CREATE TABLE dnd_racesize (
  id INTEGER PRIMARY KEY NOT NULL
  ,name char(11) NOT NULL
  ,biped_carry_factor tinyint(1) NOT NULL
  ,quadruped_carry_factor tinyint(1) NOT NULL
  ,AC_size_modifier tinyint(1) NOT NULL
  ,grapple_size_modifier tinyint(2) NOT NULL
  );''') # 11 in case want to call it "Medium-size" rather than just "Medium"
  curs.execute('''INSERT INTO dnd_racesize (name, biped_carry_factor, quadruped_carry_factor, AC_size_modifier, grapple_size_modifier) VALUES
  ("Fine",1,2,8,-16), ("Diminutive",2,4,4,-12), ("Tiny",4,6,2,-8), ("Small",6,8,1,-4), ("Medium",8,12,0,0), ("Large",16,24,-1,4), ("Huge",32,48,-2,8), ("Gargantuan",64,96,-4,12), ("Colossal",128,192,-8,16);''')
  curs.execute('''CREATE INDEX index_racesize_name ON dnd_racesize(name);''')
  """ need to not have parentheses at top level:
  sqlite> insert into blanh values (3, 4, 5);
  Error: table blanh has 1 columns but 3 values were supplied
  sqlite> insert into blanh values ( (3), (4), (5) );
  Error: table blanh has 1 columns but 3 values were supplied
  sqlite> insert into blanh values (3), (4), (5);
  sqlite> insert into blanh values 3, 4, 5;
  Error: near "3": syntax error
  """
  curs.execute('''CREATE TABLE carrying_capacity (
  strength INTEGER PRIMARY KEY NOT NULL,
  fine_biped_max_load_ounces mediumint(3) NOT NULL
  );''')
  #medium_biped_max_load_pounds mediumint(3) NOT NULL,
  curs.execute('''INSERT INTO carrying_capacity (strength, fine_biped_max_load_ounces) VALUES (?,?), (?,?), (?,?), (?,?), (?,?), (?,?), (?,?), (?,?), (?,?), (?,?);''', tuple(itertools.chain(*[(i,10*i*2) for i in range(1, 11)]) ) )
  #print(tuple(itertools.chain(*[(11,115), (12,130), (13,150), (14,175), (15,200), (16,230), (17,260), (18,300), (19,350)]) ) )
  curs.execute('''INSERT INTO carrying_capacity (strength, fine_biped_max_load_ounces) VALUES (?,?), (?,?), (?,?), (?,?), (?,?), (?,?), (?,?), (?,?), (?,?);''', tuple(itertools.chain(*[(s,2*c) for (s,c) in [(11,115), (12,130), (13,150), (14,175), (15,200), (16,230), (17,260), (18,300), (19,350)] ]) ) )
  curs.executemany('''INSERT INTO carrying_capacity (strength, fine_biped_max_load_ounces) SELECT strength+10, fine_biped_max_load_ounces*4 FROM carrying_capacity WHERE strength>=10*?;''', [(i,) for i in range(1, 9)])

  #curs.execute('''CREATE TABLE dnd_law_chaos (
  #id tinyint(1) PRIMARY KEY NOT NULL,
  #description CHAR(7) NOT NULL
  #);''')
  #curs.execute('''INSERT INTO dnd_law_chaos (id,description) VALUES (?,?), (?,?), (?,?);''',
  #             (1, "Lawful", -1, "Chaotic", 0, "Neutral") )
  curs.execute('''CREATE TABLE monster_has_alignment (
  monster_id INTEGER NOT NULL,
  good_evil tinyint(1) NOT NULL,
  law_chaos char(1) NOT NULL,
FOREIGN KEY(monster_id) REFERENCES dnd_monster(id)
  );''')
  curs.execute('''CREATE TABLE dnd_plane (
  id INTEGER PRIMARY KEY NOT NULL,
  name TEXT NOT NULL,
  layer_number int(11) DEFAULT NULL,
  parent_plane INTEGER DEFAULT NULL,
  FOREIGN KEY(parent_plane) REFERENCES dnd_plane(id)
  );''')
  curs.executemany('''INSERT INTO dnd_plane (name) VALUES (?);''', [(name,) for name in (
  'Astral Plane',
  'Ethereal Plane',
 'Plane of Shadow',
 'Elemental Plane of Air',
 'Elemental Plane of Earth',
 'Elemental Plane of Fire',
 'Elemental Plane of Water',
 'Negative Energy Plane',
 'Positive Energy Plane',
 # Lower Planes:
 'Pandemonium',
 'The Abyss',
 'Carceri',
 'Hades',
 'Gehenna',
  'Baator',
  'Acheron',
 'Mechanus',
 # Upper Planes:
  'Arcadia',
 'Celestia',
 'Bytopia',
 'Elysium',
 'Beastlands',
  'Arborea',
 'Ysgard',
 'Limbo',
 'Outlands',
 'Yggdrasil',
 'Demiplane Prison of Tharizdun',
 'Demiplane of Ectoplasm',
 'Demiplane of Filth',
 'Demiplane of Nightmares',
 'Plane of Mirrors',
 'Plane of Radiance',
 'Utzou (Void)',
 'Far Realm',
  )])
  # sqlite> select dnd_plane.name,parent.name from dnd_plane left outer join dnd_plane as parent on dnd_plane.parent_plane=parent.id;
  curs.executemany('''INSERT INTO dnd_plane (name,layer_number,parent_plane) SELECT ?,?,id FROM dnd_plane where name=?;''', [pair for pair in (
  ('Dothion', None, 'Bytopia'),
  ('Thuldanin', 2, 'Acheron'),
  ('Tintibulus', 3, 'Acheron'),
  ('Arvandor', 1, 'Arborea'),
  ('Avernus', 1, 'Baator'),
  ('Dis', 2, 'Baator'),
  ('Minauros', 3, 'Baator'),
  ('Phlegethos', 4, 'Baator'),
  ('Stygia', 5, 'Baator'),
  ('Malbolge', 6, 'Baator'),
  ('Maladomini', 7, 'Baator'),
  ('Cania', 8, 'Baator'),
  ('Nessus', 9, 'Baator'),
  ('Cathrys', 2, 'Carceri'),
  ('Colothys', 4, 'Carceri'),
  ('Lunia', 1, 'Celestia'),
  ('Mercuria', 2, 'Celestia'),
  ('Venya', 3, 'Celestia'),
  ('Solania', 4, 'Celestia'),
  ('Mertion', 5, 'Celestia'),
  ('Jovar', 6, 'Celestia'),
  ('Chronias', 7, 'Celestia'),
  ('Amoria', 1, 'Elysium'),
  ('Eronia', 2, 'Elysium'),
  ('Khalas', 1, 'Gehenna'),
  ('Torremor', 503, 'The Abyss'),
  ('Melantholep', 518, 'The Abyss'),
  ('Androlynne', 471, 'The Abyss'),
  ('Azzagrat', 45, 'The Abyss'), # Everything on the gloomy 45th layer of the Abyss is somehow doused or subdued
  ('Azzagrat', 46, 'The Abyss'), # Sunlight rises up from the ground of the 46th layer
  ('Azzagrat', 47, 'The Abyss'), # A wan cerulean sun feebly illuminates the dark sky of the 47th layer, where cold is hot and hot is cold.
  ('Belistor', 277, 'The Abyss'),
  ('The Gaping Maw', 88, 'The Abyss'),
  ('The Demonweb Pits', 66, 'The Abyss'),
  ('The Endless Maze', 600, 'The Abyss'),
  ('The Demonweb', 66, 'The Abyss'),
  ("Hollow's Heart", 176, 'The Abyss'),
  ('Hungry Tarns', 380, 'The Abyss'),
  ('Iron Wastes', 23, 'The Abyss'),
  ('Plain of Infinite Portals', 1, 'The Abyss'),
  ('Shaddonon', 49, 'The Abyss'),
  ('The Shadowsea', 89, 'The Abyss'),
  ('Shedaklah', 222, 'The Abyss'),
  ('Shendilavri', 570, 'The Abyss'),
  ('Thanatos', 113, 'The Abyss'),
  ('The Writhing Realm of Ugudenk', 177, 'The Abyss'),
  ("Yeenoghu's Realm", 422, 'The Abyss'),
  ("Zionyn", 663, 'The Abyss'),
  )])

  (
 'The Abyss (Brine Flats)', # Gaping Maw #88
 'The Abyss (Flesh Mountains)', # The Flesh Mountains, Dalmosh's Abyssal home, stretch across several layers of the Abyss.
 'The Abyss (Ice Wastes)', # probably Iron Wastes
 'The Abyss (Screaming Jungle)', # Gaping Maw #88
 'The Abyss (Stygia)', # should be Baator
 'Celestia (Empyrea)', # Empyrea, the City of Tempered Souls, on Celestia's fifth layer Mertion
 'Mechanus (Regulus)', # modron city
 r"",
 'The Abyss, Carceri, Outlands, or Pandemonium',
 'Elemental Planes of Air, Fire',
 'Elemental Planes of Air, Water',
 'Elemental Planes of Earth, Fire',
 'Elemental Planes of Earth, Water',
 'Any',
 'Any ',
 'Any Elemental Plane',
 'Any Outer Plane',
 'Any aquatic',
 'Any cold',
 'Any desert',
 'Any forest',
 'Any hills',
 'Any land',
 'Any land and underground',
 'Any marsh',
 'Any mountain',
 'Any plains',
 'Any tainted',
 'Any temperate',
 'Any warm',
 'Any warm land',
 # land probably means anything but aquatic or underground
 "Astral Plane, N'gati", # the N'gati, an ancient slumbering construct adrift in the Astral Plane
 'Atropus',
 'Carceri, Gehenna, or Hades',
 'Cold aquatic',
 'Cold deserts',
 'Cold forests',
 'Cold hills',
 'Cold marsh',
 'Cold marshes',
 'Cold mountains',
 'Cold plains',
 'Crystalline Prison', # on the Material Plane
 'Dreamscapes, Ethereal Plane', # Towers of High Sorcery
 'Forest or aquatic',
 'Forest or plains',
 'Forests and marshes',
 'Frog Swamp',
 'Inner Planes',
 'Inner Planes, aquatic',
 'Inner Planes, mountains',
 'Inner Planes, underground',
 'Lower Planes',
 'Marsh or underground',
 'Mountains and deserts',
 'Nautiloid vessel',
 'Outlands (Athenaeum Nefarious)',
 'Outlands (Sigil)',
 'Temperate aquatic',
 'Temperate deserts',
 'Temperate forests',
 'Temperate hills',
 'Temperate marshes',
 'Temperate mountains',
 'Temperate plains',
 'Underground',
 'Unknown Plane', # Keeper Fiend Folio
 'Upper Planes',
 'Urban',
 'Urban or underground',
 'Warm aquatic',
 'Warm desert',
 'Warm deserts',
 'Warm forests',
 'Warm hills',
 'Warm marshes',
 'Warm mountains',
 'Warm plains',
 'Ziggurat', # Zargon on Material
  )
  curs.execute('''CREATE INDEX index_dnd_plane_name ON dnd_plane (name);''')
  curs.execute('''CREATE TABLE monster_on_plane (
  monster_id INTEGER NOT NULL,
  plane_id INTEGER NOT NULL,
FOREIGN KEY(monster_id) REFERENCES dnd_monster(id),
FOREIGN KEY(plane_id) REFERENCES dnd_plane(id)
  );''')
  # sqlite> select dnd_monster.name,dnd_monstertype.name,dnd_plane.name from monster_plane inner join dnd_monster on monster_id=dnd_monster.id inner join dnd_plane on plane_id=dnd_plane.id inner join dnd_monstertype on dnd_monster.type_id=dnd_monstertype.id where dnd_monstertype.name="Animal" order by dnd_monstertype.name,dnd_plane.id;
  # picks up a lot that weren't noted by sqlite> select dnd_monster.name from dnd_monster inner join monster_subtype on dnd_monster.id=monster_id inner join dnd_monstertype on dnd_monster.type_id=dnd_monstertype.id inner join dnd_monstersubtype on subtype_id=dnd_monstersubtype.id where dnd_monstertype.name="Animal" and dnd_monstersubtype.name="Extraplanar";? no, those were all errors
  curs.execute('''CREATE TABLE dnd_terrain (
  id INTEGER PRIMARY KEY,
  name char(10)
  );''')
  curs.executemany('''INSERT INTO dnd_terrain(name) VALUES (?)''', [(name,) for name in TERRAIN_NAMES] )
  # A swamp is a place where the plants that make up the area covered in water are primarily woody plants or trees. Woody plants would be mangroves or cypress trees. A marsh, on the other hand, is defined as having no woody plants. The non-woody plants would be saltmarsh grasses, reeds, or sedges.
  curs.execute('''CREATE TABLE monster_on_terrain (
  monster_id INTEGER NOT NULL,
  terrain_id INTEGER NOT NULL,
FOREIGN KEY(monster_id) REFERENCES dnd_monster(id),
FOREIGN KEY(terrain_id) REFERENCES dnd_terrain(id)
  );''')
  curs.execute('''CREATE TABLE monster_in_climate (
  monster_id INTEGER NOT NULL,
  climate tinyint(1) NOT NULL,
FOREIGN KEY(monster_id) REFERENCES dnd_monster(id)
  );''')

  # could create dnd_bloodline table with feat_id...but monster table doesn't have feats
  # dnd_monsterhasfeat has been garbled by destroying and recreating dnd_monster: sqlite> select dnd_monster.name,dnd_feat.name from dnd_monsterhasfeat inner join dnd_monster on monster_id=dnd_monster.id inner join dnd_feat on feat_id=dnd_feat.id;

  curs.execute('''CREATE TABLE dnd_special_ability (
  id INTEGER PRIMARY KEY NOT NULL,
  special_attack bool NOT NULL DEFAULT 0,
  name varchar(64) NOT NULL,
  type nchar(1) DEFAULT NULL,
  description longtext DEFAULT NULL
  );''')
  # simplest to use X for Ex, U for Su, P for Sp
  curs.execute('''INSERT INTO dnd_special_ability (name, description) VALUES (?, ?)''',
               ('Cold Immunity', "A creature with cold immunity never takes cold damage. It has vulnerability to fire, which means it takes half again as much (+50%) damage as normal from fire, regardless of whether a saving throw is allowed, or if the save is a success or failure.") )
  curs.execute('''INSERT INTO dnd_special_ability (name, type, description) VALUES (?, ?, ?)''',
               ('Evasion', 'X', "If subjected to an attack that allows a Reflex save for half damage, a character with evasion takes no damage on a successful save. As with a Reflex save for any creature, a character must have room to move in order to evade. A bound character or one squeezing through an area cannot use evasion. As with a Reflex save for any creature, evasion is a reflexive ability. The character need not know that the attack is coming to use evasion.") )
  curs.execute('''CREATE INDEX index_special_ability_name ON dnd_special_ability(name);''')
  #curs.execute('''INSERT INTO dnd_rulebook (id, dnd_edition_id, name, abbr, description, year, official_url, slug, image, published) VALUES (116, 7, "Monster Manual 2", "MM2", "", 1983, "", "monster-manual-ii", NULL, NULL);''')
  # see what it actually is with select * from dnd_rulebook where id=45;
  maxNameLen = len('Olhydra (Princess of Evil Water Creatures, Princess of Watery Evil, Mistress of the Black Tide)')
  #curs.execute('''DROP TABLE dnd_monsters;''')
  curs.execute('''DROP TABLE dnd_monster;''')
  # the monster table could have some page numbers from http://archive.wizards.com/default.asp?x=dnd/arch/lists
  curs.execute('''CREATE TABLE dnd_monster (
  id INTEGER PRIMARY KEY NOT NULL,
  rulebook_id INTEGER DEFAULT NULL,
  name varchar({}) NOT NULL,
  size_id tinyint(1) NOT NULL,
  type_id tinyint(2) NOT NULL,
  hit_dice smallint(2) NOT NULL,
  land_speed smallint(2) DEFAULT NULL,
  natural_armor_bonus tinyint(2) DEFAULT NULL,
  strength tinyint(2) DEFAULT NULL,
  dexterity tinyint(2) DEFAULT NULL,
  constitution tinyint(2) DEFAULT NULL,
  intelligence tinyint(2) DEFAULT NULL,
  wisdom tinyint(2) NOT NULL,
  charisma tinyint(2) NOT NULL,
  challenge_rating tinyint(2) NOT NULL,
  level_adjustment tinyint(2) DEFAULT NULL,
FOREIGN KEY(rulebook_id) REFERENCES dnd_rulebook(id),
FOREIGN KEY(size_id) REFERENCES dnd_racesize(id)
 );'''.format(maxNameLen) )
#  law_chaos_id tinyint(1) NOT NULL,
#FOREIGN KEY(law_chaos_id) REFERENCES dnd_law_chaos(id)
  # rulebook_id should eventually be NOT NULL,
  # but will need to add the monster books to dnd_rulebook,
  # which requires dnd_edition_id and stuff
  # this pegs Dragon #309 as the start of 3.5ed http://www.enworld.org/forum/showthread.php?9651-DRAGON-Magazine-monster-index!
  # Dragon #309 (War, Incursion) from July 2003 was the first D&D 3.5 issue. https://rpg.stackexchange.com/questions/14892/in-what-issue-did-dragon-dungeon-magazine-transition-to-3-5e-rules
  curs.execute('''CREATE TABLE monster_has_subtype (
  subtype_id INTEGER NOT NULL,
  monster_id INTEGER NOT NULL,
  FOREIGN KEY(monster_id) REFERENCES dnd_monster(id),
  FOREIGN KEY(subtype_id) REFERENCES dnd_monster(id)
  );''')
  curs.execute('''CREATE TABLE monster_has_special_ability (
  monster_id INTEGER NOT NULL,
  special_ability_id INTEGER NOT NULL,
FOREIGN KEY(monster_id) REFERENCES dnd_monster(id),
FOREIGN KEY(special_ability_id) REFERENCES dnd_special_ability(id)
  );''')
  curs.execute('''CREATE TABLE monster_has_spell_like_ability (
  monster_id INTEGER NOT NULL,
  spell_id INTEGER NOT NULL,
  caster_level tinyint(2) NOT NULL,
  caster_level_scales_with_HD bool NOT NULL,
  uses_per_day tinyint(1),
  parenthetical VARCHAR(32) DEFAULT NULL,
FOREIGN KEY(monster_id) REFERENCES dnd_monster(id),
FOREIGN KEY(spell_id) REFERENCES dnd_spell(id)
  );''')
  # It's rare, but some monsters do have SLAs of different caster levels.
  curs.execute('''CREATE TABLE monster_casts_spells (
  monster_id INTEGER NOT NULL,
  character_class_id INTEGER NOT NULL,
  level tinyint(2) NOT NULL,
FOREIGN KEY(monster_id) REFERENCES dnd_monster(id),
FOREIGN KEY(character_class_id) REFERENCES dnd_characterclass(id)
  );''')
  curs.execute('''CREATE TEMPORARY TABLE dnd_characterclass_with_spells (
  id INTEGER PRIMARY KEY,
  name TEXT NOT NULL,
FOREIGN KEY(id) REFERENCES dnd_characterclass(id)
  );''')
  curs.execute('''INSERT INTO dnd_characterclass_with_spells (id, name) SELECT DISTINCT dnd_characterclass.id, dnd_characterclass.name FROM dnd_characterclass INNER JOIN dnd_spellclasslevel ON dnd_spellclasslevel.character_class_id=dnd_characterclass.id;''')
  curs.execute('''CREATE TABLE damage_reduction (
  id INTEGER PRIMARY KEY,
  bypass CHAR(64) NOT NULL,
UNIQUE(bypass)
  );''')
  curs.execute('''CREATE INDEX index_damage_reduction ON damage_reduction(bypass);''')
  curs.execute('''CREATE TABLE monster_has_damage_reduction (
  monster_id INTEGER NOT NULL,
  reduction tinyint(2) NOT NULL,
  bypass_id INTEGER DEFAULT NULL,
FOREIGN KEY(monster_id) REFERENCES dnd_monster(id),
FOREIGN KEY(bypass_id) REFERENCES damage_reduction(id)
  );''')
  curs.execute('''CREATE TABLE monster_has_fast_healing (
  monster_id INTEGER NOT NULL,
  healing tinyint(2) NOT NULL,
FOREIGN KEY(monster_id) REFERENCES dnd_monster(id)
  );''')
  curs.execute('''CREATE TABLE monster_has_spell_resistance (
  monster_id INTEGER NOT NULL,
  resistance tinyint(2) NOT NULL,
FOREIGN KEY(monster_id) REFERENCES dnd_monster(id)
  );''')
  print('about to insert psionic powers')
  insert_psionic_powers(curs)
  print('done inserting psionic powers')

  curs.execute('''CREATE TABLE min_level_to_cast_spell (
  character_class_id INTEGER,
  spell_level tinyint(1) NOT NULL,
  class_level tinyint(2) NOT NULL,
PRIMARY KEY (character_class_id, spell_level)
  );''')
  for character_class_name in ("Cleric", "Druid", "Wizard"):
    character_class_id = id_from_name(curs, 'dnd_characterclass', character_class_name)
    curs.executemany('''INSERT INTO min_level_to_cast_spell (character_class_id, spell_level, class_level) VALUES (?, ?, ?)''',
                     [(character_class_id, spell_level, 2*spell_level - 1) for spell_level in range(1, 10)])
  for character_class_name in ("Sorcerer",):
    character_class_id = id_from_name(curs, 'dnd_characterclass', character_class_name)
    curs.executemany('''INSERT INTO min_level_to_cast_spell (character_class_id, spell_level, class_level) VALUES (?, ?, ?)''',
                     [(character_class_id, spell_level, 2*spell_level) for spell_level in range(2, 10)])
  for character_class_name in ("Paladin", "Ranger"):
    character_class_id = id_from_name(curs, 'dnd_characterclass', character_class_name)
    curs.executemany('''INSERT INTO min_level_to_cast_spell (character_class_id, spell_level, class_level) VALUES (?, ?, ?)''',
                     [(character_class_id, spell_level, 4*spell_level) for spell_level in range(1, 5)])
  for character_class_name in ("Bard",):
    character_class_id = id_from_name(curs, 'dnd_characterclass', character_class_name)
    curs.executemany('''INSERT INTO min_level_to_cast_spell (character_class_id, spell_level, class_level) VALUES (?, ?, ?)''',
                     [(character_class_id, spell_level, 3*spell_level - 2) for spell_level in range(2, 7)])
  for character_class_name in ("Adept",):
    character_class_id = id_from_name(curs, 'dnd_characterclass', character_class_name)
    curs.executemany('''INSERT INTO min_level_to_cast_spell (character_class_id, spell_level, class_level) VALUES (?, ?, ?)''',
                     [(character_class_id, spell_level, 4*spell_level - 4) for spell_level in range(1, 6)])
  for character_class_name in ("Cleric", "Druid", "Wizard", "Bard"):
    character_class_id = id_from_name(curs, 'dnd_characterclass', character_class_name)
    curs.execute('''INSERT INTO min_level_to_cast_spell (character_class_id, spell_level, class_level) VALUES (?, ?, ?)''',
                 (character_class_id, 0, 1))
  curs.execute('''INSERT INTO min_level_to_cast_spell (character_class_id, spell_level, class_level) VALUES (?, ?, ?)''',
               (id_from_name(curs, 'dnd_characterclass', "Bard"), 1, 2))
  curs.execute('''INSERT INTO min_level_to_cast_spell (character_class_id, spell_level, class_level) VALUES (?, ?, ?)''',
               (id_from_name(curs, 'dnd_characterclass', "Sorcerer"), 1, 1))
  """
  sqlite> SELECT dnd_characterclass.name, spell_level, class_level FROM min_level_to_cast_spell INNER JOIN dnd_characterclass ON character_class_id=dnd_characterclass.id;
  """

  curs.execute('''CREATE TABLE spell_brings_monster (
  spell_id INTEGER NOT NULL,
  monster_id INTEGER NOT NULL,
FOREIGN KEY(spell_id) REFERENCES dnd_spell(id),
FOREIGN KEY(monster_id) REFERENCES dnd_monster(id)
  );''')
  # sqlite> INSERT INTO spell_brings_monster (spell_id,monster_id) SELECT dnd_spell.id,dnd_monster.id FROM dnd_monster INNER JOIN dnd_spell ON dnd_spell.name="Summon Monster I" AND dnd_monster.name="Elysian Thrush";
  # sqlite> select * from spell_brings_monster;
  # 2344|2441
  # manually did the command below in python3, no result, then tried in sqlite3 again:
  # sqlite> INSERT INTO spell_brings_monster (spell_id,monster_id) SELECT dnd_spell.id,dnd_monster.id FROM dnd_monster INNER JOIN dnd_spell ON dnd_spell.name="Summon Monster II" AND dnd_monster.name="Clockwork Mender";
  # sqlite> select * from spell_brings_monster;
  # 2344|2441
  # 1041|2442
  curs.executemany('''INSERT INTO spell_brings_monster (spell_id,monster_id) SELECT dnd_spell.id,dnd_monster.id FROM dnd_spell INNER JOIN dnd_monster ON dnd_spell.name=? AND dnd_monster.name=?;''', [ ('Summon Monster I', 'Elysian Thrush'), ('Summon Monster II','Devil, Lemure'), ('Summon Monster II', 'Clockwork Mender'), ('Summon Monster II', 'Fetid Fungus'), ('Summon Monster II', 'Nerra, Varoot'), ('Summon Monster II', 'Kaorti'), ('Summon Monster II', 'Howler Wasp'), ('Summon Monster II', "Ur'Epona")
  ] )
  
  # http://www.imarvintpa.com/dndlive/Index_Deities.php
  # have to follow links, but has deities with alignments and domains
  
  # SQLite database only has 165 items
  # http://www.imarvintpa.com/dndlive/Index_It_Fam.php
  # http://www.imarvintpa.com/dndlive/Index_It_Cat.php
  # http://www.imarvintpa.com/dndlive/Index_It_Sub.php
  # http://www.imarvintpa.com/dndlive/Index_It_Slot.php
  # dump to plaintext first?
  # http://www.giantitp.com/forums/showthread.php?148101-3-x-Shax-s-Indispensible-Haversack-(Equipment-Handbook)
  #  has some things Marvin doesn't like Angel Radiance http://www.imarvintpa.com/dndlive/Index_It_Source.php?Sources=Book+of+Exalted+Deeds
  
  # skill ranks! http://www.imarvintpa.com/dndlive/index.php
  # the MapTool Token link on each monster leads to XML
  # http://www.imarvintpa.com/dndlive/TokMonsters.php?ID=1
  
  print('making template tables')
  curs.execute('''CREATE TABLE dnd_template (
  id INTEGER PRIMARY KEY NOT NULL,
  name varchar({}) NOT NULL,
  rulebook_id INTEGER NOT NULL,
  page smallint(2) DEFAULT NULL,
  level_adjustment tinyint(1) DEFAULT NULL,
FOREIGN KEY(rulebook_id) REFERENCES dnd_rulebook(id)
  );'''.format(maxNameLen) )
  # webcrawl for stats? http://www.realmshelps.net/monsters/templates/index.shtml
  # Restless Prey http://archive.wizards.com/default.asp?x=dnd/fw/20030531a
  # http://www.giantitp.com/forums/showthread.php?444021-Master-Template-List-(from-the-WotC-Forum)
  curs.execute('''CREATE TABLE template_type (
  template_id INTEGER NOT NULL,
  base_type INTEGER NOT NULL,
  output_type INTEGER NOT NULL,
FOREIGN KEY(base_type) REFERENCES dnd_monstertype(id),
FOREIGN KEY(output_type) REFERENCES dnd_monstertype(id),
FOREIGN KEY(template_id) REFERENCES dnd_template(id)
  );''')
  # Savage Species type pyramid was void as soon as the 3.5 Monster Manual came out with rules which specified type changes that happen contrary to the type pyramid. For instance, Half-Dragon applied to an Aasimar changes it to Dragon type because that's what Monster Manual (the primary rule source for templates) says happens (page 146). The type pyramid says that type change doesn't occur, which is a disagreement. In the case of disagreements like this, the Primary Sources Errata Rule says Monster Manual is correct and so that Savage Species rule is no good.
  #for template in Template.findall(open('rockdeworldTemplates.txt').read() ):
  for template in Template.read_file('rockdeworldTemplates.txt'):
    template.insert_into(curs)
  print('done with templates')

  Monster.from_statblock("Ferret: CR 1/10; Diminutive animal; HD 1/4d8; hp 1; Init +2; Spd 15 ft., climb 15 ft.; AC 17, touch 16, flat-footed 15; Base Atk +0; Grp -16; Atk +6 melee (1d2-4, bite); Full Atk +6 melee (1d2-4, bite); Space/Reach 1 ft./0 ft.; SA attach; SQ scent; AL N; SV Fort +2, Ref +4, Will +1; Str 3, Dex 15, Con 10, Int 2, Wis 12, Cha 5. Skills and Feats: Balance +10, Climb +11, Hide +13, Move Silently +9, Spot +14; Weapon Finesse. Attach (Ex): On a hit with its bite attack, it automatically deals bite damage each round (AC 15 when attached).", 'DMG').insert_into(curs)
  Monster.from_statblock("Hedgehog: CR 1/10; Diminutive animal; HD 1/4d8; hp 1; Init +0; Spd 15 ft.; AC 17, touch 15, flat-footed 16; Base Atk +0; Grp -16; Atk +5 melee (1d3-4 bite); Full Atk +5 melee (1d3-4 bite); Space/Reach 1 ft./0 ft.; SA poison; SQ defensive ball; AL N; SV Fort +2, Ref +3, Will +1; Str 3, Dex 12, Con 10, Int 2, Wis 12, Cha 5. Skills and Feats: Hide +17, Listen +5, Spot +5; Weapon Finesse. Poison (Ex): When in a defensive ball (see below), spines poison foes touching the hedgehog; injury, Fortitude DC 10, initial and secondary damage 1d2 Dex. Defensive Ball (Ex): Rolls into a ball as a standard action, granting a +2 circumstance bonus on saves and AC. Unrolling is a free action.", 'DMG').insert_into(curs)
  Monster.from_statblock("Mouse: CR 1/10; Fine animal; HD 1/4d8; hp 1; Init +0; Spd 10 ft., climb 10 ft.; AC 19, touch 18, flat-footed 19; Base Atk +0; Grp -21; Atk -; Full Atk -; Space/Reach 1/2 ft./0 ft.; SA -; SQ scent; AL N; SV Fort +2, Ref +2, Will +1; Str 1, Dex 11, Con 10, Int 2, Wis 12, Cha 2. Skills and Feats: Balance +8, Climb +10, Hide +20, Move Silently +12; feat.", 'DMG').insert_into(curs)
  Monster.from_statblock("Screech Owl: CR 1/10; Diminutive animal; HD 1/4d8; hp 1; Init +3; Spd 10 ft., fly 30 ft. (average); AC 18, touch 17, flat-footed 15; Base Atk +0; Grp -15; Atk +7 melee (1d2-3, talons); Full Atk +7 melee (1d2-3, talons); Space/Reach 1 ft./0 ft.; SA -; SQ -; AL N; SV Fort +2, Ref +5, Will +2; Str 4, Dex 17, Con 10, Int 2, Wis 14, Cha 4. Skills and Feats: Listen +14, Move Silently +20, Spot +8; Weapon Finesse.", 'DMG').insert_into(curs)
  Monster.from_statblock("Thrush: CR 1/10; Diminutive animal; HD 1/4d8; hp 1; Init +2; Spd 10 ft., fly 40 ft. (average); AC 17, touch 16, flat-footed 15; Base Atk +0; Grp -17; Atk -; Full Atk -; Space/Reach 1 ft./0 ft.; SA -; SQ -; AL N; SV Fort +2, Ref +4, Will +2; Str 1, Dex 15, Con 10, Int 2, Wis 14, Cha 6. Skills and Feats: Listen +8, Spot +8; Alertness.", 'DMG').insert_into(curs)
  Monster.from_statblock("Owl, Medium: CR 1; Medium animal; HD 2d8; hp 13; Init +1; Spd 10 ft., fly 60 ft. (average); AC 14, touch 11, flat-footed 13; Base Atk +1; Grp +3; Atk +2 melee (1d4+2, talons); Full Atk +2 melee (1d4+2, talons) and +0 melee (1d6+1, bite); Space/Reach 5 ft./5 ft.; SA -; SQ -; AL N; SV Fort +4, Ref +4, Will +2; Str 14, Dex 13, Con 12, Int 2, Wis 14, Cha 4. Skills and Feats: Listen +14, Move Silently +19, Spot +14; Multiattack.", 'DMG').insert_into(curs)
  Monster.from_statblock("Raven, Small: CR 1/2; Small animal; HD 1d8; hp 5; Init +1; Spd 10 ft., fly 40 ft. (average); AC 13, touch 12, flat-footed 12; Base Atk +0; Grp -7; Atk +2 melee (1d3-3, talons); Full Atk +2 melee (1d3-3, talons); Space/Reach 5 ft./5 ft.; SA -; SQ -; AL N; SV Fort +3, Ref +3, Will +2; Str 5, Dex 13, Con 12, Int 2, Wis 14, Cha 6. Skills and Feats: Listen +6, Spot +6; Weapon Finesse.", 'DMG').insert_into(curs)

  print('about to start main monster loop')
  # I'm guessing dropwhile has no overhead after failing
  for i,row in itertools.dropwhile(lambda p: p[0]==0,
               enumerate(ODE.get_rows() ) ):
    #if i < 4300: continue
    if i % 1000 == 0: print('processing {}th row:'.format(i), row) # since row 0 was the headings, this really is the ith row
    monster = Monster(row)
    monster.insert_into(curs)
    #if i > 8000: break
  
  make_familiar_table(curs)

  conn.commit()
  make_skill_table(conn, curs)

  if False: parse_IMarvinTPA(curs, marvinCache)
  
  conn.commit()
  conn.close()
  
  for i,s in enumerate(allUnknownTerms):
    print('\n'.join(sorted(s) ), file=open('names' + str(i + 1) + '.txt', 'w') )
  for i,counter in enumerate(appearInNamesMoreThanOnce):
    print('\n'.join(str(blah) + str(s) for blah,s in counter.items() if len(s) > 1), file=open('counts' + str(i) + '.txt', 'w') )


def command_line_usage():
  parser = argparse.ArgumentParser(description='Incorporate an XLS file of monster data into a SQLite database.')
  parser.add_argument('XLSpath', metavar='XLS table',
                      nargs='?', default="Monster Compendium.xls",
                      help='path to the XLS table')
  parser.add_argument('DBpath', metavar='underlying database location',
                      nargs='?', default='dnd.sqlite',
                      help='path to the SQLite database')
  parser.add_argument('--marvin-cache', metavar='IMarvinTPA cache location', help='directory to cache IMarvinTPA skill rank data')
  args = parser.parse_args()
  if not os.path.exists(args.XLSpath):
    raise OSError("The XLS table to translate was not found at {}; try python monsters.py --help for usage.".format(args.XLSpath) )
  if not os.path.exists(args.DBpath):
    raise OSError("The underlying database was not found at {}; try python monsters.py --help for usage.".format(args.DBpath) )
  profile = cProfile.Profile()
  profile.enable()
  print('args =', args)
  create_database(args.XLSpath, DBpath=args.DBpath, marvinCache=args.marvin_cache)
  profile.disable()
  with open('cProfile.txt', 'w') as statsFile:
      stats = pstats.Stats(profile, stream=statsFile)
      stats.strip_dirs().sort_stats('cumtime').print_stats()
  #print(Monster.allEnvs)

if __name__ == "__main__":
  command_line_usage()
  

